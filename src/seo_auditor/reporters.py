# Importa JSON para exportar datos técnicos trazables.
import json

# Importa expresiones regulares para limpiar narrativa IA.
import re

# Importa contador para cálculos agregados.
from collections import Counter

# Importa la clase Path para gestionar rutas de forma robusta.
from pathlib import Path

# Importa tipos para estructuras semánticas del informe.
from typing import Any

# Importa parseo de URL para normalizar cruces entre fuentes.
from urllib.parse import urlparse

# Importa escape para sanear texto potencialmente interpretado como etiquetas XML.
from xml.sax.saxutils import escape

# Importa objetos para construir estilos de Word.
from docx import Document

# Importa utilidades de estilo de Word.
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Importa tamaño de fuente y color de Word.
from docx.shared import Pt, RGBColor

# Importa utilidades de PDF para generar un informe portable.
from reportlab.lib import colors

# Importa tamaño de página del PDF.
from reportlab.lib.pagesizes import A4

# Importa estilos tipográficos para PDF.
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

# Importa utilidades avanzadas de maquetación PDF.
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# Importa utilidades de openpyxl para Excel profesional.
from openpyxl import Workbook

# Importa gráficos de Excel.
from openpyxl.chart import BarChart, PieChart, Reference

# Importa estilos y colores de openpyxl.
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# Importa validaciones de datos para celdas editables.
from openpyxl.worksheet.datavalidation import DataValidation

# Importa herramientas de tabla de openpyxl.
from openpyxl.worksheet.table import Table as ExcelTable, TableStyleInfo

# Importa los modelos del dominio del proyecto.
from seo_auditor.models import ResultadoAuditoria


# Define el orden de severidad para visualización homogénea.
ORDEN_SEVERIDAD = ["crítica", "alta", "media", "baja", "informativa"]

# Define jerarquía fija obligatoria del informe final.
JERARQUIA_INFORME = [
    "Resumen ejecutivo",
    "KPIs principales",
    "Visibilidad orgánica real",
    "Comportamiento y conversión",
    "Páginas prioritarias",
    "Oportunidades SEO prioritarias",
    "Cruce auditoría técnica + Search Console",
    "Keyword / query mapping inicial",
    "Hallazgos críticos",
    "Quick wins",
    "Acciones técnicas",
    "Acciones de contenido",
    "Rendimiento y experiencia de usuario",
    "Indexación y rastreo",
    "Gestión de indexación",
    "Roadmap",
]

# Define secciones dependientes de Search Console para ocultación condicional.
SECCIONES_GSC = {
    "Visibilidad orgánica real",
    "Oportunidades SEO prioritarias",
    "Cruce auditoría técnica + Search Console",
    "Keyword / query mapping inicial",
}

# Define umbral mínimo de impresiones para detectar oportunidad de CTR.
OPORTUNIDAD_GSC_MIN_IMPRESIONES_CTR = 100.0

# Define umbral máximo de CTR para marcar bajo rendimiento de snippet.
OPORTUNIDAD_GSC_MAX_CTR = 0.02

# Define límite inferior de posición para oportunidades de ascenso.
OPORTUNIDAD_GSC_MIN_POSICION = 4.0

# Define límite superior de posición para oportunidades de ascenso.
OPORTUNIDAD_GSC_MAX_POSICION = 15.0

# Define mínimo de impresiones para oportunidad por posición.
OPORTUNIDAD_GSC_MIN_IMPRESIONES_POSICION = 50.0

# Define mínimo de impresiones para cruce de visibilidad y on-page.
OPORTUNIDAD_GSC_MIN_IMPRESIONES_ONPAGE = 50.0

# Define umbral de impresiones para considerar impacto alto.
OPORTUNIDAD_GSC_IMPRESIONES_ALTO_IMPACTO = 500.0

# Define paleta de colores reutilizable para dashboard Excel.
COLOR_DASHBOARD_TITULO = "1F4E78"
COLOR_KPI_TITULO_FONDO = "E6EEF8"
COLOR_KPI_VALOR_FONDO = "F8FBFF"
COLOR_BLOQUE_VISIBILIDAD = "1D4ED8"
COLOR_BLOQUE_INDEXACION = "0F766E"
COLOR_BLOQUE_INCIDENCIAS = "B91C1C"
COLOR_BLOQUE_OPORTUNIDADES = "7C3AED"
COLOR_BLOQUE_SCORE = "0B7285"

# Define total de márgenes horizontales del PDF en puntos (36+36).
PDF_HORIZONTAL_MARGIN_POINTS = 72.0

# Define umbrales de cruce GSC+Analytics para insights accionables.
CRUCE_GA_MIN_SESIONES_REBOTE = 80.0
CRUCE_GA_MIN_REBOTE_ALTO = 0.65
CRUCE_GSC_MIN_IMPRESIONES_SIN_ENGAGEMENT = 500.0
CRUCE_GA_MAX_SESIONES_SIN_ENGAGEMENT = 30.0
CRUCE_GSC_MAX_POSICION_POTENCIAL = 12.0
CRUCE_GA_MIN_SESIONES_SIN_CONVERSION = 40.0
CRUCE_GA_MAX_CONVERSIONES = 0.0

# Define umbrales y pesos de priorización para páginas estratégicas.
PRIORIDAD_GSC_MIN_IMPRESIONES_CTR = 500.0
PRIORIDAD_GSC_MAX_CTR = 0.02
PRIORIDAD_SCORE_CTR_BAJO = 40.0
PRIORIDAD_GA_MIN_SESIONES_SIN_CONVERSION = 80.0
PRIORIDAD_GA_MAX_CONVERSIONES = 0.0
PRIORIDAD_SCORE_SIN_CONVERSION = 35.0
PRIORIDAD_GSC_MIN_IMPRESIONES_ENGAGEMENT_BAJO = 400.0
PRIORIDAD_GA_MIN_REBOTE_BAJO_ENGAGEMENT = 0.65
PRIORIDAD_SCORE_ENGAGEMENT_BAJO = 25.0
PRIORIDAD_GSC_MIN_POSICION = 4.0
PRIORIDAD_GSC_MAX_POSICION = 15.0
PRIORIDAD_SCORE_POSICION_POTENCIAL = 15.0
PRIORIDAD_SCORE_MAX_HALLAZGOS = 20.0
PRIORIDAD_SCORE_HALLAZGO = 5.0

# Define umbrales para quick wins basados en Analytics.
QUICK_WIN_GA_MIN_SESIONES_REBOTE = 100.0
QUICK_WIN_GA_MIN_REBOTE = 0.65
QUICK_WIN_GA_MIN_SESIONES_CONVERSION = 80.0
QUICK_WIN_GA_MAX_CONVERSIONES = 0.0
QUICK_WINS_BLOQUE_MAXIMO = 10

# Define límites de muestras narrativas para bloques Analytics.
ANALYTICS_TOP_TRAFICO_LIMITE = 5
ANALYTICS_PEORES_REBOTE_LIMITE = 3
ANALYTICS_MEJORES_ENGAGEMENT_LIMITE = 3
ANALYTICS_CRUCES_LIMITE = 5

# Define política editorial por formato para marcadores visuales.
POLITICA_MARCADORES_POR_FORMATO = {
    "doc": {
        "✅": "Validado",
        "⚠️": "Revisión recomendada",
        "⚠": "Revisión recomendada",
        "❌": "Incidencia detectada",
        "🚨": "Alta prioridad",
        "📈": "Tendencia al alza",
        "📉": "Tendencia a la baja",
        "🎯": "Objetivo",
        "🔥": "Alta prioridad",
        "💡": "Recomendación",
    },
    "html": {
        "✅": "Validado",
        "⚠️": "Revisión recomendada",
        "⚠": "Revisión recomendada",
        "❌": "Incidencia detectada",
        "🚨": "Alta prioridad",
        "📈": "Tendencia al alza",
        "📉": "Tendencia a la baja",
        "🎯": "Objetivo",
        "🔥": "Alta prioridad",
        "💡": "Recomendación",
    },
    "excel": {
        "✅": "OK",
        "⚠️": "Alerta",
        "⚠": "Alerta",
        "❌": "Error",
        "🚨": "Crítico",
        "📈": "Alza",
        "📉": "Baja",
        "🎯": "Meta",
        "🔥": "Prioridad",
        "💡": "Idea",
    },
}


# Devuelve jerarquía visible según fuentes activas de la ejecución.
def construir_jerarquia_visible(resultado: ResultadoAuditoria) -> list[str]:
    """Filtra secciones GSC cuando la fuente no está activa."""

    # Parte de jerarquía base completa del informe.
    jerarquia = list(JERARQUIA_INFORME)

    # Filtra secciones dependientes de GSC cuando no esté activo.
    if not resultado.search_console.activo:
        # Mantiene solo secciones no dependientes de GSC.
        jerarquia = [seccion for seccion in jerarquia if seccion not in SECCIONES_GSC]

    # Oculta sección de Analytics cuando la fuente no esté activa.
    if not resultado.analytics.activo:
        # Mantiene solo secciones no dependientes de GA4.
        jerarquia = [seccion for seccion in jerarquia if seccion != "Comportamiento y conversión"]

    # Devuelve jerarquía visible final.
    return jerarquia


# Devuelve peso de orden para severidades en salida ejecutiva.
def _peso_severidad(severidad: str) -> int:
    """Convierte severidad textual en peso numérico de ordenación."""

    # Define mapa de prioridad para ordenar de mayor a menor criticidad.
    mapa = {"crítica": 0, "alta": 1, "media": 2, "baja": 3, "informativa": 4}

    # Devuelve peso conocido o fallback al final.
    return mapa.get(severidad.lower().strip(), 5)


# Devuelve color pastel de fondo para severidad en HTML.
def _color_pastel_severidad(severidad: str) -> str:
    """Asigna color suave por severidad para mejorar legibilidad visual."""

    # Devuelve rojo pastel para severidades críticas.
    if severidad.lower().strip() in {"crítica", "alta"}:
        # Retorna color rojo suave.
        return "#fde8e8"

    # Devuelve naranja pastel para severidad media.
    if severidad.lower().strip() == "media":
        # Retorna color naranja suave.
        return "#fff4e5"

    # Devuelve amarillo pastel para severidad baja.
    if severidad.lower().strip() == "baja":
        # Retorna color amarillo suave.
        return "#fef9c3"

    # Devuelve azul pastel para severidad informativa.
    return "#e6f0ff"


# Devuelve un valor de métrica en formato legible evitando `None`.
def _valor_metrica(valor: object) -> str:
    """Formatea métricas para documentos sin mostrar valores nulos."""

    # Devuelve etiqueta estándar cuando no exista dato.
    if valor in {None, ""}:
        # Retorna texto homogéneo en todos los exportadores.
        return "No disponible"

    # Devuelve el valor convertido a texto.
    return str(valor)

# Calcula anchos de columnas para PDF respetando el ancho útil del documento.
def _calcular_col_widths_pdf(columnas: list[str], filas: list[list[Any]], ancho_util: float) -> list[float]:
    """Distribuye anchos en puntos para evitar desbordes horizontales en A4."""

    # Devuelve lista vacía cuando no existan columnas.
    if not columnas:
        return []

    # Define mínimos de legibilidad por columna en puntos.
    min_por_columna = 54.0

    # Inicializa colección de pesos relativos por columna.
    limites: list[float] = []

    # Recorre columnas para estimar pesos de anchura.
    for indice, columna in enumerate(columnas):
        # Obtiene celdas existentes de la columna actual.
        celdas = [str(fila[indice]) for fila in filas if indice < len(fila)]

        # Calcula longitud estimada a partir de cabecera y muestra de valores.
        longitud = max([len(str(columna))] + [len(celda) for celda in celdas[:40]])

        # Normaliza nombre de columna para reglas de ponderación.
        nombre_columna = str(columna).lower()

        # Amplía peso en columnas narrativas con texto más largo.
        factor_narrativo = 1.8 if any(token in nombre_columna for token in ["motivo", "recomend", "observ", "descripcion", "acción"]) else 1.0

        # Acota longitud para evitar sesgo por valores extremos.
        limites.append(max(6.0, min(float(longitud), 60.0)) * factor_narrativo)

    # Calcula suma de pesos garantizando divisor válido.
    suma_pesos = sum(limites) or 1.0

    # Distribuye anchos proporcionalmente al peso estimado.
    anchos = [(peso / suma_pesos) * ancho_util for peso in limites]

    # Aplica mínimo por columna para mantener legibilidad.
    anchos = [max(min_por_columna, ancho) for ancho in anchos]

    # Recalcula suma tras aplicar mínimos.
    suma_anchos = sum(anchos)

    # Reescala anchos al espacio útil disponible.
    if suma_anchos > 0:
        escala = ancho_util / suma_anchos
        anchos = [ancho * escala for ancho in anchos]

    # Corrige posible desviación acumulada en la última columna.
    if anchos:
        delta = ancho_util - sum(anchos)
        anchos[-1] += delta

    # Devuelve anchos finales ajustados.
    return anchos

# Resuelve el periodo analizado efectivo con fallback entre fuentes temporales.
def _resolver_periodo_analizado(resultado: ResultadoAuditoria) -> tuple[str, str]:
    """Devuelve fecha inicial y final efectivas del periodo analizado."""

    # Obtiene bloque de Search Console de forma segura.
    search_console = getattr(resultado, "search_console", None)

    # Obtiene bloque de Analytics de forma segura.
    analytics = getattr(resultado, "analytics", None)

    # Obtiene fecha inicial desde Search Console si existe.
    search_console_desde = getattr(search_console, "date_from", "") or ""

    # Obtiene fecha inicial desde Analytics si existe.
    analytics_desde = getattr(analytics, "date_from", "") or ""

    # Obtiene fecha final desde Search Console si existe.
    search_console_hasta = getattr(search_console, "date_to", "") or ""

    # Obtiene fecha final desde Analytics si existe.
    analytics_hasta = getattr(analytics, "date_to", "") or ""

    # Prioriza periodo global persistido en resultado final.
    periodo_desde = resultado.periodo_date_from or search_console_desde or analytics_desde

    # Prioriza fin global persistido en resultado final.
    periodo_hasta = resultado.periodo_date_to or search_console_hasta or analytics_hasta

    # Devuelve tupla ordenada de fechas efectivas.
    return periodo_desde, periodo_hasta

# Genera una observación breve por score de rendimiento.
def _interpretacion_rendimiento(score: float | None, estrategia: str) -> str:
    """Interpreta visualmente el estado de rendimiento por estrategia."""

    # Devuelve observación neutral cuando no hay score.
    if score is None:
        # Informa ausencia de datos para análisis.
        return f"{estrategia.title()}: sin datos suficientes."

    # Evalúa score bajo de rendimiento.
    if score < 50:
        # Devuelve observación crítica.
        return f"{estrategia.title()}: rendimiento crítico, priorizar optimización."

    # Evalúa score medio de rendimiento.
    if score < 90:
        # Devuelve observación de mejora.
        return f"{estrategia.title()}: necesita mejora."

    # Devuelve observación positiva para score alto.
    return f"{estrategia.title()}: correcto."


# Autoajusta anchos, wraps y alturas para maximizar legibilidad.
def _autoajustar_hoja(hoja, columnas_wrap: set[str] | None = None, min_width: int = 10, max_width: int = 70) -> None:
    """Aplica autoajuste global de columnas y alturas de filas en una hoja."""

    # Inicializa conjunto de columnas con wrap recomendado.
    columnas_wrap = columnas_wrap or set()

    # Recorre columnas presentes para calcular ancho por contenido.
    for indice_columna in range(1, max(1, hoja.max_column) + 1):
        # Obtiene letra de columna para operar en openpyxl.
        letra_columna = get_column_letter(indice_columna)

        # Inicializa largo máximo detectado de la columna.
        largo_maximo = 0

        # Recorre celdas de la columna para medir texto real.
        for fila in range(1, max(1, hoja.max_row) + 1):
            # Lee valor de la celda actual.
            valor = hoja.cell(row=fila, column=indice_columna).value

            # Convierte valor a texto seguro para medir longitud.
            texto = str(valor) if valor is not None else ""

            # Considera la línea más larga en celdas multilínea.
            largo_texto = max((len(parte) for parte in texto.splitlines()), default=0)

            # Guarda máximo observado en columna.
            largo_maximo = max(largo_maximo, largo_texto)

            # Aplica wrap recomendado en columnas críticas.
            if texto and (fila == 1 or hoja.cell(row=1, column=indice_columna).value in columnas_wrap):
                # Preserva alineación superior para lectura.
                hoja.cell(row=fila, column=indice_columna).alignment = Alignment(wrap_text=True, vertical="top")

        # Calcula ancho final con márgenes y límites.
        ancho_final = min(max_width, max(min_width, largo_maximo + 3))

        # Aplica ancho calculado de columna.
        hoja.column_dimensions[letra_columna].width = ancho_final

    # Recorre filas para ajustar altura por contenido y saltos de línea.
    for indice_fila in range(2, max(2, hoja.max_row) + 1):
        # Calcula altura base por número de líneas visibles.
        lineas_maximas = 1

        # Recorre columnas de la fila.
        for indice_columna in range(1, max(1, hoja.max_column) + 1):
            # Lee valor de celda para estimar altura.
            valor = hoja.cell(row=indice_fila, column=indice_columna).value
            texto = str(valor) if valor is not None else ""
            lineas_maximas = max(lineas_maximas, texto.count("\n") + (1 if texto else 0))

        # Ajusta altura de fila con límite razonable.
        hoja.row_dimensions[indice_fila].height = min(120, max(20, 16 * lineas_maximas))


# Renderiza un bloque ejecutivo con cabecera y filas de texto legibles.
def _renderizar_bloque_dashboard(hoja, celda_inicio: str, titulo: str, lineas: list[str], color_bloque: str, columnas_ancho: int = 6) -> int:
    """Pinta un bloque del dashboard con jerarquía visual y alto dinámico."""

    # Obtiene columna y fila inicial para posicionar el bloque.
    columna_inicio = hoja[celda_inicio].column
    fila_inicio = hoja[celda_inicio].row

    # Calcula la última columna del bloque según ancho deseado.
    columna_fin = columna_inicio + max(1, columnas_ancho) - 1

    # Convierte índices de columna a letras de Excel.
    letra_inicio = get_column_letter(columna_inicio)
    letra_fin = get_column_letter(columna_fin)

    # Fusiona fila de cabecera para el título del bloque.
    hoja.merge_cells(f"{letra_inicio}{fila_inicio}:{letra_fin}{fila_inicio}")

    # Escribe el título del bloque en su cabecera.
    hoja[f"{letra_inicio}{fila_inicio}"] = titulo

    # Estiliza cabecera con color sólido y texto blanco.
    hoja[f"{letra_inicio}{fila_inicio}"].font = Font(size=12, bold=True, color="FFFFFF")
    hoja[f"{letra_inicio}{fila_inicio}"].alignment = Alignment(horizontal="left", vertical="center")
    hoja[f"{letra_inicio}{fila_inicio}"].fill = PatternFill(fill_type="solid", fgColor=color_bloque)

    # Recorre líneas de contenido para mostrarlas bajo la cabecera.
    for desplazamiento, linea in enumerate(lineas, start=1):
        # Calcula la fila de destino para la línea actual.
        fila_actual = fila_inicio + desplazamiento

        # Fusiona celdas horizontales para una lectura cómoda.
        hoja.merge_cells(f"{letra_inicio}{fila_actual}:{letra_fin}{fila_actual}")

        # Escribe línea de contenido del bloque.
        hoja[f"{letra_inicio}{fila_actual}"] = linea

        # Estiliza línea con fondo suave y texto legible.
        hoja[f"{letra_inicio}{fila_actual}"].font = Font(size=10, color="1F2937")
        hoja[f"{letra_inicio}{fila_actual}"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        hoja[f"{letra_inicio}{fila_actual}"].fill = PatternFill(fill_type="solid", fgColor="F8FAFC")

    # Devuelve la fila final pintada para facilitar composición.
    return fila_inicio + len(lineas)


# Construye quick wins agrupados por URL para capa ejecutiva clara.
def _construir_quick_wins(filas: list[dict], limite: int = 7) -> list[dict[str, object]]:
    """Devuelve quick wins agrupados por URL con impacto máximo y esfuerzo mínimo."""

    # Define escala de impacto para seleccionar el máximo por URL.
    escala_impacto = {"Muy alto": 4, "Alto": 3, "Medio": 2, "Bajo": 1}

    # Define escala de esfuerzo para seleccionar el mínimo por URL.
    escala_esfuerzo = {"Bajo": 1, "Medio": 2, "Alto": 3}

    # Inicializa contenedor de agrupación por URL.
    agrupados: dict[str, dict[str, object]] = {}

    # Recorre filas para construir quick wins agrupados.
    for fila in filas:
        # Obtiene URL base de la fila.
        url = str(fila.get("url", "")).strip()

        # Obtiene problema principal.
        problema = str(fila.get("problema", "")).strip()

        # Obtiene recomendación asociada.
        recomendacion = str(fila.get("recomendacion", "")).strip()

        # Obtiene impacto declarado.
        impacto = str(fila.get("impacto", "")).strip()

        # Obtiene esfuerzo declarado.
        esfuerzo = str(fila.get("esfuerzo", "")).strip()

        # Descarta filas incompletas para evitar tarjetas vacías.
        if not url or not problema or not recomendacion:
            # Continúa con la siguiente fila útil.
            continue

        # Filtra quick wins priorizando acciones de bajo esfuerzo.
        if esfuerzo not in {"Bajo", "Medio"} or impacto not in {"Muy alto", "Alto", "Medio"}:
            # Continúa cuando no cumpla criterio ejecutivo.
            continue

        # Inicializa estructura de la URL cuando no exista.
        if url not in agrupados:
            # Crea tarjeta base de quick win por URL.
            agrupados[url] = {"url": url, "problemas": [], "recomendaciones": [], "impacto_maximo": impacto or "Medio", "esfuerzo_maximo": esfuerzo or "Medio", "prioridad": 0}

        # Obtiene tarjeta actual de la URL.
        tarjeta = agrupados[url]

        # Añade problema evitando duplicados.
        if problema not in tarjeta["problemas"]:
            # Inserta problema único en la tarjeta.
            tarjeta["problemas"].append(problema)

        # Añade recomendación evitando duplicados.
        if recomendacion not in tarjeta["recomendaciones"]:
            # Inserta recomendación única en la tarjeta.
            tarjeta["recomendaciones"].append(recomendacion)

        # Actualiza impacto máximo detectado.
        if escala_impacto.get(impacto, 0) > escala_impacto.get(str(tarjeta["impacto_maximo"]), 0):
            # Guarda nuevo impacto máximo de la URL.
            tarjeta["impacto_maximo"] = impacto

        # Actualiza esfuerzo mínimo detectado.
        if escala_esfuerzo.get(esfuerzo, 99) < escala_esfuerzo.get(str(tarjeta["esfuerzo_maximo"]), 99):
            # Guarda nuevo esfuerzo mínimo de la URL.
            tarjeta["esfuerzo_maximo"] = esfuerzo

        # Calcula prioridad combinada para orden final.
        tarjeta["prioridad"] = escala_impacto.get(str(tarjeta["impacto_maximo"]), 0) * 10 - escala_esfuerzo.get(str(tarjeta["esfuerzo_maximo"]), 99)

    # Convierte agrupación a lista para ordenar.
    quick_wins = list(agrupados.values())

    # Ordena quick wins por prioridad descendente y URL estable.
    quick_wins_ordenados = sorted(quick_wins, key=lambda item: (-int(item.get("prioridad", 0)), str(item.get("url", ""))))

    # Devuelve subconjunto máximo solicitado.
    return quick_wins_ordenados[:limite]


# Construye fuente combinada de quick wins técnicos y de visibilidad real.
def construir_filas_quick_wins(resultado: ResultadoAuditoria) -> list[dict]:
    """Combina incidencias técnicas y oportunidades GSC para priorización real."""

    # Construye filas técnicas base de la auditoría.
    filas = construir_filas(resultado)

    # Recorre oportunidades GSC para convertirlas al mismo esquema.
    for item in construir_oportunidades_gsc(resultado):
        # Añade oportunidad GSC como quick win comparable.
        filas.append(
            {
                "url": item.get("url", ""),
                "problema": item.get("oportunidad", ""),
                "recomendacion": item.get("accion_recomendada", ""),
                "impacto": item.get("impacto", "Medio"),
                "esfuerzo": item.get("esfuerzo", "Medio"),
            }
        )

    # Devuelve colección combinada de quick wins.
    return filas


# Devuelve color de tarjeta según impacto para salida visual coherente.
def _color_por_impacto(impacto: str) -> str:
    """Asigna un color sobrio por nivel de impacto."""

    # Mapea impacto muy alto a rojo suave.
    if impacto == "Muy alto":
        # Devuelve rojo suave para destacar urgencia.
        return "#FDECEA"

    # Mapea impacto alto a naranja suave.
    if impacto == "Alto":
        # Devuelve naranja suave de prioridad.
        return "#FCE8E6"

    # Mapea impacto medio a amarillo suave.
    if impacto == "Medio":
        # Devuelve amarillo suave para atención moderada.
        return "#FFF4E5"

    # Devuelve azul suave para resto de casos.
    return "#E8F0FE"


# Calcula incidencias agrupadas para capa ejecutiva sin perder detalle técnico.
def _calcular_incidencias_agrupadas(filas: list[dict]) -> dict[str, int]:
    """Agrupa incidencias por tipología para resumen ejecutivo coherente."""

    # Inicializa contador agregado por familia de problema.
    agregadas: Counter[str] = Counter()

    # Recorre filas de incidencias técnicas.
    for fila in filas:
        # Obtiene descripción del problema actual.
        problema = str(fila.get("problema", "")).lower().strip()

        # Descarta filas informativas sin problema.
        if not problema:
            # Continúa al siguiente registro.
            continue

        # Agrupa normalizaciones de slash para evitar sobredimensionar resumen.
        if "slash final" in problema:
            # Cuenta una familia ejecutiva de normalización de URL.
            agregadas["Normalización URL (slash final)"] += 1
            # Continúa para evitar doble conteo.
            continue

        # Agrupa incidencias de imágenes sin alt en una misma familia.
        if "imágenes sin atributo alt" in problema or "imagen sin atributo alt" in problema:
            # Cuenta familia multimedia/accesibilidad.
            agregadas["Imágenes sin ALT"] += 1
            # Continúa para evitar doble conteo.
            continue

        # Agrupa problemas de titles en una familia ejecutiva.
        if "title" in problema:
            # Cuenta familia de títulos on-page.
            agregadas["Titles"] += 1
            # Continúa para evitar doble conteo.
            continue

        # Agrupa problemas de meta description en su familia.
        if "meta description" in problema:
            # Cuenta familia de metas on-page.
            agregadas["Meta descriptions"] += 1
            # Continúa para evitar doble conteo.
            continue

        # Agrupa problemas de h1 y headings.
        if "h1" in problema:
            # Cuenta familia de headings.
            agregadas["Headings"] += 1
            # Continúa para evitar doble conteo.
            continue

        # Registra resto de problemas como categoría literal normalizada.
        agregadas[problema[:90]] += 1

    # Devuelve incidencias agrupadas para capa ejecutiva.
    return dict(agregadas)


# Construye un prefijo de nombre de archivo legible y consistente.
def construir_prefijo_archivo(resultado: ResultadoAuditoria) -> str:
    """Genera un prefijo de naming profesional para entregables."""

    # Limpia el nombre del cliente para convertirlo en parte de archivo.
    cliente = resultado.cliente.lower().replace(" ", "_")

    # Devuelve el prefijo homogéneo de archivos exportados.
    return f"informe_seo_{cliente}_{resultado.fecha_ejecucion}"


# Sanea un texto libre para que reportlab no lo interprete como marcado inválido.
def sanear_texto_para_pdf(texto: str) -> str:
    """Limpia y escapa texto dinámico para render seguro en PDF."""

    # Normaliza capa editorial y evita emojis conflictivos.
    texto_normalizado = sanitizar_texto_final_exportable(texto, formato="doc").replace("\r\n", "\n").replace("\r", "\n")

    # Escapa caracteres especiales para evitar errores del parser XML interno.
    texto_escapado = escape(texto_normalizado)

    # Devuelve texto saneado para inserción segura en PDF.
    return texto_escapado


# Limpia marcas markdown para evitar texto crudo en DOCX/PDF.
def limpiar_markdown_crudo(texto: str) -> str:
    """Elimina marcas markdown frecuentes y deja texto natural legible."""

    # Elimina separadores horizontales de markdown.
    texto_limpio = re.sub(r"^\s*---+\s*$", "", texto, flags=re.MULTILINE)

    # Elimina encabezados markdown conservando el contenido.
    texto_limpio = re.sub(r"^\s*#{1,6}\s*", "", texto_limpio, flags=re.MULTILINE)

    # Elimina negritas y cursivas markdown.
    texto_limpio = texto_limpio.replace("**", "").replace("__", "")

    # Elimina bloques de código en línea.
    texto_limpio = texto_limpio.replace("`", "")

    # Devuelve el texto con limpieza por línea.
    return "\n".join(linea.strip() for linea in texto_limpio.splitlines())


# Sustituye marcadores visuales por texto editorial según formato de salida.
def reemplazar_emojis_problematicos(texto: str, formato: str = "doc") -> str:
    """Evita glifos conflictivos y aplica una política editorial coherente por formato."""

    # Normaliza nombre de formato con fallback documental.
    formato_normalizado = formato.strip().lower() if formato else "doc"

    # Resuelve mapa de sustitución por formato o fallback a documento.
    mapa_formato = POLITICA_MARCADORES_POR_FORMATO.get(formato_normalizado, POLITICA_MARCADORES_POR_FORMATO["doc"])

    # Inicializa salida con el texto original.
    salida = texto

    # Reemplaza cada marcador por su etiqueta editorial definida.
    for emoji, etiqueta in mapa_formato.items():
        # Sustituye ocurrencias para mejorar compatibilidad tipográfica.
        salida = salida.replace(emoji, etiqueta)

    # Elimina variaciones de selector para evitar símbolos residuales.
    salida = salida.replace("\ufe0f", "")

    # Devuelve texto sin glifos conflictivos para el formato solicitado.
    return salida


# Bloquea placeholders residuales para evitar fugas tipo [TOKEN_MAYUSCULA].
def bloquear_placeholders_residuales(texto: str) -> str:
    """Neutraliza tokens residuales entre corchetes con texto legible sin corchetes."""

    # Convierte placeholders mayúsculos a texto editorial legible.
    texto_limpio = re.sub(r"\[([A-ZÁÉÍÓÚÜÑ_]+)\]", lambda coincidencia: coincidencia.group(1).replace("_", " ").title(), texto)

    # Elimina corchetes residuales de cualquier marcador no previsto.
    texto_limpio = re.sub(r"\[([^\]]+)\]", r"\1", texto_limpio)

    # Devuelve contenido sin placeholders técnicos en corchetes.
    return texto_limpio


# Normaliza contenido narrativo para una capa semántica estable.
def sanitizar_texto_editorial(texto: str, formato: str = "doc") -> str:
    """Limpia residuos markdown y homogeniza estilo narrativo del informe."""

    # Limpia markdown crudo de entrada.
    texto_limpio = limpiar_markdown_crudo(texto)

    # Sustituye emojis no seguros por etiquetas legibles.
    texto_limpio = reemplazar_emojis_problematicos(texto_limpio, formato=formato)

    # Colapsa espacios múltiples preservando saltos.
    texto_limpio = re.sub(r"[ \t]+", " ", texto_limpio)

    # Normaliza separación tras porcentajes y tiempos comunes.
    texto_limpio = re.sub(r"(\d)\s*%", r"\1%", texto_limpio)
    texto_limpio = re.sub(r"(\d)\s*s\b", r"\1 s", texto_limpio)
    texto_limpio = re.sub(r"(\d)\s*ms\b", r"\1 ms", texto_limpio)
    # Bloquea placeholders residuales para no exportar tokens técnicos.
    texto_limpio = bloquear_placeholders_residuales(texto_limpio)

    # Devuelve texto editorial consistente.
    return texto_limpio.strip()


# Ejecuta sanitización final obligatoria para salidas documentales exportables.
def sanitizar_texto_final_exportable(texto: str, formato: str) -> str:
    """Aplica política final por formato y bloquea placeholders antes de exportar."""

    # Normaliza y sanea texto con la política del formato objetivo.
    texto_limpio = sanitizar_texto_editorial(texto, formato=formato)

    # Devuelve texto ya saneado por la política del formato objetivo.
    return texto_limpio


# Convierte un valor genérico en texto/celda seguro para exportación a Excel.
def sanitizar_valor_excel(valor: Any, limite: int = 96) -> Any:
    """Normaliza strings para Excel con versión corta y compatible."""

    # Devuelve valores no textuales sin transformación.
    if not isinstance(valor, str):
        # Mantiene tipos numéricos y booleanos para cálculos.
        return valor

    # Sanea contenido textual siguiendo política específica de Excel.
    valor_limpio = sanitizar_texto_final_exportable(valor, formato="excel")

    # Recorta longitud máxima para mejorar legibilidad de celdas.
    return valor_limpio[:limite]


# Construye un bloque semántico base reutilizable por todos los exportadores.
def _crear_bloque_semantico(tipo: str, titulo: str, subtitulo: str = "", **campos: Any) -> dict[str, Any]:
    """Crea un bloque neutral con estructura explícita para render final."""

    # Inicializa bloque con propiedades troncales.
    bloque: dict[str, Any] = {
        "tipo_bloque": tipo,
        "titulo": titulo,
        "subtitulo": subtitulo,
        "parrafos": [],
        "listas": [],
        "tablas": [],
        "tarjetas": [],
        "notas": [],
    }

    # Inserta campos complementarios enviados por el llamador.
    bloque.update(campos)

    # Devuelve bloque semántico listo para consumo.
    return bloque


# Convierte narrativa IA en secciones internas estructuradas.
def construir_secciones_desde_ia(texto_ia: str | None) -> list[dict[str, object]]:
    """Transforma texto IA en secciones semánticas sin markdown crudo."""

    # Devuelve colección vacía cuando no exista contenido IA.
    if not texto_ia:
        # Retorna una lista vacía para simplificar flujo.
        return []

    # Limpia markdown para evitar arrastre en entregables.
    texto_limpio = limpiar_markdown_crudo(texto_ia)

    # Inicializa resultado de secciones intermedias.
    secciones: list[dict[str, object]] = []

    # Inicializa sección por defecto de arranque.
    seccion_actual: dict[str, object] = {"titulo": "Resumen ejecutivo", "tipo": "parrafos", "items": []}

    # Recorre líneas del texto limpio.
    for linea in texto_limpio.splitlines():
        # Descarta líneas vacías.
        if not linea.strip():
            # Continúa con la siguiente línea útil.
            continue

        # Detecta títulos de sección de forma tolerante.
        if re.match(r"^(\d+[\).]\s*)?(resumen|kpis|hallazgos|quick wins|acciones técnicas|acciones de contenido|rendimiento|roadmap)", linea.lower()):
            # Guarda sección previa cuando tenga contenido.
            if seccion_actual["items"]:
                # Inserta sección construida en la lista final.
                secciones.append(seccion_actual)

            # Limpia numeración inicial del título.
            titulo = re.sub(r"^\d+[\).]\s*", "", linea).strip().title()

            # Crea nueva sección activa.
            seccion_actual = {"titulo": titulo, "tipo": "parrafos", "items": []}

            # Continúa sin añadir la línea como contenido.
            continue

        # Detecta viñetas con prefijos habituales.
        if linea.startswith(("- ", "* ", "• ")):
            # Cambia tipo de sección a viñetas.
            seccion_actual["tipo"] = "vinetas"

            # Añade item sin el prefijo de viñeta.
            seccion_actual["items"].append(linea[2:].strip())

            # Continúa con la siguiente línea.
            continue

        # Detecta listas numeradas y las normaliza.
        if re.match(r"^\d+[\).]\s+", linea):
            # Cambia tipo de sección a lista numerada.
            seccion_actual["tipo"] = "numerada"

            # Añade item eliminando el prefijo numérico.
            seccion_actual["items"].append(re.sub(r"^\d+[\).]\s*", "", linea).strip())

            # Continúa con la siguiente línea.
            continue

        # Segmenta párrafos excesivamente largos para hacerlos más escaneables.
        if len(linea.strip()) > 220 and ". " in linea:
            # Reparte el contenido en frases cortas manteniendo contexto.
            for fragmento in [parte.strip() for parte in linea.split(". ") if parte.strip()]:
                # Añade frase con cierre de puntuación consistente.
                seccion_actual["items"].append(fragmento if fragmento.endswith(".") else f"{fragmento}.")
        else:
            # Añade línea como párrafo normal.
            seccion_actual["items"].append(linea.strip())

    # Añade última sección pendiente si tiene contenido.
    if seccion_actual["items"]:
        # Inserta sección final en la colección.
        secciones.append(seccion_actual)

    # Devuelve secciones intermedias para renderizado estable.
    return secciones


# Convierte el resultado a una estructura tabular cómoda para exportación.
def construir_filas(resultado: ResultadoAuditoria) -> list[dict]:
    """Convierte el resultado de auditoría a una lista de filas tabulares."""

    # Inicializa la colección de filas tabulares.
    filas: list[dict] = []

    # Recorre cada URL auditada para convertirla a filas de incidencias.
    for item in resultado.resultados:
        # Construye una base común para evitar duplicación de campos.
        fila_base = {
            "url": item.url,
            "url_final": item.url_final,
            "tipo": item.tipo,
            "estado_http": item.estado_http,
            "redirecciona": "Sí" if item.redirecciona else "No",
            "title": item.title,
            "h1": item.h1,
            "meta_description": item.meta_description,
            "canonical": item.canonical or "",
            "noindex": "Sí" if item.noindex else "No",
            "palabras": getattr(item, "palabras", 0),
            "calidad_contenido": getattr(item, "calidad_contenido", "baja"),
            "thin_content": "Sí" if getattr(item, "thin_content", False) else "No",
            "densidad_texto": getattr(item, "densidad_texto", 0.0),
            "ratio_texto_html": getattr(item, "ratio_texto_html", 0.0),
            "h1_unico": "Sí" if getattr(item, "h1_unico", False) else "No",
            "estructura_headings_correcta": "Sí" if getattr(item, "estructura_headings_correcta", True) else "No",
            "imagenes_sin_alt": getattr(item, "imagenes_sin_alt", 0),
            "lazy_load_detectado": "Sí" if getattr(item, "lazy_load_detectado", False) else "No",
            "estado": "Pendiente",
            "resuelto": "No",
            "responsable": "",
        }

        # Inserta fila de control cuando no haya hallazgos.
        if not item.hallazgos:
            # Crea copia independiente de la fila base.
            fila = fila_base.copy()

            # Completa datos por defecto para URL sin incidencias.
            fila.update({"problema": "", "recomendacion": "", "severidad": "informativa", "area": "Calidad", "impacto": "Bajo", "esfuerzo": "Bajo", "prioridad": "P4", "observaciones": item.error or "Sin incidencias críticas"})
            # Añade categoría ejecutiva derivada para reporting.
            fila["categoria"] = "Calidad"

            # Añade fila informativa final a la colección.
            filas.append(fila)

        # Recorre hallazgos cuando existan para crear filas detalladas.
        for hallazgo in item.hallazgos:
            # Crea copia independiente de la fila base.
            fila = fila_base.copy()

            # Completa datos de hallazgo para esta fila.
            fila.update({"problema": hallazgo.descripcion, "recomendacion": hallazgo.recomendacion, "severidad": hallazgo.severidad, "area": hallazgo.area, "impacto": hallazgo.impacto, "esfuerzo": hallazgo.esfuerzo, "prioridad": hallazgo.prioridad, "observaciones": item.error or ""})
            # Añade categoría operativa alineada al tipo de hallazgo.
            fila["categoria"] = hallazgo.tipo

            # Añade una fila por hallazgo detectado.
            filas.append(fila)

    # Devuelve las filas preparadas para cualquier exportador.
    return filas


# Convierte resultados de PageSpeed a filas de seguimiento.
def construir_filas_contenido_consolidado(resultado: ResultadoAuditoria) -> list[dict]:
    """Devuelve filas de contenido sin duplicar URLs por incidencia."""

    # Inicializa colección consolidada indexada por URL.
    contenido_por_url: dict[str, dict[str, object]] = {}

    # Recorre resultados técnicos por URL para crear base consolidada.
    for item in resultado.resultados:
        # Inserta registro único por URL con métricas on-page.
        contenido_por_url[item.url] = {
            "url": item.url,
            "palabras": item.palabras,
            "calidad_contenido": item.calidad_contenido,
            "h1": item.h1,
            "title": item.title,
            "meta_description": item.meta_description,
            "imagenes_sin_alt": item.imagenes_sin_alt,
            "thin_content": "Sí" if item.thin_content else "No",
            "densidad_texto": item.densidad_texto,
            "ratio_texto_html": item.ratio_texto_html,
            "incidencias_url": len(item.hallazgos),
        }

    # Devuelve colección consolidada ordenada por URL para estabilidad.
    return [contenido_por_url[url] for url in sorted(contenido_por_url.keys())]


# Convierte resultados de PageSpeed a filas de seguimiento.
def construir_filas_rendimiento(resultado: ResultadoAuditoria) -> list[dict]:
    """Construye filas para hoja Rendimiento y anexo técnico de performance."""

    # Inicializa lista de filas de rendimiento.
    filas: list[dict] = []

    # Recorre cada resultado de PageSpeed disponible.
    for item in resultado.rendimiento:
        # Construye una base común para evitar duplicación de campos.
        fila_base = {
            "url": item.url,
            "estrategia": item.estrategia,
            "performance_score": item.performance_score,
            "accessibility_score": item.accessibility_score,
            "best_practices_score": item.best_practices_score,
            "seo_score": item.seo_score,
            "lcp": item.lcp,
            "cls": item.cls,
            "inp": item.inp,
            "fcp": item.fcp,
            "speed_index": item.speed_index,
            "estado": "Pendiente",
            "resuelto": "No",
            "responsable": "",
        }

        # Crea una fila base cuando no haya oportunidades.
        if not item.oportunidades:
            # Crea copia independiente de la fila base.
            fila = fila_base.copy()

            # Completa datos por defecto para esta ejecución.
            fila.update({"oportunidad": "", "descripcion": item.error or "Sin oportunidades destacadas", "ahorro_estimado": "", "severidad": "informativa", "recomendacion": "Mantener monitorización continua", "observaciones": item.error or ""})

            # Añade fila base de métricas sin oportunidad concreta.
            filas.append(fila)

        # Recorre oportunidades para crear filas accionables.
        for oportunidad in item.oportunidades:
            # Crea copia independiente de la fila base.
            fila = fila_base.copy()

            # Completa datos específicos de la oportunidad.
            fila.update({"oportunidad": oportunidad.titulo, "descripcion": oportunidad.descripcion, "ahorro_estimado": oportunidad.ahorro_estimado, "severidad": oportunidad.severidad, "recomendacion": f"Aplicar mejora: {oportunidad.titulo}", "observaciones": ""})

            # Añade fila por oportunidad detectada.
            filas.append(fila)

    # Devuelve filas listas para Excel.
    return filas


# Convierte datos de Search Console por página a filas tabulares.
def construir_filas_search_console_paginas(resultado: ResultadoAuditoria) -> list[dict]:
    """Construye filas por página cuando Search Console esté activo."""

    # Obtiene bloque GSC del resultado consolidado.
    datos_gsc = resultado.search_console

    # Devuelve lista vacía cuando GSC no esté activo.
    if not datos_gsc.activo:
        # Retorna colección vacía para degradación elegante.
        return []

    # Inicializa colección de filas por página.
    filas: list[dict] = []

    # Recorre métricas por página para crear filas.
    for item in datos_gsc.paginas:
        # Añade fila tipada con métricas base.
        filas.append(
            {
                "url": item.url,
                "clics": round(item.clicks, 2),
                "impresiones": round(item.impresiones, 2),
                "ctr": round(item.ctr, 4),
                "posicion_media": round(item.posicion_media, 2),
                "dispositivo": item.dispositivo,
                "pais": item.pais,
            }
        )

    # Devuelve filas listas para reportes.
    return filas


# Convierte datos de Search Console por query a filas tabulares.
def construir_filas_search_console_queries(resultado: ResultadoAuditoria) -> list[dict]:
    """Construye filas por query cuando Search Console esté activo."""

    # Obtiene bloque GSC del resultado consolidado.
    datos_gsc = resultado.search_console

    # Devuelve lista vacía cuando GSC no esté activo.
    if not datos_gsc.activo:
        # Retorna colección vacía para degradación elegante.
        return []

    # Inicializa colección de filas por query.
    filas: list[dict] = []

    # Recorre métricas por query para crear filas.
    for item in datos_gsc.queries:
        # Añade fila tipada con métricas base.
        filas.append(
            {
                "query": item.query,
                "clics": round(item.clicks, 2),
                "impresiones": round(item.impresiones, 2),
                "ctr": round(item.ctr, 4),
                "posicion_media": round(item.posicion_media, 2),
                "url_asociada": item.url_asociada,
            }
        )

    # Devuelve filas listas para reportes.
    return filas


# Convierte datos de Analytics por página a filas tabulares.
def construir_filas_analytics_paginas(resultado: ResultadoAuditoria) -> list[dict]:
    """Construye filas por página cuando Analytics esté activo."""

    # Obtiene bloque Analytics del resultado consolidado.
    datos_analytics = resultado.analytics

    # Devuelve lista vacía cuando Analytics no esté activo.
    if not datos_analytics.activo:
        # Retorna colección vacía para degradación elegante.
        return []

    # Inicializa colección de filas por página.
    filas: list[dict] = []

    # Recorre métricas por página para crear filas.
    for item in datos_analytics.paginas:
        # Añade fila tipada con métricas base.
        filas.append(
            {
                "url": item.url,
                "sesiones": round(item.sesiones, 2),
                "usuarios": round(item.usuarios, 2),
                "rebote": round(item.rebote, 4),
                "duracion": round(item.duracion_media, 2),
                "conversiones": round(item.conversiones, 2),
                "calidad_trafico": item.calidad_trafico,
            }
        )

    # Devuelve filas listas para reportes.
    return filas


# Normaliza una URL o ruta para cruce entre GSC (URL completa) y GA4 (pagePath).
def _clave_url_cruce(url_o_ruta: str) -> str:
    """Genera una clave comparable común entre distintas fuentes de datos."""

    # Limpia espacios y corta cuando no exista valor.
    valor = str(url_o_ruta or "").strip()
    if not valor:
        # Devuelve cadena vacía para descartar filas inválidas.
        return ""

    # Parsea valor para extraer componentes cuando venga URL completa.
    parseada = urlparse(valor)

    # Usa path de URL completa o valor original cuando ya sea ruta.
    path = (parseada.path or valor).strip()

    # Normaliza slash inicial.
    if not path.startswith("/"):
        path = f"/{path}"

    # Elimina slash final no raíz para evitar falsos negativos.
    if path != "/" and path.endswith("/"):
        path = path.rstrip("/")

    # Devuelve clave comparable en minúsculas.
    return path.lower()


# Construye señales de cruce entre Search Console y Analytics por URL.
def construir_cruces_gsc_analytics(resultado: ResultadoAuditoria) -> list[dict]:
    """Relaciona visibilidad y comportamiento para detectar oportunidades reales."""

    # Requiere que ambas fuentes estén activas para cruzar datos.
    if not (resultado.search_console.activo and resultado.analytics.activo):
        # Retorna colección vacía sin romper flujo.
        return []

    # Construye índice de Search Console por URL.
    gsc_por_url = {_clave_url_cruce(item.url): item for item in resultado.search_console.paginas if _clave_url_cruce(item.url)}

    # Construye índice de Analytics por URL.
    analytics_por_url = {_clave_url_cruce(item.url): item for item in resultado.analytics.paginas if _clave_url_cruce(item.url)}

    # Calcula conjunto común de URLs disponibles en ambas fuentes.
    urls_cruzables = sorted(set(gsc_por_url.keys()) & set(analytics_por_url.keys()))

    # Inicializa colección de cruces.
    cruces: list[dict] = []

    # Construye índice de gestión de indexación para enriquecer insights de negocio.
    indice_indexacion = {_clave_url_cruce(item.url): item for item in resultado.gestion_indexacion if _clave_url_cruce(item.url)}

    # Recorre URLs comunes y deriva insights.
    for url in urls_cruzables:
        # Obtiene métricas de Search Console para la URL.
        metrica_gsc = gsc_por_url[url]

        # Obtiene métricas de Analytics para la URL.
        metrica_analytics = analytics_por_url[url]

        # Inicializa insights detectados por URL.
        insights: list[str] = []

        # Detecta tráfico activo con mala experiencia de usuario.
        if metrica_analytics.sesiones >= CRUCE_GA_MIN_SESIONES_REBOTE and metrica_analytics.rebote >= CRUCE_GA_MIN_REBOTE_ALTO:
            # Inserta insight de UX degradada.
            insights.append("Tráfico alto con mala experiencia (rebote elevado).")

        # Detecta visibilidad orgánica sin engagement equivalente.
        if metrica_gsc.impresiones >= CRUCE_GSC_MIN_IMPRESIONES_SIN_ENGAGEMENT and metrica_analytics.sesiones <= CRUCE_GA_MAX_SESIONES_SIN_ENGAGEMENT:
            # Inserta insight de baja captación/enganche.
            insights.append("Alta visibilidad sin engagement suficiente.")

        # Detecta potencial orgánico sin conversión.
        if (
            metrica_gsc.posicion_media <= CRUCE_GSC_MAX_POSICION_POTENCIAL
            and metrica_analytics.sesiones >= CRUCE_GA_MIN_SESIONES_SIN_CONVERSION
            and metrica_analytics.conversiones <= CRUCE_GA_MAX_CONVERSIONES
        ):
            # Inserta insight de bajo rendimiento de conversión.
            insights.append("Potencial SEO sin conversión (revisar intención/CTA).")

        # Detecta impresiones elevadas con CTR bajo para optimización de snippet.
        if metrica_gsc.impresiones >= OPORTUNIDAD_GSC_MIN_IMPRESIONES_CTR and metrica_gsc.ctr <= OPORTUNIDAD_GSC_MAX_CTR:
            # Inserta insight de bajo rendimiento de clic.
            insights.append("Muchas impresiones y CTR bajo (optimizar snippet).")

        # Detecta conversiones en URLs que no deberían indexarse.
        decision_indexacion = indice_indexacion.get(url)
        if decision_indexacion and decision_indexacion.clasificacion == "NO_INDEXAR" and metrica_analytics.conversiones > 0:
            # Inserta insight de conflicto entre negocio e indexación.
            insights.append("Convierte pero está marcada como NO_INDEXAR (validar estrategia).")

        # Añade fila de cruce con métricas comparables.
        cruces.append(
            {
                "url": metrica_gsc.url or metrica_analytics.url or url,
                "impresiones": round(metrica_gsc.impresiones, 2),
                "sesiones": round(metrica_analytics.sesiones, 2),
                "ctr": round(metrica_gsc.ctr, 4),
                "rebote": round(metrica_analytics.rebote, 4),
                "posicion_media": round(metrica_gsc.posicion_media, 2),
                "duracion_media": round(metrica_analytics.duracion_media, 2),
                "conversiones": round(metrica_analytics.conversiones, 2),
                "insights": insights,
            }
        )

    # Devuelve cruces priorizados por impresiones y sesiones.
    return sorted(cruces, key=lambda item: (-(item["impresiones"]), -(item["sesiones"])))


# Convierte decisiones de indexación inteligente a filas tabulares.
def construir_filas_gestion_indexacion(resultado: ResultadoAuditoria) -> list[dict]:
    """Construye filas operativas de gestión de indexación para informes y Excel."""

    # Inicializa lista de filas de gestión.
    filas: list[dict] = []

    # Recorre decisiones calculadas en la auditoría.
    for decision in resultado.gestion_indexacion:
        # Añade fila con columnas del modelo solicitado.
        filas.append(
            {
                "url": decision.url,
                "clasificacion": decision.clasificacion,
                "motivo": decision.motivo,
                "accion_recomendada": decision.accion_recomendada,
                "prioridad": decision.prioridad,
            }
        )

    # Devuelve filas finales de gestión de indexación.
    return filas


# Calcula score y explicación estructurada para priorización SEO por URL.
def calcular_score_prioridad_pagina(gsc: object | None, ga: object | None, tecnico: object | None) -> dict[str, object]:
    """Evalúa señales actuales de priorización y devuelve trazabilidad explicable."""

    # Inicializa score acumulado final.
    score_prioridad = 0.0

    # Inicializa lista de motivos legibles para el informe.
    razones: list[str] = []

    # Inicializa detalle por componente para evolución futura del motor.
    componentes = {
        "potencial_seo": 0.0,
        "oportunidad_ctr": 0.0,
        "oportunidad_conversion": 0.0,
        "friccion_tecnica_onpage": 0.0,
        "esfuerzo_estimado": 0.0,
    }

    # Evalúa oportunidad de CTR con alta visibilidad.
    if gsc and gsc.impresiones >= PRIORIDAD_GSC_MIN_IMPRESIONES_CTR and gsc.ctr <= PRIORIDAD_GSC_MAX_CTR:
        # Suma peso asociado al componente de CTR.
        componentes["oportunidad_ctr"] += PRIORIDAD_SCORE_CTR_BAJO

        # Registra motivo explicativo para la capa ejecutiva.
        razones.append("Muchas impresiones y CTR bajo.")

    # Evalúa oportunidad por tráfico sin conversión.
    if ga and ga.sesiones >= PRIORIDAD_GA_MIN_SESIONES_SIN_CONVERSION and ga.conversiones <= PRIORIDAD_GA_MAX_CONVERSIONES:
        # Suma peso asociado a componente de conversión.
        componentes["oportunidad_conversion"] += PRIORIDAD_SCORE_SIN_CONVERSION

        # Registra motivo explicativo para la capa ejecutiva.
        razones.append("Sesiones altas sin conversión.")

    # Evalúa fricción de engagement entre visibilidad y rebote.
    if gsc and ga and gsc.impresiones >= PRIORIDAD_GSC_MIN_IMPRESIONES_ENGAGEMENT_BAJO and ga.rebote >= PRIORIDAD_GA_MIN_REBOTE_BAJO_ENGAGEMENT:
        # Suma peso asociado a fricción de comportamiento.
        componentes["friccion_tecnica_onpage"] += PRIORIDAD_SCORE_ENGAGEMENT_BAJO

        # Registra motivo explicativo para la capa ejecutiva.
        razones.append("Alta visibilidad con engagement bajo.")

    # Evalúa potencial de crecimiento por posición media escalable.
    if gsc and PRIORIDAD_GSC_MIN_POSICION <= gsc.posicion_media <= PRIORIDAD_GSC_MAX_POSICION:
        # Suma peso asociado al potencial SEO de ascenso.
        componentes["potencial_seo"] += PRIORIDAD_SCORE_POSICION_POTENCIAL

        # Registra motivo explicativo para la capa ejecutiva.
        razones.append("Posición 4-15 con potencial de crecimiento.")

    # Evalúa fricción técnica por volumen de hallazgos detectados.
    if tecnico and tecnico.hallazgos:
        # Calcula contribución máxima del bloque técnico.
        contribucion_hallazgos = min(PRIORIDAD_SCORE_MAX_HALLAZGOS, len(tecnico.hallazgos) * PRIORIDAD_SCORE_HALLAZGO)

        # Suma contribución al componente técnico/on-page.
        componentes["friccion_tecnica_onpage"] += contribucion_hallazgos

        # Registra motivo explicativo para la capa ejecutiva.
        razones.append(f"Acumula {len(tecnico.hallazgos)} hallazgos on-page/técnicos.")

    # Suma componentes para obtener score final.
    score_prioridad = round(sum(float(valor) for valor in componentes.values()), 1)

    # Devuelve estructura trazable de priorización.
    return {
        "score_prioridad": score_prioridad,
        "motivos": razones,
        "componentes": componentes,
    }


# Construye priorización de páginas combinando visibilidad, tráfico y conversión.
def construir_paginas_prioritarias(resultado: ResultadoAuditoria, limite: int = 10) -> list[dict]:
    """Ordena páginas por potencial SEO y negocio usando GSC, GA4 y on-page."""

    # Construye índices de datos por URL normalizada para cruces consistentes.
    gsc_por_clave = {_clave_url_cruce(item.url): item for item in resultado.search_console.paginas if _clave_url_cruce(item.url)}
    ga_por_clave = {_clave_url_cruce(item.url): item for item in resultado.analytics.paginas if _clave_url_cruce(item.url)}
    tecnico_por_clave = {_clave_url_cruce(item.url): item for item in resultado.resultados if _clave_url_cruce(item.url)}

    # Calcula universo de URLs disponibles en cualquier fuente.
    claves = sorted(set(gsc_por_clave.keys()) | set(ga_por_clave.keys()) | set(tecnico_por_clave.keys()))

    # Inicializa colección de páginas priorizadas.
    prioridades: list[dict] = []

    # Recorre cada URL candidata para asignar puntuación de priorización.
    for clave in claves:
        # Obtiene métricas por fuente cuando existan.
        gsc = gsc_por_clave.get(clave)
        ga = ga_por_clave.get(clave)
        tecnico = tecnico_por_clave.get(clave)

        # Resuelve URL legible priorizando formato completo.
        url = (gsc.url if gsc else "") or (ga.url if ga else "") or (tecnico.url if tecnico else clave)

        # Calcula score y detalle explicable desde función dedicada.
        evaluacion_prioridad = calcular_score_prioridad_pagina(gsc=gsc, ga=ga, tecnico=tecnico)

        # Obtiene score final consolidado.
        score_prioridad = float(evaluacion_prioridad["score_prioridad"])

        # Obtiene motivos legibles para narrativa y tablas.
        razones = list(evaluacion_prioridad["motivos"])

        # Descarta URLs sin señales de priorización claras.
        if score_prioridad <= 0:
            continue

        # Añade fila consolidada de priorización.
        prioridades.append(
            {
                "url": url,
                "prioridad_score": round(score_prioridad, 1),
                "impresiones": round(gsc.impresiones, 2) if gsc else 0.0,
                "ctr": round(gsc.ctr, 4) if gsc else 0.0,
                "posicion_media": round(gsc.posicion_media, 2) if gsc else 0.0,
                "sesiones": round(ga.sesiones, 2) if ga else 0.0,
                "conversiones": round(ga.conversiones, 2) if ga else 0.0,
                "hallazgos_onpage": len(tecnico.hallazgos) if tecnico else 0,
                "motivos": razones,
                "explicacion_prioridad": evaluacion_prioridad["componentes"],
            }
        )

    # Devuelve páginas priorizadas por score descendente.
    return sorted(prioridades, key=lambda item: (-float(item["prioridad_score"]), -float(item["impresiones"]), -float(item["sesiones"])))[:limite]


# Construye oportunidades SEO accionables cruzando técnica y Search Console.
def construir_oportunidades_gsc(resultado: ResultadoAuditoria) -> list[dict]:
    """Genera oportunidades reales de crecimiento cuando GSC esté activo."""

    # Obtiene bloque GSC del resultado consolidado.
    datos_gsc = resultado.search_console

    # Devuelve lista vacía cuando GSC no esté activo.
    if not datos_gsc.activo:
        # Retorna colección vacía por degradación elegante.
        return []

    # Construye índice rápido por URL de resultados técnicos.
    indice_resultados = {item.url: item for item in resultado.resultados}

    # Inicializa colección de oportunidades priorizadas.
    oportunidades: list[dict] = []

    # Recorre páginas con datos de Search Console.
    for metrica in datos_gsc.paginas:
        # Obtiene resultado técnico asociado si existe.
        pagina = indice_resultados.get(metrica.url)

        # Inicializa flags de problema on-page.
        problema_onpage = ""

        # Evalúa problema on-page cuando exista página técnica asociada.
        if pagina:
            # Detecta ausencia de meta description.
            if not pagina.meta_description.strip():
                # Define problema principal para recomendación.
                problema_onpage = "meta description ausente"
            # Detecta title potencialmente mejorable por longitud.
            elif len(pagina.title.strip()) > 60 or len(pagina.title.strip()) < 15:
                # Define problema principal para recomendación.
                problema_onpage = "title mejorable"
            # Detecta thin content sobre URL con visibilidad real.
            elif pagina.thin_content:
                # Define problema principal para recomendación.
                problema_onpage = "thin content con visibilidad"
            # Detecta debilidad estructural de headings.
            elif not pagina.h1_unico or not pagina.estructura_headings_correcta:
                # Define problema principal para recomendación.
                problema_onpage = "headings deficientes"

        # Inicializa etiqueta de oportunidad detectada.
        oportunidad = ""

        # Clasifica oportunidad por impresiones altas y CTR bajo.
        if metrica.impresiones >= OPORTUNIDAD_GSC_MIN_IMPRESIONES_CTR and metrica.ctr <= OPORTUNIDAD_GSC_MAX_CTR:
            # Define oportunidad prioritaria de CTR.
            oportunidad = "alto volumen con CTR bajo"
        # Clasifica oportunidad por posición de crecimiento rápido.
        elif OPORTUNIDAD_GSC_MIN_POSICION <= metrica.posicion_media <= OPORTUNIDAD_GSC_MAX_POSICION and metrica.impresiones >= OPORTUNIDAD_GSC_MIN_IMPRESIONES_POSICION:
            # Define oportunidad prioritaria por posición.
            oportunidad = "posición 4-15 con potencial"
        # Clasifica oportunidad por visibilidad y debilidad on-page.
        elif metrica.impresiones >= OPORTUNIDAD_GSC_MIN_IMPRESIONES_ONPAGE and problema_onpage:
            # Define oportunidad de mejora on-page.
            oportunidad = "visibilidad desaprovechada por on-page"

        # Descarta filas sin patrón claro de oportunidad.
        if not oportunidad:
            # Continúa con siguiente URL analizada.
            continue

        # Calcula impacto de negocio estimado por impresiones.
        impacto = "Alto" if metrica.impresiones >= OPORTUNIDAD_GSC_IMPRESIONES_ALTO_IMPACTO else "Medio"

        # Calcula esfuerzo estimado por tipo de problema.
        esfuerzo = "Bajo" if problema_onpage in {"meta description ausente", "title mejorable", "headings deficientes"} else "Medio"

        # Calcula prioridad final combinando impacto y esfuerzo.
        prioridad = "P1" if impacto == "Alto" and esfuerzo == "Bajo" else "P2"

        # Añade oportunidad consolidada para reporting.
        oportunidades.append(
            {
                "url": metrica.url,
                "query_principal": "",
                "clics": round(metrica.clicks, 2),
                "impresiones": round(metrica.impresiones, 2),
                "ctr": round(metrica.ctr, 4),
                "posicion_media": round(metrica.posicion_media, 2),
                "problema_onpage_detectado": problema_onpage or "sin señal on-page crítica",
                "oportunidad": oportunidad,
                "accion_recomendada": "Optimizar snippet y contenido focalizado en intención de búsqueda principal.",
                "impacto": impacto,
                "esfuerzo": esfuerzo,
                "prioridad": prioridad,
                "estado": "Pendiente",
                "resuelto": "No",
                "responsable": "",
                "observaciones": "",
            }
        )

    # Ordena oportunidades por prioridad e impresiones descendentes.
    return sorted(oportunidades, key=lambda item: (item["prioridad"], -float(item["impresiones"])))


# Calcula métricas agregadas de Search Console para reutilización transversal.
def _calcular_metricas_gsc(filas_gsc_paginas: list[dict]) -> dict[str, float]:
    """Devuelve clics, impresiones, CTR y posición media desde filas GSC."""

    # Calcula clics totales agregados de Search Console.
    clics_totales = round(sum(float(fila.get("clics", 0.0)) for fila in filas_gsc_paginas), 2)

    # Calcula impresiones totales agregadas de Search Console.
    impresiones_totales = round(sum(float(fila.get("impresiones", 0.0)) for fila in filas_gsc_paginas), 2)

    # Calcula CTR medio ponderado por impresiones.
    ctr_medio = round((clics_totales / max(1.0, impresiones_totales)), 4) if impresiones_totales > 0 else 0.0

    # Calcula posición media de páginas con datos GSC.
    posicion_media = (
        round(sum(float(fila.get("posicion_media", 0.0)) for fila in filas_gsc_paginas) / max(1, len(filas_gsc_paginas)), 2)
        if filas_gsc_paginas
        else 0.0
    )

    # Devuelve conjunto de métricas agregadas de GSC.
    return {
        "clics_totales": clics_totales,
        "impresiones_totales": impresiones_totales,
        "ctr_medio": ctr_medio,
        "posicion_media": posicion_media,
    }


# Garantiza que la carpeta de salida exista antes de escribir archivos.
def asegurar_directorio(path_salida: Path) -> None:
    """Crea la carpeta de salida si no existe."""

    # Crea el directorio y sus padres sin fallar si ya existen.
    path_salida.mkdir(parents=True, exist_ok=True)


# Calcula métricas ejecutivas agregadas para informes.
def calcular_metricas(resultado: ResultadoAuditoria) -> dict[str, int | float | dict[str, int] | str]:
    """Calcula métricas ejecutivas reutilizables por todos los reportes."""

    # Inicializa acumuladores principales.
    total_incidencias = 0

    # Inicializa contadores de URLs.
    urls_sanas = 0

    # Inicializa distribución por severidad.
    severidades: Counter[str] = Counter()

    # Inicializa distribución por tipo.
    tipos: Counter[str] = Counter()

    # Inicializa distribución por área.
    areas: Counter[str] = Counter()

    # Recorre resultados para consolidar datos.
    for item in resultado.resultados:
        # Suma incidencias de la URL actual.
        total_incidencias += len(item.hallazgos)

        # Cuenta URL sana cuando no hay hallazgos.
        if not item.hallazgos:
            # Incrementa contador de URLs sin incidencias.
            urls_sanas += 1

        # Recorre hallazgos para completar distribuciones.
        for hallazgo in item.hallazgos:
            # Incrementa la severidad detectada.
            severidades[hallazgo.severidad] += 1

            # Incrementa el tipo detectado.
            tipos[hallazgo.tipo] += 1

            # Incrementa el área detectada.
            areas[hallazgo.area] += 1

    # Define pesos de scoring por severidad.
    pesos = {"crítica": 10.0, "alta": 6.0, "media": 3.0, "baja": 1.0, "informativa": 0.5}

    # Calcula penalización ponderada total.
    penalizacion = sum(severidades.get(severidad, 0) * peso for severidad, peso in pesos.items())

    # Calcula máximo teórico por volumen auditado.
    max_penalizacion = max(10.0, float(resultado.total_urls) * 10.0)

    # Calcula score SEO global acotado entre 5 y 100.
    score = round(max(5.0, min(100.0, 100.0 - (penalizacion / max_penalizacion) * 100.0)), 1)

    # Construye filas técnicas para agregación ejecutiva coherente.
    filas = construir_filas(resultado)

    # Construye filas GSC por página para métricas agregadas.
    filas_gsc_paginas = construir_filas_search_console_paginas(resultado)

    # Construye filas de gestión de indexación para KPIs.
    filas_gestion_indexacion = construir_filas_gestion_indexacion(resultado)

    # Calcula métricas agregadas de Search Console.
    metricas_gsc = _calcular_metricas_gsc(filas_gsc_paginas)

    # Obtiene métricas agregadas de Analytics desde resumen ya calculado en la carga.
    metricas_analytics = (
        {
            "sesiones_totales": float(resultado.analytics.resumen.sesiones_totales),
            "usuarios_totales": float(resultado.analytics.resumen.usuarios_totales),
            "rebote_medio": float(resultado.analytics.resumen.rebote_medio),
            "duracion_media": float(resultado.analytics.resumen.duracion_media),
            "conversiones": float(resultado.analytics.resumen.conversiones),
        }
        if resultado.analytics.activo
        else {"sesiones_totales": 0.0, "usuarios_totales": 0.0, "rebote_medio": 0.0, "duracion_media": 0.0, "conversiones": 0.0}
    )

    # Calcula incidencias agrupadas para capa ejecutiva.
    incidencias_agrupadas = _calcular_incidencias_agrupadas(filas)

    # Cuenta total de incidencias agrupadas para resumen ejecutivo.
    total_agrupadas = sum(incidencias_agrupadas.values())

    # Calcula score de bloque de indexación y arquitectura.
    score_indexacion = round(max(5.0, min(100.0, 100.0 - ((areas.get("Indexación", 0) + areas.get("Arquitectura", 0)) * 4.0))), 1)

    # Calcula score de bloque de contenido on-page.
    score_contenido = round(max(5.0, min(100.0, 100.0 - (areas.get("Contenido", 0) * 3.5))), 1)

    # Calcula score de bloque multimedia/accesibilidad básica.
    score_multimedia = round(max(5.0, min(100.0, 100.0 - (incidencias_agrupadas.get("Imágenes sin ALT", 0) * 5.0))), 1)

    # Obtiene scores de rendimiento disponibles.
    scores_rendimiento = [item.performance_score for item in resultado.rendimiento if isinstance(item.performance_score, (int, float))]

    # Calcula score de rendimiento con fallback al global si no hay datos.
    score_rendimiento = round(sum(scores_rendimiento) / len(scores_rendimiento), 1) if scores_rendimiento else float(resultado.score_rendimiento or score)

    # Resuelve score técnico preferente desde resultado consolidado.
    score_tecnico = float(resultado.score_tecnico) if isinstance(resultado.score_tecnico, (int, float)) else score_indexacion

    # Resuelve score de contenido preferente desde resultado consolidado.
    score_contenido_final = float(resultado.score_contenido) if isinstance(resultado.score_contenido, (int, float)) else score_contenido

    # Resuelve score global preferente desde resultado consolidado.
    seo_score_global = (
        float(resultado.seo_score_global)
        if isinstance(resultado.seo_score_global, (int, float))
        else round((score_tecnico * 0.4) + (score_contenido_final * 0.4) + (score_rendimiento * 0.2), 1)
    )

    # Devuelve métricas calculadas.
    return {
        "total_urls": resultado.total_urls,
        "total_incidencias": total_incidencias,
        "total_incidencias_agrupadas": total_agrupadas,
        "incidencias_agrupadas": incidencias_agrupadas,
        "severidades": dict(severidades),
        "tipos": dict(tipos),
        "areas": dict(areas),
        "urls_sanas": urls_sanas,
        "score": seo_score_global,
        "score_tecnico": score_tecnico,
        "score_contenido": score_contenido_final,
        "score_rendimiento": score_rendimiento,
        "score_bloques": {
            "indexacion_arquitectura": {"score": score_indexacion, "peso": 0.35},
            "contenido_onpage": {"score": score_contenido, "peso": 0.30},
            "rendimiento": {"score": score_rendimiento, "peso": 0.20},
            "multimedia_accesibilidad": {"score": score_multimedia, "peso": 0.15},
        },
        "gsc": metricas_gsc,
        "analytics": metricas_analytics,
        "formula_score": "Score = 100 - (penalización_ponderada/(total_urls*10))*100",
        "gestion_indexacion": {
            "indexable": len([fila for fila in filas_gestion_indexacion if fila.get("clasificacion") == "INDEXABLE"]),
            "revisar": len([fila for fila in filas_gestion_indexacion if fila.get("clasificacion") == "REVISAR"]),
            "no_indexar": len([fila for fila in filas_gestion_indexacion if fila.get("clasificacion") == "NO_INDEXAR"]),
        },
    }


# Construye la narrativa final sin duplicar secciones.
def _construir_bloques_narrativos(resultado: ResultadoAuditoria) -> dict[str, list[str]]:
    """Genera bloques narrativos únicos siguiendo jerarquía editorial obligatoria."""

    # Construye diccionario base con todas las secciones requeridas.
    bloques = {titulo: [] for titulo in JERARQUIA_INFORME}

    # Obtiene secciones procesadas desde IA cuando existan.
    secciones_ia = construir_secciones_desde_ia(resultado.resumen_ia)

    # Recorre secciones IA para mapearlas a la jerarquía interna.
    for seccion in secciones_ia:
        # Obtiene título normalizado de sección.
        titulo = str(seccion["titulo"]).strip().lower()

        # Recorre jerarquía para encontrar destino compatible.
        for destino in construir_jerarquia_visible(resultado):
            # Compara por inclusión simple para robustez.
            if destino.lower() in titulo or titulo in destino.lower():
                # Inserta solo textos no vacíos en el bloque destino.
                bloques[destino].extend([str(item).strip() for item in seccion["items"] if str(item).strip()])

    # Construye vistas tabulares para generar fallback de secciones obligatorias.
    filas = construir_filas(resultado)

    # Construye vista tabular de rendimiento para secciones de UX.
    filas_rendimiento = construir_filas_rendimiento(resultado)

    # Calcula métricas agregadas para reutilización narrativa.
    metricas = calcular_metricas(resultado)

    # Obtiene métricas agregadas de Search Console.
    metricas_gsc = metricas.get("gsc", {}) if isinstance(metricas.get("gsc", {}), dict) else {}

    # Obtiene métricas agregadas de Analytics.
    metricas_analytics = metricas.get("analytics", {}) if isinstance(metricas.get("analytics", {}), dict) else {}

    # Construye filas de Analytics por página para insights narrativos.
    filas_analytics_paginas = construir_filas_analytics_paginas(resultado)

    # Construye cruces GSC + Analytics cuando ambas fuentes estén activas.
    filas_cruce_gsc_analytics = construir_cruces_gsc_analytics(resultado)

    # Resuelve periodo efectivo con helper común para evitar duplicación.
    periodo_desde, periodo_hasta = _resolver_periodo_analizado(resultado)

    # Conserva estado previo para no anular fallback de resumen.
    resumen_tenia_contenido = bool(bloques["Resumen ejecutivo"])

    # Construye línea estándar de periodo cuando exista información temporal.
    linea_periodo = f"Periodo analizado: {periodo_desde} - {periodo_hasta}" if periodo_desde and periodo_hasta else ""

    # Inserta periodo en resumen ejecutivo una única vez cuando aplique.
    if linea_periodo and linea_periodo not in bloques["Resumen ejecutivo"]:
        # Prioriza mostrar el periodo al inicio del resumen.
        bloques["Resumen ejecutivo"].insert(0, linea_periodo)

    # Construye fallback de resumen ejecutivo cuando IA no aporta contenido.
    if not resumen_tenia_contenido:
        # Añade resumen automático con datos técnicos.
        bloques["Resumen ejecutivo"].append(f"Se auditaron {resultado.total_urls} URLs con fuentes activas: {', '.join(resultado.fuentes_activas)}.")

    # Construye fallback de KPIs principales para mantener consistencia de jerarquía.
    if not bloques["KPIs principales"]:
        # Añade recordatorio de que los KPIs se muestran en tabla.
        bloques["KPIs principales"].append("Los indicadores clave se presentan en la tabla KPI de esta sección.")

    # Construye fallback de visibilidad orgánica real cuando GSC esté activo.
    if not bloques["Visibilidad orgánica real"] and resultado.search_console.activo:
        # Construye filas de páginas de Search Console.
        filas_gsc_paginas = construir_filas_search_console_paginas(resultado)

        # Construye filas de queries de Search Console.
        filas_gsc_queries = construir_filas_search_console_queries(resultado)

        # Obtiene clics totales desde métricas centralizadas.
        clics_totales = float(metricas_gsc.get("clics_totales", 0.0))

        # Obtiene impresiones totales desde métricas centralizadas.
        impresiones_totales = float(metricas_gsc.get("impresiones_totales", 0.0))

        # Obtiene CTR medio desde métricas centralizadas.
        ctr_medio = float(metricas_gsc.get("ctr_medio", 0.0))

        # Obtiene posición media desde métricas centralizadas.
        posicion_media = float(metricas_gsc.get("posicion_media", 0.0))

        # Inserta resumen de visibilidad principal.
        bloques["Visibilidad orgánica real"].append(f"Clics: {clics_totales} | Impresiones: {impresiones_totales} | CTR medio: {ctr_medio} | Posición media: {posicion_media}.")

        # Inserta top páginas por impresiones.
        for fila in sorted(filas_gsc_paginas, key=lambda item: float(item.get('impresiones', 0.0)), reverse=True)[:5]:
            # Añade línea de top página.
            bloques["Visibilidad orgánica real"].append(f"Top página: {fila.get('url', '')} | impresiones={fila.get('impresiones', 0)} | clics={fila.get('clics', 0)}.")

        # Inserta top queries por impresiones.
        for fila in sorted(filas_gsc_queries, key=lambda item: float(item.get('impresiones', 0.0)), reverse=True)[:5]:
            # Añade línea de top query.
            bloques["Visibilidad orgánica real"].append(f"Top query: {fila.get('query', '')} | impresiones={fila.get('impresiones', 0)} | clics={fila.get('clics', 0)}.")

    # Construye fallback de oportunidades SEO prioritarias cuando GSC esté activo.
    if not bloques["Oportunidades SEO prioritarias"] and resultado.search_console.activo:
        # Obtiene oportunidades GSC cruzadas.
        oportunidades_gsc = construir_oportunidades_gsc(resultado)

        # Recorre oportunidades priorizadas para narrativa.
        for fila in oportunidades_gsc[:8]:
            # Añade oportunidad resumida con acción.
            bloques["Oportunidades SEO prioritarias"].append(
                f"{fila.get('url', '')}: {fila.get('oportunidad', '')} | CTR={fila.get('ctr', 0)} | posición={fila.get('posicion_media', 0)} | acción={fila.get('accion_recomendada', '')}"
            )

    # Construye fallback de cruce técnico + Search Console.
    if not bloques["Cruce auditoría técnica + Search Console"] and resultado.search_console.activo:
        # Obtiene oportunidades GSC cruzadas.
        oportunidades_gsc = construir_oportunidades_gsc(resultado)

        # Recorre oportunidades para explicar cruce técnico.
        for fila in oportunidades_gsc[:8]:
            # Añade lectura de cruce técnico-negocio.
            bloques["Cruce auditoría técnica + Search Console"].append(
                f"{fila.get('url', '')}: visibilidad real con {fila.get('problema_onpage_detectado', 'sin problema')}."
            )

    # Construye fallback de keyword/query mapping inicial.
    if not bloques["Keyword / query mapping inicial"] and resultado.search_console.activo:
        # Recorre filas de query para mostrar mapping inicial.
        for fila in construir_filas_search_console_queries(resultado)[:10]:
            # Añade línea de mapping básico.
            bloques["Keyword / query mapping inicial"].append(
                f"query={fila.get('query', '')} | clics={fila.get('clics', 0)} | impresiones={fila.get('impresiones', 0)} | posición={fila.get('posicion_media', 0)}"
            )

    # Construye fallback de hallazgos críticos cuando no exista IA útil.
    if not bloques["Hallazgos críticos"]:
        # Filtra hallazgos críticos para el bloque ejecutivo.
        hallazgos_criticos = [fila for fila in filas if str(fila.get("severidad", "")).lower() in {"crítica", "alta"} and fila.get("problema")]

        # Inserta resumen compacto de hallazgos prioritarios.
        for fila in hallazgos_criticos[:5]:
            # Añade línea ejecutiva de hallazgo.
            bloques["Hallazgos críticos"].append(f"[{fila['severidad']}] {fila['problema']} ({fila['url']})")

    # Construye fallback de quick wins cuando no exista narrativa IA.
    if not bloques["Quick wins"]:
        # Calcula quick wins deduplicados y accionables.
        quick_wins = _construir_quick_wins(construir_filas_quick_wins(resultado), limite=8)

        # Inserta recomendaciones rápidas priorizadas.
        for fila in quick_wins:
            # Añade quick win concreto.
            bloques["Quick wins"].append(f"{', '.join(fila['recomendaciones'][:2])} ({fila['url']})")

        # Obtiene decisiones de indexación de alta prioridad para quick wins.
        quick_wins_indexacion = [
            fila for fila in construir_filas_gestion_indexacion(resultado) if fila.get("clasificacion") in {"NO_INDEXAR", "REVISAR"} and fila.get("prioridad") in {"Alta", "Media"}
        ]

        # Añade quick wins de indexación inteligente.
        for fila in quick_wins_indexacion[:4]:
            # Inserta quick win de gestión de indexación.
            bloques["Quick wins"].append(f"[Indexación] {fila.get('accion_recomendada', '')} ({fila.get('url', '')})")

        # Añade quick wins basados en señales de Analytics.
        for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("sesiones", 0.0)), reverse=True):
            # Detecta páginas con tráfico alto y mala experiencia.
            if (
                float(fila.get("sesiones", 0.0)) >= QUICK_WIN_GA_MIN_SESIONES_REBOTE
                and float(fila.get("rebote", 0.0)) >= QUICK_WIN_GA_MIN_REBOTE
            ):
                bloques["Quick wins"].append(f"[Analytics] Mejorar UX en {fila.get('url', '')} por rebote elevado.")

            # Detecta páginas con tráfico sin conversión.
            if (
                float(fila.get("sesiones", 0.0)) >= QUICK_WIN_GA_MIN_SESIONES_CONVERSION
                and float(fila.get("conversiones", 0.0)) <= QUICK_WIN_GA_MAX_CONVERSIONES
            ):
                bloques["Quick wins"].append(f"[Analytics] Optimizar contenido/CTA en {fila.get('url', '')} por ausencia de conversiones.")

            # Limita tamaño final de quick wins para mantener legibilidad.
            if len(bloques["Quick wins"]) >= QUICK_WINS_BLOQUE_MAXIMO:
                break

    # Construye fallback de acciones técnicas cuando no haya bloque IA.
    if not bloques["Acciones técnicas"]:
        # Filtra recomendaciones técnicas.
        acciones_tecnicas = [fila for fila in filas if fila.get("area") in {"Infraestructura", "Indexación", "Arquitectura"} and fila.get("recomendacion")]

        # Inserta recomendaciones técnicas de mayor prioridad.
        for fila in acciones_tecnicas[:5]:
            # Añade acción técnica priorizada.
            bloques["Acciones técnicas"].append(f"{fila['prioridad']}: {fila['recomendacion']}")

        # Añade fallback genérico cuando no existan acciones técnicas específicas.
        if not bloques["Acciones técnicas"]:
            # Inserta recomendación de revisión técnica base.
            bloques["Acciones técnicas"].append("Revisar cobertura de rastreo, estado HTTP y canonicals para mantener estabilidad técnica.")

    # Construye fallback de acciones de contenido cuando no haya bloque IA.
    if not bloques["Acciones de contenido"]:
        # Filtra recomendaciones del área de contenido.
        acciones_contenido = [fila for fila in filas if fila.get("area") == "Contenido" and fila.get("recomendacion")]

        # Inserta recomendaciones editoriales de mayor prioridad.
        for fila in acciones_contenido[:5]:
            # Añade acción de contenido priorizada.
            bloques["Acciones de contenido"].append(f"{fila['prioridad']}: {fila['recomendacion']}")

    # Construye fallback de rendimiento y experiencia de usuario.
    if not bloques["Rendimiento y experiencia de usuario"]:
        # Filtra filas con métricas reales de rendimiento.
        filas_metricas_validas = [fila for fila in filas_rendimiento if isinstance(fila.get("performance_score"), (int, float))]

        # Inserta resumen de PageSpeed cuando exista información válida.
        for fila in filas_metricas_validas[:5]:
            # Añade línea compacta con score y oportunidad.
            bloques["Rendimiento y experiencia de usuario"].append(f"{fila['url']} [{fila['estrategia']}] score={fila['performance_score']} oportunidad={fila['oportunidad'] or 'sin oportunidad destacada'}")

        # Añade mensaje profesional cuando no hay datos válidos.
        if not bloques["Rendimiento y experiencia de usuario"]:
            # Determina si hubo fallos explícitos de PageSpeed.
            hubo_fallo_pagespeed = "pagespeed" in resultado.fuentes_fallidas or bool(resultado.pagespeed_estado)

            # Inserta mensaje específico según disponibilidad.
            bloques["Rendimiento y experiencia de usuario"].append(
                "No se pudieron obtener métricas de PageSpeed en esta ejecución por timeout o error de la API."
                if hubo_fallo_pagespeed
                else "No se han recibido datos de PageSpeed en esta ejecución."
            )

    # Construye fallback de comportamiento y conversión cuando Analytics esté activo.
    if not bloques["Comportamiento y conversión"] and resultado.analytics.activo:
        # Inserta resumen agregado de comportamiento.
        bloques["Comportamiento y conversión"].append(
            "Sesiones: "
            f"{metricas_analytics.get('sesiones_totales', 0.0)} | "
            f"Usuarios: {metricas_analytics.get('usuarios_totales', 0.0)} | "
            f"Rebote medio: {metricas_analytics.get('rebote_medio', 0.0)} | "
            f"Duración media: {metricas_analytics.get('duracion_media', 0.0)}s | "
            f"Conversiones: {metricas_analytics.get('conversiones', 0.0)}."
        )

        # Añade páginas con más tráfico.
        for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("sesiones", 0.0)), reverse=True)[:ANALYTICS_TOP_TRAFICO_LIMITE]:
            bloques["Comportamiento y conversión"].append(
                f"Top tráfico: {fila.get('url', '')} | sesiones={fila.get('sesiones', 0)} | rebote={fila.get('rebote', 0)}."
            )

        # Añade páginas con peor rebote.
        for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("rebote", 0.0)), reverse=True)[:ANALYTICS_PEORES_REBOTE_LIMITE]:
            bloques["Comportamiento y conversión"].append(
                f"Peor rebote: {fila.get('url', '')} | rebote={fila.get('rebote', 0)} | sesiones={fila.get('sesiones', 0)}."
            )

        # Añade páginas con mejor engagement.
        mejores_engagement = sorted(
            filas_analytics_paginas,
            key=lambda item: (float(item.get("duracion", 0.0)), -float(item.get("rebote", 1.0))),
            reverse=True,
        )[:ANALYTICS_MEJORES_ENGAGEMENT_LIMITE]
        for fila in mejores_engagement:
            bloques["Comportamiento y conversión"].append(
                f"Mejor engagement: {fila.get('url', '')} | duración={fila.get('duracion', 0)}s | rebote={fila.get('rebote', 0)}."
            )

        # Añade páginas con tráfico y sin conversión como foco de negocio.
        contador_sin_conversion = 0
        for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("sesiones", 0.0)), reverse=True):
            if float(fila.get("sesiones", 0.0)) >= CRUCE_GA_MIN_SESIONES_SIN_CONVERSION and float(fila.get("conversiones", 0.0)) <= 0.0:
                bloques["Comportamiento y conversión"].append(
                    f"Tráfico sin conversión: {fila.get('url', '')} | sesiones={fila.get('sesiones', 0)} | conversiones={fila.get('conversiones', 0)}."
                )
                contador_sin_conversion += 1
                if contador_sin_conversion >= 3:
                    break

        # Añade insights derivados del cruce GSC + Analytics.
        for fila in filas_cruce_gsc_analytics[:ANALYTICS_CRUCES_LIMITE]:
            for insight in fila.get("insights", []):
                bloques["Comportamiento y conversión"].append(f"{fila.get('url', '')}: {insight}")

    # Construye fallback de páginas prioritarias combinando señales SEO y negocio.
    if not bloques["Páginas prioritarias"]:
        # Obtiene shortlist consolidada de páginas a priorizar.
        paginas_prioritarias = construir_paginas_prioritarias(resultado, limite=8)

        # Recorre páginas para justificar priorización.
        for pagina in paginas_prioritarias:
            bloques["Páginas prioritarias"].append(
                f"{pagina.get('url', '')} | score_prioridad={pagina.get('prioridad_score', 0)} | "
                f"impresiones={pagina.get('impresiones', 0)} | CTR={pagina.get('ctr', 0)} | "
                f"posición={pagina.get('posicion_media', 0)} | sesiones={pagina.get('sesiones', 0)} | "
                f"conversiones={pagina.get('conversiones', 0)} | motivos={'; '.join(pagina.get('motivos', []))}"
            )

    # Construye fallback de indexación y rastreo con datos técnicos reales.
    if not bloques["Indexación y rastreo"]:
        # Obtiene resumen de indexación si está disponible.
        resumen_indexacion = resultado.indexacion_rastreo if isinstance(resultado.indexacion_rastreo, dict) else {}

        # Lee disponibilidad de robots en formato seguro.
        robots_disponible = bool(resumen_indexacion.get("robots_disponible", False))

        # Lee volumen de URLs bloqueadas detectadas.
        urls_bloqueadas = resumen_indexacion.get("urls_bloqueadas", [])

        # Lee volumen de incoherencias sitemap vs robots.
        incoherencias = resumen_indexacion.get("incoherencias_sitemap_robots", [])

        # Inserta línea de estado principal de rastreo.
        bloques["Indexación y rastreo"].append(f"robots.txt disponible: {'sí' if robots_disponible else 'no'}; URLs bloqueadas: {len(urls_bloqueadas)}.")

        # Inserta incoherencias cuando existan.
        if incoherencias:
            # Añade recomendación concreta basada en incoherencias.
            bloques["Indexación y rastreo"].append("Existen incoherencias entre sitemap y robots: revisar URLs bloqueadas listadas en sitemap.")
        else:
            # Añade estado estable cuando no hay incoherencias detectadas.
            bloques["Indexación y rastreo"].append("No se detectaron incoherencias críticas entre sitemap y robots en esta ejecución.")

    # Construye fallback de gestión de indexación inteligente.
    if not bloques["Gestión de indexación"]:
        # Inicializa agrupación por clasificación para evitar múltiples pasadas.
        filas_por_clasificacion: dict[str, list[dict]] = {"NO_INDEXAR": [], "REVISAR": [], "INDEXABLE": []}

        # Recorre filas de gestión de indexación para agruparlas en una pasada.
        for fila in construir_filas_gestion_indexacion(resultado):
            # Obtiene clasificación de la fila actual.
            clasificacion = str(fila.get("clasificacion", "")).strip()

            # Añade la fila cuando la clasificación sea reconocida.
            if clasificacion in filas_por_clasificacion:
                # Inserta fila en su clasificación correspondiente.
                filas_por_clasificacion[clasificacion].append(fila)

        # Calcula totales por clasificación.
        total_no_indexar = len(filas_por_clasificacion["NO_INDEXAR"])
        total_revisar = len(filas_por_clasificacion["REVISAR"])
        total_indexable = len(filas_por_clasificacion["INDEXABLE"])

        # Añade resumen global de gestión de indexación.
        bloques["Gestión de indexación"].append(
            f"Resumen global: indexables={total_indexable}, revisar={total_revisar}, no indexar={total_no_indexar}."
        )

        # Inserta bloque de no indexables con prioridad.
        for fila in filas_por_clasificacion["NO_INDEXAR"][:5]:
            # Añade línea de URL no indexable.
            bloques["Gestión de indexación"].append(f"NO_INDEXAR: {fila.get('url', '')} | motivo={fila.get('motivo', '')}.")

        # Inserta bloque de URLs a revisar.
        for fila in filas_por_clasificacion["REVISAR"][:5]:
            # Añade línea de URL a revisar.
            bloques["Gestión de indexación"].append(f"REVISAR: {fila.get('url', '')} | motivo={fila.get('motivo', '')}.")

        # Inserta recomendación clara de ejecución.
        bloques["Gestión de indexación"].append(
            "Recomendación: priorizar exclusión de URLs transaccionales y corregir señales SEO en URLs con potencial antes de solicitar indexación."
        )

    # Construye fallback de roadmap cuando IA no lo entregue.
    if not bloques["Roadmap"]:
        # Añade fase 1 centrada en quick wins con impacto real.
        bloques["Roadmap"].append("Fase 1 (0-30 días): quick wins técnicos con impacto real y corrección de páginas con impresiones altas y problemas on-page.")
        # Añade fase 2 centrada en crecimiento de posiciones y CTR.
        bloques["Roadmap"].append("Fase 2 (31-60 días): optimización de páginas en posición media 4-15, mejora de CTR y ampliación de contenido con mayor oportunidad.")
        # Añade fase 3 centrada en estrategia escalable.
        bloques["Roadmap"].append("Fase 3 (61-90 días): estrategia de contenido, arquitectura temática, enlazado interno y monitorización continua.")
        # Añade fase específica de gobernanza de indexación.
        bloques["Roadmap"].append("Fase 2.5 (45-75 días): implantar gestión de indexación inteligente con limpieza de URLs no indexables y revisión priorizada de URLs en estado REVISAR.")

    # Verifica presencia de fase de medio plazo en roadmap.
    contiene_medio_plazo = any("60" in linea or "medio plazo" in linea.lower() for linea in bloques["Roadmap"])

    # Añade fallback obligatorio de medio plazo cuando falte.
    if not contiene_medio_plazo:
        # Inserta bloque estándar de medio plazo con foco en rendimiento.
        bloques["Roadmap"].append("Fase Medio Plazo: optimización de Time to Interactive, reducción de recursos JS/CSS, mejora de LCP mobile y optimización de renderizado.")

    # Devuelve bloques listos para render narrativo.
    return bloques


# Exporta el resultado técnico en formato JSON.
def exportar_json(resultado: ResultadoAuditoria, path_salida: Path) -> Path:
    """Genera un JSON técnico con todos los resultados y metadatos."""

    # Garantiza la existencia de la carpeta de salida.
    asegurar_directorio(path_salida)

    # Define la ruta del archivo JSON técnico.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}_tecnico.json"

    # Construye contenido serializable completo.
    contenido = {
        "sitemap": resultado.sitemap,
        "cliente": resultado.cliente,
        "gestor": resultado.gestor,
        "fecha_ejecucion": resultado.fecha_ejecucion,
        "fuentes_activas": resultado.fuentes_activas,
        "fuentes_fallidas": resultado.fuentes_fallidas,
        "total_urls": resultado.total_urls,
        "resumen_ia": resultado.resumen_ia,
        "metricas": calcular_metricas(resultado),
        "resultados": construir_filas(resultado),
        "rendimiento": construir_filas_rendimiento(resultado),
        "search_console_paginas": construir_filas_search_console_paginas(resultado),
        "search_console_queries": construir_filas_search_console_queries(resultado),
        "analytics": {
            "paginas": construir_filas_analytics_paginas(resultado),
            "resumen": {
                "sesiones_totales": resultado.analytics.resumen.sesiones_totales,
                "usuarios_totales": resultado.analytics.resumen.usuarios_totales,
                "rebote_medio": resultado.analytics.resumen.rebote_medio,
                "duracion_media": resultado.analytics.resumen.duracion_media,
                "conversiones": resultado.analytics.resumen.conversiones,
            },
        },
        "cruce_gsc_analytics": construir_cruces_gsc_analytics(resultado),
        "oportunidades_gsc": construir_oportunidades_gsc(resultado),
        "gestion_indexacion": construir_filas_gestion_indexacion(resultado),
        "pagespeed_estado": resultado.pagespeed_estado,
    }

    # Escribe el JSON con codificación UTF-8 legible.
    destino.write_text(json.dumps(contenido, ensure_ascii=False, indent=2), encoding="utf-8")

    # Devuelve la ruta final del archivo generado.
    return destino


# Exporta el detalle tabular a Excel.
def exportar_excel(resultado: ResultadoAuditoria, path_salida: Path) -> Path:
    """Genera un Excel profesional con dashboard, hoja de errores y hoja de rendimiento."""

    # Garantiza la existencia de la carpeta de salida.
    asegurar_directorio(path_salida)

    # Define la ruta del archivo Excel final.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}.xlsx"

    # Crea libro de trabajo.
    libro = Workbook()

    # Prepara hoja de KPIs como primera pestaña ejecutiva.
    hoja_kpis = libro.active

    # Renombra hoja principal a KPIs.
    hoja_kpis.title = "KPIs"

    # Crea hoja dashboard analítica como segunda pestaña.
    hoja_dashboard = libro.create_sheet("Dashboard")

    # Crea hojas auxiliares de trabajo.
    hoja_errores = libro.create_sheet("Errores")

    # Crea hoja específica de contenido on-page.
    hoja_contenido = libro.create_sheet("Contenido")

    # Crea hoja específica de rendimiento.
    hoja_rendimiento = libro.create_sheet("Rendimiento")

    # Crea hoja específica de gestión de indexación inteligente.
    hoja_indexacion = libro.create_sheet("Indexacion")

    # Crea hoja específica de Analytics.
    hoja_analytics = libro.create_sheet("Analytics")

    # Crea hoja de métricas Search Console por páginas.
    hoja_gsc_paginas = libro.create_sheet("Search_Console_Paginas")

    # Crea hoja de métricas Search Console por queries.
    hoja_gsc_queries = libro.create_sheet("Search_Console_Queries")

    # Crea hoja de oportunidades basadas en Search Console.
    hoja_oportunidades_gsc = libro.create_sheet("Oportunidades_GSC")

    # Crea hoja auxiliar para KPIs y rangos de gráficos.
    hoja_aux = libro.create_sheet("AuxDashboard")

    # Construye filas técnicas y de rendimiento.
    filas = construir_filas(resultado)

    # Construye filas de contenido consolidadas por URL.
    filas_contenido = construir_filas_contenido_consolidado(resultado)

    # Construye filas de rendimiento.
    filas_rendimiento = construir_filas_rendimiento(resultado)

    # Construye filas de Search Console por página.
    filas_gsc_paginas = construir_filas_search_console_paginas(resultado)

    # Construye filas de Search Console por query.
    filas_gsc_queries = construir_filas_search_console_queries(resultado)

    # Construye filas de Analytics por página.
    filas_analytics_paginas = construir_filas_analytics_paginas(resultado)

    # Construye oportunidades cruzadas entre técnica y GSC.
    filas_oportunidades_gsc = construir_oportunidades_gsc(resultado)

    # Construye filas de gestión de indexación inteligente.
    filas_gestion_indexacion = construir_filas_gestion_indexacion(resultado)

    # Define columnas fijas de la hoja de errores para mantener estructura estable.
    columnas_errores = ["url", "url_final", "tipo", "estado_http", "redirecciona", "title", "h1", "meta_description", "canonical", "noindex", "problema", "recomendacion", "severidad", "categoria", "area", "impacto", "esfuerzo", "prioridad", "estado", "resuelto", "responsable", "observaciones"]

    # Usa columnas fijas como encabezados para escritura tabular.
    encabezados = columnas_errores

    # Recorre encabezados de errores.
    for columna, encabezado in enumerate(encabezados, start=1):
        # Escribe cada encabezado.
        hoja_errores.cell(row=1, column=columna, value=encabezado)

    # Escribe contenido de errores.
    for fila_indice, fila in enumerate(filas, start=2):
        # Recorre columnas de la fila.
        for columna, encabezado in enumerate(encabezados, start=1):
            # Escribe valor de celda.
            hoja_errores.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Aplica formato básico de encabezado en errores.
    for celda in hoja_errores[1]:
        # Define estilo de encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Configura ancho de columnas de errores con foco en legibilidad.
    anchos_errores = {"A": 48, "B": 48, "C": 14, "D": 12, "E": 12, "F": 26, "G": 26, "H": 36, "I": 36, "J": 10, "K": 52, "L": 52, "M": 12, "N": 14, "O": 18, "P": 12, "Q": 12, "R": 12, "S": 14, "T": 10, "U": 20, "V": 36}

    # Recorre anchos definidos y los aplica.
    for columna, ancho in anchos_errores.items():
        # Asigna ancho explícito de columna.
        hoja_errores.column_dimensions[columna].width = ancho

    # Ajusta celdas para legibilidad.
    for fila in hoja_errores.iter_rows(min_row=2, max_row=max(2, len(filas) + 1), min_col=1, max_col=max(1, len(encabezados))):
        # Recorre celdas de la fila.
        for celda in fila:
            # Activa ajuste automático de texto.
            celda.alignment = Alignment(wrap_text=True, vertical="top")

    # Ajusta altura de filas para lectura de textos largos.
    for indice_fila in range(2, max(3, len(filas) + 2)):
        # Aplica altura moderada para evitar cortes.
        hoja_errores.row_dimensions[indice_fila].height = 38

    # Aplica color suave por severidad en la hoja de errores.
    colores_severidad = {"crítica": "FADBD8", "alta": "FADBD8", "media": "F8C471", "baja": "FCF3CF", "informativa": "D6EAF8"}

    # Recorre filas para aplicar color según severidad.
    for indice_fila in range(2, max(3, len(filas) + 2)):
        # Calcula índice de columna de severidad de forma estable.
        indice_columna_severidad = columnas_errores.index("severidad") + 1

        # Obtiene severidad desde su columna real.
        severidad = str(hoja_errores.cell(row=indice_fila, column=indice_columna_severidad).value or "").lower()

        # Resuelve color asociado a severidad.
        color = colores_severidad.get(severidad, "FFFFFF")

        # Recorre columnas de la fila para aplicar el color.
        for indice_columna in range(1, max(1, len(encabezados)) + 1):
            # Obtiene celda actual para formato.
            celda = hoja_errores.cell(row=indice_fila, column=indice_columna)

            # Aplica relleno suave según severidad.
            celda.fill = PatternFill(fill_type="solid", fgColor=color)

    # Aplica validación de datos Sí/No en columna resuelto.
    validacion_resuelto = DataValidation(type="list", formula1='"Sí,No"', allow_blank=False)

    # Registra validación en hoja.
    hoja_errores.add_data_validation(validacion_resuelto)

    # Aplica validación al rango de seguimiento.
    validacion_resuelto.add(f"T2:T{max(2, len(filas) + 1)}")

    # Activa filtros en la hoja de errores.
    hoja_errores.auto_filter.ref = f"A1:V{max(2, len(filas) + 1)}"

    # Congela paneles para mejorar navegación.
    hoja_errores.freeze_panes = "A2"

    # Define columnas de hoja de contenido para seguimiento editorial.
    columnas_contenido = ["url", "palabras", "calidad_contenido", "h1", "title", "meta_description", "imagenes_sin_alt", "thin_content", "densidad_texto", "ratio_texto_html", "incidencias_url"]

    # Recorre encabezados de contenido.
    for columna, encabezado in enumerate(columnas_contenido, start=1):
        # Escribe encabezado en hoja de contenido.
        hoja_contenido.cell(row=1, column=columna, value=encabezado)

    # Recorre filas consolidadas para poblar hoja de contenido sin duplicados.
    for fila_indice, fila in enumerate(filas_contenido, start=2):
        # Escribe columnas de contenido en cada fila.
        for columna, encabezado in enumerate(columnas_contenido, start=1):
            # Inserta valor correspondiente con fallback vacío.
            hoja_contenido.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de contenido.
    for celda in hoja_contenido[1]:
        # Define estilo visual coherente con el resto de hojas.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo del encabezado.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Ajusta formato visual de la hoja de contenido.
    for indice in range(1, len(columnas_contenido) + 1):
        # Configura ancho homogéneo en todas las columnas.
        hoja_contenido.column_dimensions[chr(64 + indice)].width = 24

    # Activa filtros en la hoja de contenido.
    hoja_contenido.auto_filter.ref = f"A1:K{max(2, len(filas_contenido) + 1)}"

    # Congela paneles de contenido para navegación.
    hoja_contenido.freeze_panes = "A2"

    # Activa ajuste de texto y alineación superior en contenido.
    for fila in hoja_contenido.iter_rows(min_row=2, max_row=max(2, len(filas_contenido) + 1), min_col=1, max_col=len(columnas_contenido)):
        # Recorre celdas de la fila actual.
        for celda in fila:
            # Aplica alineación legible y multilinea.
            celda.alignment = Alignment(wrap_text=True, vertical="top")

    # Escribe tabla de rendimiento con esquema obligatorio.
    columnas_rendimiento = ["url", "estrategia", "performance_score", "accessibility_score", "best_practices_score", "seo_score", "lcp", "cls", "inp", "fcp", "speed_index", "oportunidad", "descripcion", "ahorro_estimado", "severidad", "recomendacion", "estado", "resuelto", "responsable", "observaciones"]

    # Recorre encabezados de rendimiento.
    for columna, encabezado in enumerate(columnas_rendimiento, start=1):
        # Escribe encabezado en hoja de rendimiento.
        hoja_rendimiento.cell(row=1, column=columna, value=encabezado)

    # Recorre filas de rendimiento para poblar hoja.
    for fila_indice, fila in enumerate(filas_rendimiento, start=2):
        # Recorre columnas obligatorias.
        for columna, encabezado in enumerate(columnas_rendimiento, start=1):
            # Escribe valor de columna o vacío por defecto.
            hoja_rendimiento.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de rendimiento.
    for celda in hoja_rendimiento[1]:
        # Define negrita de encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica fondo corporativo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Ajusta lectura de hoja rendimiento.
    for indice in range(1, len(columnas_rendimiento) + 1):
        # Asigna ancho homogéneo por columna.
        hoja_rendimiento.column_dimensions[chr(64 + indice)].width = 22

    # Activa wrap y alineación superior en rendimiento.
    for fila in hoja_rendimiento.iter_rows(min_row=2, max_row=max(2, len(filas_rendimiento) + 1), min_col=1, max_col=len(columnas_rendimiento)):
        # Recorre celdas de fila.
        for celda in fila:
            # Configura alineación legible.
            celda.alignment = Alignment(wrap_text=True, vertical="top")

    # Ajusta alturas en rendimiento para mantener legibilidad.
    for indice_fila in range(2, max(3, len(filas_rendimiento) + 2)):
        # Aplica altura de lectura a filas.
        hoja_rendimiento.row_dimensions[indice_fila].height = 34

    # Activa filtros en la hoja de rendimiento solo cuando no exista tabla.
    if not filas_rendimiento:
        # Aplica filtro simple sobre el encabezado cuando no hay tabla.
        hoja_rendimiento.auto_filter.ref = "A1:T1"

    # Congela paneles en rendimiento.
    hoja_rendimiento.freeze_panes = "A2"

    # Define columnas de Search Console por página.
    columnas_gsc_paginas = ["url", "clics", "impresiones", "ctr", "posicion_media", "dispositivo", "pais"]

    # Escribe encabezados de Search Console páginas.
    for columna, encabezado in enumerate(columnas_gsc_paginas, start=1):
        # Escribe encabezado en la hoja.
        hoja_gsc_paginas.cell(row=1, column=columna, value=encabezado)

    # Escribe filas de Search Console páginas.
    for fila_indice, fila in enumerate(filas_gsc_paginas, start=2):
        # Recorre columnas para poblar la hoja.
        for columna, encabezado in enumerate(columnas_gsc_paginas, start=1):
            # Inserta valor con fallback vacío.
            hoja_gsc_paginas.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de Search Console páginas.
    for celda in hoja_gsc_paginas[1]:
        # Define estilo corporativo del encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo homogéneo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Activa filtros y congelación para hoja de páginas GSC.
    hoja_gsc_paginas.auto_filter.ref = f"A1:G{max(2, len(filas_gsc_paginas) + 1)}"
    hoja_gsc_paginas.freeze_panes = "A2"

    # Define columnas de Search Console por queries.
    columnas_gsc_queries = ["query", "clics", "impresiones", "ctr", "posicion_media", "url_asociada"]

    # Escribe encabezados de Search Console queries.
    for columna, encabezado in enumerate(columnas_gsc_queries, start=1):
        # Escribe encabezado en la hoja.
        hoja_gsc_queries.cell(row=1, column=columna, value=encabezado)

    # Escribe filas de Search Console queries.
    for fila_indice, fila in enumerate(filas_gsc_queries, start=2):
        # Recorre columnas para poblar la hoja.
        for columna, encabezado in enumerate(columnas_gsc_queries, start=1):
            # Inserta valor con fallback vacío.
            hoja_gsc_queries.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de Search Console queries.
    for celda in hoja_gsc_queries[1]:
        # Define estilo corporativo del encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo homogéneo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Activa filtros y congelación para hoja de queries GSC.
    hoja_gsc_queries.auto_filter.ref = f"A1:F{max(2, len(filas_gsc_queries) + 1)}"
    hoja_gsc_queries.freeze_panes = "A2"

    # Define columnas de Analytics por página.
    columnas_analytics = ["url", "sesiones", "usuarios", "rebote", "duracion", "conversiones", "calidad_trafico"]

    # Escribe encabezados de Analytics.
    for columna, encabezado in enumerate(columnas_analytics, start=1):
        # Escribe encabezado en la hoja.
        hoja_analytics.cell(row=1, column=columna, value=encabezado)

    # Escribe filas de Analytics.
    for fila_indice, fila in enumerate(filas_analytics_paginas, start=2):
        # Recorre columnas para poblar la hoja.
        for columna, encabezado in enumerate(columnas_analytics, start=1):
            # Inserta valor con fallback vacío.
            hoja_analytics.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de Analytics.
    for celda in hoja_analytics[1]:
        # Define estilo corporativo del encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo homogéneo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Ajusta ancho de columnas para hoja Analytics.
    anchos_analytics = {"A": 56, "B": 14, "C": 14, "D": 12, "E": 12, "F": 14, "G": 16}

    # Recorre configuración de anchos y la aplica.
    for columna, ancho in anchos_analytics.items():
        # Asigna ancho explícito para la columna.
        hoja_analytics.column_dimensions[columna].width = ancho

    # Activa filtros y congelación para hoja de Analytics.
    hoja_analytics.auto_filter.ref = f"A1:G{max(2, len(filas_analytics_paginas) + 1)}"
    hoja_analytics.freeze_panes = "A2"

    # Define columnas de oportunidades GSC.
    columnas_oportunidades_gsc = [
        "url",
        "query_principal",
        "clics",
        "impresiones",
        "ctr",
        "posicion_media",
        "problema_onpage_detectado",
        "oportunidad",
        "accion_recomendada",
        "impacto",
        "esfuerzo",
        "prioridad",
        "estado",
        "resuelto",
        "responsable",
        "observaciones",
    ]

    # Escribe encabezados de oportunidades GSC.
    for columna, encabezado in enumerate(columnas_oportunidades_gsc, start=1):
        # Escribe encabezado en la hoja.
        hoja_oportunidades_gsc.cell(row=1, column=columna, value=encabezado)

    # Escribe filas de oportunidades GSC.
    for fila_indice, fila in enumerate(filas_oportunidades_gsc, start=2):
        # Recorre columnas para poblar la hoja.
        for columna, encabezado in enumerate(columnas_oportunidades_gsc, start=1):
            # Inserta valor con fallback vacío.
            hoja_oportunidades_gsc.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de oportunidades GSC.
    for celda in hoja_oportunidades_gsc[1]:
        # Define estilo corporativo del encabezado.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo homogéneo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Activa filtros y congelación para hoja de oportunidades GSC.
    hoja_oportunidades_gsc.auto_filter.ref = f"A1:P{max(2, len(filas_oportunidades_gsc) + 1)}"
    hoja_oportunidades_gsc.freeze_panes = "A2"

    # Define columnas de la hoja de indexación inteligente.
    columnas_indexacion = ["url", "clasificacion", "motivo", "accion_recomendada", "prioridad"]

    # Escribe encabezados de indexación.
    for columna, encabezado in enumerate(columnas_indexacion, start=1):
        # Escribe encabezado en la hoja.
        hoja_indexacion.cell(row=1, column=columna, value=encabezado)

    # Recorre filas de indexación para poblar la hoja.
    for fila_indice, fila in enumerate(filas_gestion_indexacion, start=2):
        # Recorre columnas de indexación.
        for columna, encabezado in enumerate(columnas_indexacion, start=1):
            # Inserta el valor correspondiente.
            hoja_indexacion.cell(row=fila_indice, column=columna, value=sanitizar_valor_excel(fila.get(encabezado, "")))

    # Estiliza encabezado de indexación.
    for celda in hoja_indexacion[1]:
        # Aplica estilo de encabezado homogéneo.
        celda.font = Font(bold=True, color="FFFFFF")

        # Aplica color corporativo.
        celda.fill = PatternFill(fill_type="solid", fgColor="1F4E78")

    # Ajusta ancho de columnas en hoja de indexación.
    anchos_indexacion = {"A": 56, "B": 16, "C": 62, "D": 62, "E": 14}

    # Recorre configuración de anchos y la aplica.
    for columna, ancho in anchos_indexacion.items():
        # Asigna ancho explícito para la columna.
        hoja_indexacion.column_dimensions[columna].width = ancho

    # Activa filtros en hoja de indexación.
    hoja_indexacion.auto_filter.ref = f"A1:E{max(2, len(filas_gestion_indexacion) + 1)}"

    # Congela cabecera para navegación.
    hoja_indexacion.freeze_panes = "A2"

    # Calcula métricas de dashboard.
    metricas = calcular_metricas(resultado)

    # Calcula scores medios mobile y desktop desde ejecuciones únicas.
    scores_mobile = [item.performance_score for item in resultado.rendimiento if item.estrategia == "mobile" and isinstance(item.performance_score, (int, float))]

    # Calcula scores de desktop desde ejecuciones únicas.
    scores_desktop = [item.performance_score for item in resultado.rendimiento if item.estrategia == "desktop" and isinstance(item.performance_score, (int, float))]

    # Obtiene media mobile segura.
    score_medio_mobile = round(sum(scores_mobile) / len(scores_mobile), 1) if scores_mobile else 0.0

    # Obtiene media desktop segura.
    score_medio_desktop = round(sum(scores_desktop) / len(scores_desktop), 1) if scores_desktop else 0.0

    # Resuelve periodo efectivo con helper común para evitar duplicación.
    periodo_desde, periodo_hasta = _resolver_periodo_analizado(resultado)

    # Construye etiqueta editorial del periodo en formato único.
    periodo_texto = f"{periodo_desde} - {periodo_hasta}" if periodo_desde and periodo_hasta else "No disponible"

    # Define KPIs del dashboard.
    hoja_dashboard["A1"] = "Dashboard SEO Ejecutivo"

    # Aplica estilo de título.
    hoja_dashboard["A1"].font = Font(size=20, bold=True, color=COLOR_DASHBOARD_TITULO)
    hoja_dashboard["A2"] = f"Cliente: {resultado.cliente} | Fecha: {resultado.fecha_ejecucion} | Periodo: {periodo_texto}"
    hoja_dashboard["A2"].font = Font(size=11, color="4B5563")

    # Calcula total de oportunidades de rendimiento.
    total_oportunidades = len([fila for fila in filas_rendimiento if fila.get("oportunidad")])

    # Calcula total de incidencias críticas.
    incidencias_criticas = metricas["severidades"].get("crítica", 0)

    # Calcula total de incidencias altas.
    incidencias_altas = metricas["severidades"].get("alta", 0)

    # Calcula total de incidencias medias.
    incidencias_medias = metricas["severidades"].get("media", 0)

    # Calcula total de incidencias bajas.
    incidencias_bajas = metricas["severidades"].get("baja", 0)

    # Calcula porcentaje de URLs con incidencias.
    porcentaje_urls_con_incidencia = round(((metricas["total_urls"] - metricas["urls_sanas"]) / max(1, metricas["total_urls"])) * 100.0, 1)

    # Calcula porcentaje de incidencias marcadas como resueltas en tracking.
    incidencias_resueltas = len([fila for fila in filas if str(fila.get("resuelto", "")).strip().lower() == "sí"])

    # Calcula porcentaje real de incidencias resueltas.
    porcentaje_resueltas = round((incidencias_resueltas / max(1, metricas["total_incidencias"])) * 100.0, 1)

    # Cuenta imágenes sin alt dentro del detalle técnico.
    total_imagenes_sin_alt = len([fila for fila in filas if "imágenes sin atributo alt" in str(fila.get("problema", "")).lower()])

    # Cuenta titles demasiado largos dentro del detalle técnico.
    total_titles_largos = len([fila for fila in filas if "title demasiado largo" in str(fila.get("problema", "")).lower()])

    # Cuenta metas vacías dentro del detalle técnico.
    total_metas_vacias = len([fila for fila in filas if "no tiene meta description" in str(fila.get("problema", "")).lower()])

    # Cuenta páginas con problemas de H1.
    total_h1_problematico = len([fila for fila in filas if "h1" in str(fila.get("problema", "")).lower()])

    # Obtiene métricas GSC centralizadas para evitar duplicación.
    metricas_gsc = metricas.get("gsc", {}) if isinstance(metricas.get("gsc", {}), dict) else {}

    # Obtiene métricas Analytics centralizadas para evitar duplicación.
    metricas_analytics = metricas.get("analytics", {}) if isinstance(metricas.get("analytics", {}), dict) else {}

    # Obtiene clics totales de GSC desde métrica centralizada.
    clics_totales_gsc = float(metricas_gsc.get("clics_totales", 0.0))

    # Obtiene impresiones totales de GSC desde métrica centralizada.
    impresiones_totales_gsc = float(metricas_gsc.get("impresiones_totales", 0.0))

    # Obtiene CTR medio de GSC desde métrica centralizada.
    ctr_medio_gsc = float(metricas_gsc.get("ctr_medio", 0.0))

    # Obtiene posición media de GSC desde métrica centralizada.
    posicion_media_gsc = float(metricas_gsc.get("posicion_media", 0.0))

    # Cuenta páginas con oportunidad real por GSC.
    total_oportunidades_gsc = len(filas_oportunidades_gsc)

    # Obtiene sesiones totales de Analytics desde métrica centralizada.
    sesiones_totales_analytics = float(metricas_analytics.get("sesiones_totales", 0.0))

    # Obtiene usuarios totales de Analytics desde métrica centralizada.
    usuarios_totales_analytics = float(metricas_analytics.get("usuarios_totales", 0.0))

    # Obtiene rebote medio de Analytics desde métrica centralizada.
    rebote_medio_analytics = float(metricas_analytics.get("rebote_medio", 0.0))

    # Obtiene duración media de Analytics desde métrica centralizada.
    duracion_media_analytics = float(metricas_analytics.get("duracion_media", 0.0))

    # Obtiene score por bloques para dashboard explicable.
    score_bloques = metricas.get("score_bloques", {}) if isinstance(metricas.get("score_bloques", {}), dict) else {}

    # Obtiene resumen de gestión de indexación para KPIs.
    resumen_gestion_indexacion = metricas.get("gestion_indexacion", {}) if isinstance(metricas.get("gestion_indexacion", {}), dict) else {}

    # Define KPIs ejecutivos mínimos obligatorios de la pestaña KPIs.
    kpis_dashboard = [
        ("Total URLs", metricas["total_urls"]),
        ("Total incidencias", metricas["total_incidencias"]),
        ("Score global", metricas["score"]),
        ("Score técnico", metricas["score_tecnico"]),
        ("Score contenido", metricas["score_contenido"]),
        ("Score rendimiento", metricas["score_rendimiento"]),
        ("Clics GSC", clics_totales_gsc),
        ("Impresiones GSC", impresiones_totales_gsc),
        ("CTR medio", ctr_medio_gsc),
        ("Posición media", posicion_media_gsc),
        ("Sesiones GA4", sesiones_totales_analytics),
        ("Usuarios GA4", usuarios_totales_analytics),
        ("Rebote medio", rebote_medio_analytics),
        ("Duración media", duracion_media_analytics),
        ("Conversiones", float(metricas_analytics.get("conversiones", 0.0))),
        ("Indexables", resumen_gestion_indexacion.get("indexable", 0)),
        ("Revisar", resumen_gestion_indexacion.get("revisar", 0)),
        ("No indexar", resumen_gestion_indexacion.get("no_indexar", 0)),
    ]

    # Renderiza cabecera de la hoja KPI ejecutiva.
    hoja_kpis["A1"] = "KPIs SEO ejecutivos"
    hoja_kpis["A1"].font = Font(size=18, bold=True, color=COLOR_DASHBOARD_TITULO)
    hoja_kpis["A2"] = f"Cliente: {resultado.cliente} | Fecha: {resultado.fecha_ejecucion} | Periodo: {periodo_texto}"
    hoja_kpis["A2"].font = Font(size=11, color="4B5563")
    hoja_kpis["A4"] = "KPI"
    hoja_kpis["B4"] = "Valor"
    hoja_kpis["A4"].font = Font(bold=True, color="FFFFFF")
    hoja_kpis["B4"].font = Font(bold=True, color="FFFFFF")
    hoja_kpis["A4"].fill = PatternFill(fill_type="solid", fgColor=COLOR_DASHBOARD_TITULO)
    hoja_kpis["B4"].fill = PatternFill(fill_type="solid", fgColor=COLOR_DASHBOARD_TITULO)

    # Escribe KPIs ejecutivos en la parte visible inicial sin scroll horizontal.
    for indice, (titulo_kpi, valor_kpi) in enumerate(kpis_dashboard, start=5):
        hoja_kpis[f"A{indice}"] = titulo_kpi
        hoja_kpis[f"B{indice}"] = valor_kpi
        hoja_kpis[f"A{indice}"].fill = PatternFill(fill_type="solid", fgColor=COLOR_KPI_TITULO_FONDO)
        hoja_kpis[f"A{indice}"].font = Font(bold=True, color=COLOR_DASHBOARD_TITULO)
        hoja_kpis[f"B{indice}"].fill = PatternFill(fill_type="solid", fgColor=COLOR_KPI_VALOR_FONDO)
        hoja_kpis[f"A{indice}"].alignment = Alignment(vertical="top", wrap_text=True)
        hoja_kpis[f"B{indice}"].alignment = Alignment(vertical="top", wrap_text=True)

    # Configura formato y navegación de la hoja KPI.
    hoja_kpis.column_dimensions["A"].width = 34
    hoja_kpis.column_dimensions["B"].width = 18
    hoja_kpis.freeze_panes = "A5"

    # Define tarjetas visuales principales en primera franja del dashboard.
    tarjetas_superiores = [
        ("Total URLs", metricas["total_urls"], "E8F0FE", "1F4E78"),
        ("Total incidencias", metricas["total_incidencias"], "FDECEC", "9B1C1C"),
        ("Score global", metricas["score"], "E7F8EF", "166534"),
        ("URLs sanas", metricas["urls_sanas"], "ECFDF3", "065F46"),
        ("CTR medio GSC", ctr_medio_gsc, "EFF6FF", "1D4ED8"),
        ("Posición media GSC", posicion_media_gsc, "FEF3C7", "92400E"),
        ("Sesiones GA4", sesiones_totales_analytics, "ECFEFF", "155E75"),
        ("Rebote GA4", rebote_medio_analytics, "FEF2F2", "991B1B"),
        ("Duración GA4", duracion_media_analytics, "F0FDF4", "166534"),
        ("Score técnico", resultado.score_tecnico, "F3E8FF", "6B21A8"),
        ("Score contenido", resultado.score_contenido, "FFF7ED", "9A3412"),
        ("Score rendimiento", resultado.score_rendimiento, "E0F2FE", "0C4A6E"),
    ]

    # Renderiza tarjetas KPI grandes en una fila superior de alto impacto visual.
    for indice, (titulo, valor, color_fondo, color_texto) in enumerate(tarjetas_superiores):
        # Calcula bloque de dos columnas por tarjeta.
        col_inicio = 1 + (indice * 2)
        col_fin = col_inicio + 1
        letra_inicio = get_column_letter(col_inicio)
        letra_fin = get_column_letter(col_fin)

        # Fusiona celdas para construir tarjeta amplia.
        hoja_dashboard.merge_cells(f"{letra_inicio}3:{letra_fin}3")
        hoja_dashboard.merge_cells(f"{letra_inicio}4:{letra_fin}6")

        # Escribe título y valor de tarjeta.
        hoja_dashboard[f"{letra_inicio}3"] = titulo
        hoja_dashboard[f"{letra_inicio}4"] = valor

        # Estiliza título de tarjeta.
        hoja_dashboard[f"{letra_inicio}3"].font = Font(size=10, bold=True, color=color_texto)
        hoja_dashboard[f"{letra_inicio}3"].alignment = Alignment(horizontal="center", vertical="center")
        hoja_dashboard[f"{letra_inicio}3"].fill = PatternFill(fill_type="solid", fgColor=color_fondo)

        # Estiliza valor principal de tarjeta.
        hoja_dashboard[f"{letra_inicio}4"].font = Font(size=18, bold=True, color=color_texto)
        hoja_dashboard[f"{letra_inicio}4"].alignment = Alignment(horizontal="center", vertical="center")
        hoja_dashboard[f"{letra_inicio}4"].fill = PatternFill(fill_type="solid", fgColor=color_fondo)

    # Obtiene top páginas por impresiones para bloque de visibilidad.
    top_paginas = ", ".join(str(fila.get("url", "")) for fila in sorted(filas_gsc_paginas, key=lambda item: float(item.get("impresiones", 0.0)), reverse=True)[:3]) or "Sin datos"

    # Obtiene top queries por impresiones para bloque de visibilidad.
    top_queries = ", ".join(str(fila.get("query", "")) for fila in sorted(filas_gsc_queries, key=lambda item: float(item.get("impresiones", 0.0)), reverse=True)[:3]) or "Sin datos"

    # Obtiene top páginas por sesiones de Analytics.
    top_paginas_analytics = ", ".join(str(fila.get("url", "")) for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("sesiones", 0.0)), reverse=True)[:3]) or "Sin datos"

    # Obtiene páginas con peor comportamiento en Analytics.
    peores_paginas_analytics = ", ".join(str(fila.get("url", "")) for fila in sorted(filas_analytics_paginas, key=lambda item: float(item.get("rebote", 0.0)), reverse=True)[:3]) or "Sin datos"

    # Construye distribuciones para gráficos de dashboard.
    severidades_rend = Counter([str(fila.get("severidad", "informativa")).lower() for fila in filas_rendimiento if fila.get("oportunidad")])

    # Construye distribución por tipo de mejora.
    tipos_mejora = Counter([str(fila.get("oportunidad", "Sin clasificar")).strip() for fila in filas_rendimiento if fila.get("oportunidad")])

    # Construye distribución de severidad técnica.
    severidades_tecnicas = Counter([str(fila.get("severidad", "informativa")).lower() for fila in filas if fila.get("problema")])

    # Construye distribución de incidencias por área para bloque ejecutivo.
    incidencias_por_area = Counter([str(fila.get("area", "Sin área")).strip() for fila in filas if fila.get("problema")])

    # Construye distribución de incidencias por tipo para bloque ejecutivo.
    incidencias_por_tipo = Counter([str(fila.get("tipo", "Sin tipo")).strip() for fila in filas if fila.get("problema")])

    # Obtiene incidencias agrupadas por familia desde métricas.
    incidencias_agrupadas_data = metricas.get("incidencias_agrupadas")
    incidencias_por_familia = incidencias_agrupadas_data if isinstance(incidencias_agrupadas_data, dict) else {}

    # Escribe bloque de severidad de rendimiento en hoja auxiliar.
    hoja_aux["A1"] = "Severidad rendimiento"
    hoja_aux["B1"] = "Cantidad"

    # Rellena tabla de severidad de rendimiento.
    for indice, severidad in enumerate(["crítica", "alta", "media", "baja", "informativa"], start=2):
        # Escribe etiqueta de severidad.
        hoja_aux[f"A{indice}"] = severidad

        # Escribe cantidad de oportunidades para la severidad.
        hoja_aux[f"B{indice}"] = severidades_rend.get(severidad, 0)

    # Escribe bloque de severidad técnica en hoja auxiliar.
    hoja_aux["D1"] = "Severidad técnica"
    hoja_aux["E1"] = "Cantidad"

    # Rellena tabla de severidad técnica.
    for indice, severidad in enumerate(["crítica", "alta", "media", "baja", "informativa"], start=2):
        # Escribe etiqueta de severidad técnica.
        hoja_aux[f"D{indice}"] = severidad

        # Escribe cantidad de incidencias técnicas.
        hoja_aux[f"E{indice}"] = severidades_tecnicas.get(severidad, 0)

    # Escribe bloque de tipos de mejora en hoja auxiliar.
    hoja_aux["G1"] = "Tipo de mejora"
    hoja_aux["H1"] = "Cantidad"

    # Escribe bloque de score por bloques en hoja auxiliar.
    hoja_aux["J1"] = "Bloque score"
    hoja_aux["K1"] = "Score"

    # Escribe bloque de comparación técnico vs agrupado en hoja auxiliar.
    hoja_aux["M1"] = "Tipo conteo"
    hoja_aux["N1"] = "Cantidad"
    hoja_aux["M2"] = "Incidencias técnicas"
    hoja_aux["N2"] = metricas["total_incidencias"]
    hoja_aux["M3"] = "Incidencias agrupadas"
    hoja_aux["N3"] = metricas.get("total_incidencias_agrupadas", 0)

    # Define pares de score por bloques para visualización.
    bloques_score = [
        ("Indexación/arquitectura", score_bloques.get("indexacion_arquitectura", {}).get("score", 0)),
        ("Contenido on-page", score_bloques.get("contenido_onpage", {}).get("score", 0)),
        ("Rendimiento", score_bloques.get("rendimiento", {}).get("score", 0)),
        ("Multimedia/accesibilidad", score_bloques.get("multimedia_accesibilidad", {}).get("score", 0)),
    ]

    # Recorre scores por bloque para poblar auxiliar.
    for indice, (bloque, score_bloque) in enumerate(bloques_score, start=2):
        # Escribe nombre del bloque de score.
        hoja_aux[f"J{indice}"] = bloque

        # Escribe score del bloque.
        hoja_aux[f"K{indice}"] = score_bloque

    # Recorre tipos de mejora más frecuentes.
    for indice, (tipo, cantidad) in enumerate(tipos_mejora.most_common(6), start=2):
        # Escribe etiqueta del tipo de mejora.
        hoja_aux[f"G{indice}"] = tipo

        # Escribe cantidad del tipo de mejora.
        hoja_aux[f"H{indice}"] = cantidad

    # Escribe fallback cuando no existan oportunidades.
    if not tipos_mejora:
        # Inserta etiqueta por defecto de mejoras.
        hoja_aux["G2"] = "Sin oportunidades"

        # Inserta valor por defecto.
        hoja_aux["H2"] = 1

    # Calcula total de quick wins para sección de oportunidades.
    total_quick_wins = len(_construir_quick_wins(filas, limite=7))

    # Renderiza bloque de oportunidades SEO priorizadas.
    oportunidades_pos_4_15 = len([fila for fila in filas_oportunidades_gsc if "4-15" in str(fila.get("oportunidad", ""))])
    oportunidades_ctr_bajo = len([fila for fila in filas_oportunidades_gsc if "ctr bajo" in str(fila.get("oportunidad", "")).lower()])
    mayor_potencial = ", ".join(str(fila.get("url", "")) for fila in filas_oportunidades_gsc[:3]) or "Sin datos"

    # Define desglose de score por bloques para lectura rápida.
    lineas_score_bloques = [
        f"Indexación / arquitectura: {score_bloques.get('indexacion_arquitectura', {}).get('score', 0)}",
        f"Contenido on-page: {score_bloques.get('contenido_onpage', {}).get('score', 0)}",
        f"Rendimiento: {score_bloques.get('rendimiento', {}).get('score', 0)}",
        f"Multimedia / accesibilidad: {score_bloques.get('multimedia_accesibilidad', {}).get('score', 0)}",
    ]

    # Renderiza bloque de visibilidad orgánica real.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "D9",
        "Visibilidad orgánica real",
        [
            f"Clics: {clics_totales_gsc}",
            f"Impresiones: {impresiones_totales_gsc}",
            f"CTR medio: {ctr_medio_gsc}",
            f"Posición media: {posicion_media_gsc}",
            f"Top páginas: {top_paginas}",
            f"Top queries: {top_queries}",
        ],
        COLOR_BLOQUE_VISIBILIDAD,
    )

    # Renderiza bloque de score por bloques.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "D17",
        "Score por bloques",
        lineas_score_bloques,
        COLOR_BLOQUE_SCORE,
    )

    # Renderiza bloque de oportunidades de negocio.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "D23",
        "Oportunidades",
        [
            f"Páginas con CTR bajo: {oportunidades_ctr_bajo}",
            f"Páginas en posición 4-15: {oportunidades_pos_4_15}",
            f"Quick wins detectados: {total_quick_wins}",
            f"Páginas de mayor potencial: {mayor_potencial}",
        ],
        COLOR_BLOQUE_OPORTUNIDADES,
    )

    # Renderiza bloque de comportamiento de usuario basado en Analytics.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "D28",
        "Comportamiento Analytics",
        [
            f"Sesiones totales: {sesiones_totales_analytics}",
            f"Usuarios totales: {usuarios_totales_analytics}",
            f"Rebote medio: {rebote_medio_analytics}",
            f"Duración media: {duracion_media_analytics}",
            f"Conversiones: {float(metricas_analytics.get('conversiones', 0.0))}",
            f"Top páginas tráfico: {top_paginas_analytics}",
            f"Páginas peor comportamiento: {peores_paginas_analytics}",
        ],
        "0E7490",
    )

    # Renderiza bloque de gestión de indexación.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "L9",
        "Gestión de indexación",
        [
            f"Indexable: {resumen_gestion_indexacion.get('indexable', 0)}",
            f"Revisar: {resumen_gestion_indexacion.get('revisar', 0)}",
            f"No indexar: {resumen_gestion_indexacion.get('no_indexar', 0)}",
        ],
        COLOR_BLOQUE_INDEXACION,
    )

    # Renderiza bloque de incidencias agrupadas y severidad.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "L14",
        "Incidencias",
        [
            f"Por severidad: crítica={incidencias_criticas}, alta={incidencias_altas}, media={incidencias_medias}, baja={incidencias_bajas}",
            "Por área: " + ", ".join(f"{k}={v}" for k, v in incidencias_por_area.most_common(3)),
            "Por tipo: " + ", ".join(f"{k}={v}" for k, v in incidencias_por_tipo.most_common(3)),
            "Por familia: " + ", ".join(f"{k}={v}" for k, v in list(incidencias_por_familia.items())[:3]),
        ],
        COLOR_BLOQUE_INCIDENCIAS,
    )

    # Renderiza bloque de ranking de páginas y queries prioritarias.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "L20",
        "Top páginas y queries",
        [
            f"Top páginas GSC: {top_paginas}",
            f"Top queries GSC: {top_queries}",
            f"Top páginas GA4: {top_paginas_analytics}",
        ],
        "334155",
    )

    # Renderiza bloque de páginas prioritarias por cruce SEO + negocio.
    _renderizar_bloque_dashboard(
        hoja_dashboard,
        "L30",
        "Páginas prioritarias",
        [
            f"{item.get('url', '')} (score={item.get('prioridad_score', 0)})"
            for item in construir_paginas_prioritarias(resultado, limite=4)
        ]
        or ["Sin señales suficientes para priorizar."],
        "7C3AED",
    )

    # Crea gráfico de distribución de severidad de rendimiento.
    grafico_severidad_rend = PieChart()

    # Define título del gráfico de severidad de rendimiento.
    grafico_severidad_rend.title = "Severidad oportunidades de rendimiento"

    # Carga datos del gráfico desde hoja auxiliar.
    grafico_severidad_rend.add_data(Reference(hoja_aux, min_col=2, min_row=1, max_row=6), titles_from_data=True)

    # Carga categorías del gráfico desde hoja auxiliar.
    grafico_severidad_rend.set_categories(Reference(hoja_aux, min_col=1, min_row=2, max_row=6))

    # Fija tamaño visual del gráfico.
    grafico_severidad_rend.width = 7.2

    # Fija altura visual del gráfico.
    grafico_severidad_rend.height = 5.4

    # Inserta gráfico en posición explícita.
    hoja_dashboard.add_chart(grafico_severidad_rend, "D30")

    # Crea gráfico de severidad técnica.
    grafico_severidad_tecnica = BarChart()

    # Define título del gráfico de severidad técnica.
    grafico_severidad_tecnica.title = "Incidencias técnicas por severidad"

    # Carga datos del gráfico desde hoja auxiliar.
    grafico_severidad_tecnica.add_data(Reference(hoja_aux, min_col=5, min_row=1, max_row=6), titles_from_data=True)

    # Carga categorías del gráfico desde hoja auxiliar.
    grafico_severidad_tecnica.set_categories(Reference(hoja_aux, min_col=4, min_row=2, max_row=6))

    # Ajusta ancho del gráfico.
    grafico_severidad_tecnica.width = 7.8

    # Ajusta alto del gráfico.
    grafico_severidad_tecnica.height = 5.4

    # Inserta gráfico en posición explícita.
    hoja_dashboard.add_chart(grafico_severidad_tecnica, "L24")

    # Crea gráfico de tipos de mejora de rendimiento.
    grafico_tipos_mejora = BarChart()

    # Define título del gráfico de tipos de mejora.
    grafico_tipos_mejora.title = "Top tipos de mejora PageSpeed"

    # Determina fila final de tipos de mejora.
    max_fila_tipos = max(2, 1 + min(6, len(tipos_mejora) if tipos_mejora else 1))

    # Carga datos del gráfico desde hoja auxiliar.
    grafico_tipos_mejora.add_data(Reference(hoja_aux, min_col=8, min_row=1, max_row=max_fila_tipos), titles_from_data=True)

    # Carga categorías del gráfico desde hoja auxiliar.
    grafico_tipos_mejora.set_categories(Reference(hoja_aux, min_col=7, min_row=2, max_row=max_fila_tipos))

    # Ajusta ancho del gráfico.
    grafico_tipos_mejora.width = 16.0

    # Ajusta alto del gráfico.
    grafico_tipos_mejora.height = 5.8

    # Inserta gráfico en posición explícita.
    hoja_dashboard.add_chart(grafico_tipos_mejora, "D46")

    # Crea gráfico de score por bloques para explicar score global.
    grafico_score_bloques = BarChart()

    # Define título del gráfico de score por bloques.
    grafico_score_bloques.title = "Score SEO por bloques"

    # Carga datos del gráfico de score por bloques.
    grafico_score_bloques.add_data(Reference(hoja_aux, min_col=11, min_row=1, max_row=5), titles_from_data=True)

    # Carga categorías del gráfico de score por bloques.
    grafico_score_bloques.set_categories(Reference(hoja_aux, min_col=10, min_row=2, max_row=5))

    # Ajusta tamaño visual del gráfico.
    grafico_score_bloques.width = 7.5

    # Ajusta alto visual del gráfico.
    grafico_score_bloques.height = 5.4

    # Inserta gráfico de score por bloques en dashboard.
    hoja_dashboard.add_chart(grafico_score_bloques, "L38")

    # Crea gráfico de comparación entre incidencias técnicas y agrupadas.
    grafico_agrupadas = PieChart()

    # Define título del gráfico de comparación.
    grafico_agrupadas.title = "Incidencias técnicas vs agrupadas"

    # Carga datos para gráfico de comparación.
    grafico_agrupadas.add_data(Reference(hoja_aux, min_col=14, min_row=1, max_row=3), titles_from_data=True)

    # Carga categorías para gráfico de comparación.
    grafico_agrupadas.set_categories(Reference(hoja_aux, min_col=13, min_row=2, max_row=3))

    # Ajusta ancho del gráfico comparativo.
    grafico_agrupadas.width = 7.2

    # Ajusta alto del gráfico comparativo.
    grafico_agrupadas.height = 5.4

    # Inserta gráfico comparativo en dashboard.
    hoja_dashboard.add_chart(grafico_agrupadas, "D60")

    # Oculta la hoja auxiliar para no contaminar entregable final.
    hoja_aux.sheet_state = "hidden"

    # Crea tabla Excel en hoja rendimiento cuando haya filas.
    if filas_rendimiento:
        # Define rango dinámico de tabla de rendimiento.
        rango = f"A1:T{len(filas_rendimiento) + 1}"

        # Crea tabla de rendimiento.
        tabla = ExcelTable(displayName="TablaRendimiento", ref=rango)

        # Define estilo de tabla.
        tabla.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)

        # Inserta tabla en hoja.
        hoja_rendimiento.add_table(tabla)

    # Ajusta anchos base del dashboard para tarjetas y bloques visuales.
    for indice_columna in range(1, 21):
        # Aplica ancho uniforme para permitir tarjetas amplias.
        hoja_dashboard.column_dimensions[get_column_letter(indice_columna)].width = 17

    # Congela paneles del dashboard para mantener visibilidad de KPIs.
    hoja_dashboard.freeze_panes = "A7"

    # Define columnas con texto extenso para wrap y legibilidad total.
    columnas_largas = {"url", "problema", "recomendacion", "observaciones", "query", "accion_recomendada"}

    # Aplica autoajuste global de todas las hojas de datos.
    hojas_a_ajustar = [
        hoja_kpis,
        hoja_dashboard,
        hoja_errores,
        hoja_contenido,
        hoja_rendimiento,
        hoja_indexacion,
        hoja_analytics,
        hoja_gsc_paginas,
        hoja_gsc_queries,
        hoja_oportunidades_gsc,
        hoja_aux,
    ]

    # Recorre hojas de datos para aplicar ajuste homogéneo.
    for hoja in hojas_a_ajustar:
        # Autoajusta legibilidad según contenido y columnas largas.
        _autoajustar_hoja(hoja, columnas_wrap=columnas_largas)

    # Guarda libro final en disco.
    libro.save(destino)

    # Devuelve ruta de salida.
    return destino


# Exporta un informe ejecutivo en Word.
def construir_modelo_semantico_informe(resultado: ResultadoAuditoria) -> dict[str, Any]:
    """Construye una única capa semántica neutral para DOCX/PDF/HTML."""

    # Calcula métricas globales reutilizables.
    metricas = calcular_metricas(resultado)

    # Resuelve periodo efectivo con helper común para evitar duplicación.
    periodo_desde, periodo_hasta = _resolver_periodo_analizado(resultado)

    # Construye narrativa base consolidada.
    bloques_narrativos = _construir_bloques_narrativos(resultado)

    # Construye datasets tabulares para secciones clave.
    filas_analytics = sorted(construir_filas_analytics_paginas(resultado), key=lambda item: float(item.get("sesiones", 0.0)), reverse=True)[:10]
    filas_indexacion = construir_filas_gestion_indexacion(resultado)[:12]
    filas_prioritarias = construir_paginas_prioritarias(resultado, limite=10)
    quick_wins = _construir_quick_wins(construir_filas_quick_wins(resultado), limite=10)

    # Agrupa datos de rendimiento por URL y estrategia para tablas comparables.
    rendimiento_por_url: dict[str, dict[str, object]] = {}

    # Recorre resultados de rendimiento para agregarlos por URL.
    for item in resultado.rendimiento:
        # Inicializa contenedor por URL cuando no exista.
        rendimiento_por_url.setdefault(item.url, {})

        # Guarda registro por estrategia de ejecución.
        rendimiento_por_url[item.url][item.estrategia] = item

    # Define colección de secciones semánticas.
    secciones: list[dict[str, Any]] = []

    # Crea sección semántica de portada.
    secciones.append(
        _crear_bloque_semantico(
            tipo="portada",
            titulo="Portada",
            parrafos=[
                "INFORME DE AUDITORÍA SEO",
                f"Cliente: {resultado.cliente}",
                f"Gestor: {resultado.gestor}",
                f"Sitemap: {resultado.sitemap}",
                f"Fecha: {resultado.fecha_ejecucion}",
                f"Periodo analizado: {periodo_desde} - {periodo_hasta}" if periodo_desde and periodo_hasta else "Periodo analizado: No disponible",
            ],
        )
    )

    # Recorre jerarquía visible para montar bloques homogéneos.
    for titulo_seccion in construir_jerarquia_visible(resultado):
        # Inicializa bloque semántico de la sección actual.
        bloque = _crear_bloque_semantico(tipo="seccion", titulo=titulo_seccion)

        # Inserta contenido tabular de KPIs.
        if titulo_seccion == "KPIs principales":
            bloque["tablas"].append(
                {
                    "titulo": "KPIs principales",
                    "columnas": ["KPI", "Valor"],
                    "filas": [
                        ["Total URLs", str(metricas["total_urls"])],
                        ["Total incidencias", str(metricas["total_incidencias"])],
                        ["Score SEO", str(metricas["score"])],
                        ["Fuentes activas", ", ".join(resultado.fuentes_activas)],
                    ],
                }
            )

        # Inserta tarjetas semánticas de quick wins.
        elif titulo_seccion == "Quick wins":
            for item in quick_wins:
                bloque["tarjetas"].append(
                    {
                        "titulo": f"URL: {item['url']}",
                        "subtitulo": f"Impacto: {item['impacto_maximo']} | Esfuerzo: {item['esfuerzo_maximo']}",
                        "listas": [
                            {"titulo": "Problemas", "items": [sanitizar_texto_editorial(str(valor)) for valor in item["problemas"]]},
                            {"titulo": "Acciones recomendadas", "items": [sanitizar_texto_editorial(str(valor)) for valor in item["recomendaciones"]]},
                        ],
                    }
                )

        # Inserta tabla de comportamiento y conversión.
        elif titulo_seccion == "Comportamiento y conversión":
            bloque["tablas"].append(
                {
                    "titulo": "Comportamiento y conversión",
                    "columnas": ["URL", "Sesiones", "Usuarios", "Rebote", "Duración", "Conversiones"],
                    "filas": [
                        [
                            str(fila.get("url", "")),
                            str(fila.get("sesiones", 0)),
                            str(fila.get("usuarios", 0)),
                            str(fila.get("rebote", 0)),
                            str(fila.get("duracion", 0)),
                            str(fila.get("conversiones", 0)),
                        ]
                        for fila in filas_analytics
                    ],
                }
            )

        # Inserta tabla de gestión de indexación.
        elif titulo_seccion == "Gestión de indexación":
            bloque["tablas"].append(
                {
                    "titulo": "Gestión de indexación",
                    "columnas": ["URL", "Clasificación", "Motivo", "Acción recomendada", "Prioridad"],
                    "filas": [
                        [
                            str(fila.get("url", "")),
                            str(fila.get("clasificacion", "")),
                            str(fila.get("motivo", "")),
                            str(fila.get("accion_recomendada", "")),
                            str(fila.get("prioridad", "")),
                        ]
                        for fila in filas_indexacion
                    ],
                }
            )

        # Inserta tabla de páginas prioritarias con explicación estructurada.
        elif titulo_seccion == "Páginas prioritarias":
            bloque["tablas"].append(
                {
                    "titulo": "Páginas prioritarias",
                    "columnas": ["URL", "Score", "Impresiones", "CTR", "Sesiones", "Conversiones", "Motivos"],
                    "filas": [
                        [
                            str(fila.get("url", "")),
                            str(fila.get("prioridad_score", 0)),
                            str(fila.get("impresiones", 0)),
                            str(fila.get("ctr", 0)),
                            str(fila.get("sesiones", 0)),
                            str(fila.get("conversiones", 0)),
                            "; ".join(fila.get("motivos", [])),
                        ]
                        for fila in filas_prioritarias
                    ],
                }
            )

        # Inserta tabla detallada de rendimiento y oportunidades.
        elif titulo_seccion == "Rendimiento y experiencia de usuario":
            # Inicializa filas de métricas para formato vertical.
            filas_rendimiento: list[list[str]] = []

            # Recorre muestra de URLs para mantener legibilidad.
            for url_actual, bloque_rend in list(rendimiento_por_url.items())[:4]:
                # Recorre ambas estrategias en orden estándar.
                for estrategia in ["mobile", "desktop"]:
                    # Obtiene registro por estrategia cuando exista.
                    item = bloque_rend.get(estrategia)

                    # Construye métricas verticales por bloque.
                    metricas_item = [
                        ("Performance", _valor_metrica(getattr(item, "performance_score", None))),
                        ("SEO", _valor_metrica(getattr(item, "seo_score", None))),
                        ("LCP", _valor_metrica(getattr(item, "lcp", None))),
                        ("CLS", _valor_metrica(getattr(item, "cls", None))),
                        ("INP", _valor_metrica(getattr(item, "inp", None))),
                        ("FCP", _valor_metrica(getattr(item, "fcp", None))),
                        ("Speed Index", _valor_metrica(getattr(item, "speed_index", None))),
                    ]

                    # Recorre cada métrica para añadir filas.
                    for indice_metrica, (nombre_metrica, valor_metrica) in enumerate(metricas_item):
                        # Añade fila con observación en primera línea del bloque.
                        filas_rendimiento.append(
                            [
                                f"{url_actual} [{estrategia}]" if indice_metrica == 0 else "",
                                nombre_metrica,
                                str(valor_metrica),
                                _interpretacion_rendimiento(getattr(item, "performance_score", None), estrategia) if indice_metrica == 0 else "",
                            ]
                        )

            # Inserta tabla de métricas de rendimiento.
            bloque["tablas"].append(
                {
                    "titulo": "Rendimiento por métrica",
                    "columnas": ["URL / Estrategia", "Métrica", "Valor", "Observación"],
                    "filas": filas_rendimiento,
                }
            )

            # Inicializa filas de oportunidades priorizadas.
            filas_oportunidades: list[list[str]] = []

            # Recorre resultados con oportunidades reales.
            for item in [fila for fila in resultado.rendimiento if fila.oportunidades][:6]:
                # Recorre oportunidades acotadas por bloque.
                for oportunidad in item.oportunidades[:4]:
                    # Añade oportunidad priorizada en formato tabular.
                    filas_oportunidades.append(
                        [
                            f"{item.url} [{item.estrategia}]",
                            str(oportunidad.titulo),
                            str(oportunidad.severidad),
                            _valor_metrica(oportunidad.ahorro_estimado),
                        ]
                    )

            # Inserta tabla de oportunidades si existen datos.
            if filas_oportunidades:
                bloque["tablas"].append(
                    {
                        "titulo": "Oportunidades PageSpeed priorizadas",
                        "columnas": ["URL / Estrategia", "Oportunidad", "Severidad", "Ahorro estimado"],
                        "filas": filas_oportunidades,
                    }
                )

        # Inserta contenido narrativo sanitizado como fallback estándar.
        else:
            bloque["parrafos"] = [sanitizar_texto_final_exportable(str(linea), formato="doc") for linea in bloques_narrativos.get(titulo_seccion, [])[:10]]

        # Añade nota estándar cuando no haya contenido explícito.
        if not bloque["parrafos"] and not bloque["tablas"] and not bloque["tarjetas"]:
            bloque["notas"].append("No disponible para esta ejecución.")

        # Inserta bloque en modelo final.
        secciones.append(bloque)

    # Construye anexo técnico en bloque separado.
    secciones.append(
        _crear_bloque_semantico(
            tipo="anexo",
            titulo="Anexo técnico",
            listas=[
                {
                    "titulo": "URLs auditadas (muestra)",
                    "items": [
                        sanitizar_texto_editorial(
                            f"{item.url} | HTTP={item.estado_http} | redirección={'Sí' if item.redirecciona else 'No'} | noindex={'Sí' if item.noindex else 'No'}"
                        )
                        for item in resultado.resultados[:20]
                    ],
                }
            ],
        )
    )

    # Construye etiqueta de periodo para reutilización consistente.
    periodo_texto = f"{periodo_desde} - {periodo_hasta}" if periodo_desde and periodo_hasta else "No disponible"

    # Devuelve modelo semántico único para todos los renderizadores.
    return {
        "meta": {
            "cliente": resultado.cliente,
            "gestor": resultado.gestor,
            "fecha_ejecucion": resultado.fecha_ejecucion,
            "fecha": resultado.fecha_ejecucion,
            "sitemap": resultado.sitemap,
            "periodo_desde": periodo_desde,
            "periodo_hasta": periodo_hasta,
            "periodo_texto": periodo_texto,
        },
        "secciones": secciones,
    }


# Inserta una tabla semántica en Word de forma homogénea.
def _renderizar_tabla_word(documento: Document, tabla_semantica: dict[str, Any]) -> None:
    """Renderiza una tabla semántica en DOCX manteniendo formato corporativo."""

    # Obtiene columnas definidas para la tabla.
    columnas = list(tabla_semantica.get("columnas", []))

    # Obtiene filas definidas para la tabla.
    filas = list(tabla_semantica.get("filas", []))

    # Crea tabla base con encabezados.
    tabla = documento.add_table(rows=1, cols=max(1, len(columnas)))

    # Aplica estilo de rejilla corporativa.
    tabla.style = "Table Grid"

    # Inserta encabezados de columna.
    for indice, columna in enumerate(columnas):
        tabla.rows[0].cells[indice].text = sanitizar_texto_final_exportable(str(columna), formato="doc")

    # Inserta filas de datos.
    for fila in filas:
        celdas = tabla.add_row().cells
        for indice, valor in enumerate(fila):
            if indice < len(celdas):
                celdas[indice].text = sanitizar_texto_final_exportable(str(valor), formato="doc")


# Obtiene subtables PDF según el tipo de tabla para priorizar legibilidad.
def _resolver_subtablas_pdf(tabla_semantica: dict[str, Any]) -> list[dict[str, Any]]:
    """Divide tablas muy anchas en bloques verticales semánticos."""

    # Obtiene título de la tabla semántica actual.
    titulo = str(tabla_semantica.get("titulo", "")).strip()

    # Divide páginas prioritarias en dos bloques para evitar horizontal extrema.
    if titulo == "Páginas prioritarias":
        columnas = list(tabla_semantica.get("columnas", []))
        filas = [list(fila) for fila in tabla_semantica.get("filas", [])]
        indice = {str(col): pos for pos, col in enumerate(columnas)}

        # Define agrupaciones de columnas por bloque.
        grupo_1 = ["URL", "Score", "Impresiones", "CTR"]
        grupo_2 = ["URL", "Sesiones", "Conversiones", "Motivos"]

        # Construye subtablas manteniendo orden semántico.
        subtables = []
        for sufijo, grupo in [(" (impacto)", grupo_1), (" (acciones)", grupo_2)]:
            columnas_grupo = [col for col in grupo if col in indice]
            filas_grupo = [[fila[indice[col]] if indice[col] < len(fila) else "" for col in columnas_grupo] for fila in filas]
            subtables.append({"titulo": f"{titulo}{sufijo}", "columnas": columnas_grupo, "filas": filas_grupo})
        return subtables

    # Divide rendimiento por métrica para priorizar observación con wrap.
    if titulo == "Rendimiento por métrica":
        columnas = list(tabla_semantica.get("columnas", []))
        filas = [list(fila) for fila in tabla_semantica.get("filas", [])]
        indice = {str(col): pos for pos, col in enumerate(columnas)}
        grupo_1 = ["URL / Estrategia", "Métrica", "Valor"]
        grupo_2 = ["URL / Estrategia", "Observación"]
        subtables = []
        for sufijo, grupo in [(" (métricas)", grupo_1), (" (insights)", grupo_2)]:
            columnas_grupo = [col for col in grupo if col in indice]
            filas_grupo = [[fila[indice[col]] if indice[col] < len(fila) else "" for col in columnas_grupo] for fila in filas]
            subtables.append({"titulo": f"{titulo}{sufijo}", "columnas": columnas_grupo, "filas": filas_grupo})
        return subtables

    # Mantiene estructura original para el resto de tablas.
    return [tabla_semantica]

# Inserta una tabla semántica en PDF de forma homogénea.
def _renderizar_tabla_pdf(tabla_semantica: dict[str, Any], estilos: dict[str, Any]) -> list[Any]:
    """Renderiza tabla semántica para ReportLab usando estilo corporativo."""

    # Define ancho útil A4 descontando márgenes horizontales.
    ancho_util = float(A4[0] - PDF_HORIZONTAL_MARGIN_POINTS)

    # Define estilo de celdas para wrapping real en contenido.
    estilo_celda = ParagraphStyle(name="TablaCelda", parent=estilos["BodyText"], fontSize=8.5, leading=11)

    # Define estilo de cabecera centrado.
    estilo_cabecera = ParagraphStyle(name="TablaCabecera", parent=estilos["BodyText"], fontSize=8.5, leading=10, textColor=colors.white)

    # Inicializa lista de elementos tabulares.
    elementos_tabla: list[Any] = []

    # Recorre subtables resueltas por semántica.
    for bloque in _resolver_subtablas_pdf(tabla_semantica):
        # Construye matriz con encabezado y filas.
        columnas = list(bloque.get("columnas", []))
        filas = [list(fila) for fila in bloque.get("filas", [])]
        datos = [columnas] + filas

        # Garantiza matriz mínima para evitar errores de ReportLab.
        if not datos:
            datos = [["Dato"], ["No disponible"]]
            columnas = ["Dato"]
            filas = [["No disponible"]]

        # Crea anchos explícitos para evitar auto-layout extremo.
        col_widths = _calcular_col_widths_pdf(columnas, filas, ancho_util=ancho_util)

        # Construye encabezado con Paragraph para wrapping real.
        cabecera = [Paragraph(f"<b>{sanear_texto_para_pdf(str(valor))}</b>", estilo_cabecera) for valor in datos[0]]

        # Construye filas de contenido con wrapping real por celda.
        cuerpo = [[Paragraph(sanear_texto_para_pdf(str(valor)), estilo_celda) for valor in fila] for fila in datos[1:]]
        datos_render = [cabecera] + cuerpo

        # Crea tabla con repetición de encabezado.
        tabla = Table(datos_render, repeatRows=1, colWidths=col_widths)

        # Aplica estilo corporativo unificado.
        tabla.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9D9D9")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        # Acumula tabla y espaciado para evitar bloque plano.
        elementos_tabla.append(tabla)
        elementos_tabla.append(Spacer(1, 10))

    # Devuelve elementos tabulares maquetados.
    return elementos_tabla

def exportar_word(resultado: ResultadoAuditoria, path_salida: Path) -> Path:
    """Genera un DOCX corporativo usando estructura semántica intermedia."""

    # Garantiza carpeta de salida.
    asegurar_directorio(path_salida)

    # Define ruta del documento Word.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}.docx"

    # Crea documento vacío.
    documento = Document()

    # Construye modelo semántico único para reutilizar metadatos editoriales.
    modelo = construir_modelo_semantico_informe(resultado)

    # Configura estilo base de documento.
    documento.styles["Normal"].font.name = "Calibri"

    # Configura tamaño base de fuente.
    documento.styles["Normal"].font.size = Pt(11)

    # Inserta portada.
    titulo = documento.add_paragraph("INFORME DE AUDITORÍA SEO")

    # Centra título de portada.
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Aplica estilo visual a portada.
    titulo.runs[0].font.color.rgb = RGBColor(31, 78, 120)

    # Aplica tamaño destacado a portada.
    titulo.runs[0].font.size = Pt(26)

    # Extrae metadatos semánticos para portada consistente.
    meta_modelo = modelo.get("meta", {})

    # Renderiza metadatos en tabla fija para conservar jerarquía visual.
    tabla_meta = documento.add_table(rows=6, cols=2)
    tabla_meta.style = "Table Grid"
    filas_meta = [
        ("Cliente", str(meta_modelo.get("cliente", resultado.cliente))),
        ("Gestor", str(meta_modelo.get("gestor", resultado.gestor))),
        ("Fecha de ejecución", str(meta_modelo.get("fecha_ejecucion", resultado.fecha_ejecucion))),
        ("Periodo desde", str(meta_modelo.get("periodo_desde") or "No disponible")),
        ("Periodo hasta", str(meta_modelo.get("periodo_hasta") or "No disponible")),
        ("Sitemap", str(meta_modelo.get("sitemap", resultado.sitemap))),
    ]
    for indice, (etiqueta, valor) in enumerate(filas_meta):
        tabla_meta.cell(indice, 0).text = sanitizar_texto_final_exportable(etiqueta, formato="doc")
        tabla_meta.cell(indice, 1).text = sanitizar_texto_final_exportable(valor, formato="doc")

    # Añade línea editorial con el periodo consolidado para visibilidad ejecutiva.
    documento.add_paragraph(f"Periodo analizado: {sanitizar_texto_final_exportable(str(meta_modelo.get('periodo_texto', 'No disponible')), formato='doc')}")

    # Inserta salto de página tras portada.
    documento.add_page_break()

    # Recorre secciones semánticas excluyendo portada ya pintada.
    for seccion in modelo["secciones"]:
        if seccion["tipo_bloque"] == "portada":
            continue
        if seccion["titulo"] == "Anexo técnico":
            documento.add_page_break()
        documento.add_heading(sanitizar_texto_final_exportable(str(seccion["titulo"]), formato="doc"), level=1)
        for tabla in seccion.get("tablas", []):
            _renderizar_tabla_word(documento, tabla)
        for tarjeta in seccion.get("tarjetas", []):
            documento.add_paragraph(sanitizar_texto_final_exportable(str(tarjeta.get("titulo", "")), formato="doc"))
            documento.add_paragraph(sanitizar_texto_final_exportable(str(tarjeta.get("subtitulo", "")), formato="doc"))
            for lista in tarjeta.get("listas", []):
                documento.add_paragraph(sanitizar_texto_final_exportable(str(lista.get("titulo", "")), formato="doc"))
                for item in lista.get("items", []):
                    documento.add_paragraph(sanitizar_texto_final_exportable(str(item), formato="doc"), style="List Bullet")
        for linea in seccion.get("parrafos", []):
            documento.add_paragraph(sanitizar_texto_final_exportable(str(linea), formato="doc"))
        for nota in seccion.get("notas", []):
            documento.add_paragraph(sanitizar_texto_final_exportable(str(nota), formato="doc"))

    # Guarda documento final.
    documento.save(destino)

    # Devuelve ruta del Word generado.
    return destino


# Exporta un PDF con estructura similar a Word.
def exportar_pdf(resultado: ResultadoAuditoria, path_salida: Path) -> Path:
    """Genera un PDF profesional respetando la misma jerarquía documental."""

    # Garantiza carpeta de salida.
    asegurar_directorio(path_salida)

    # Define ruta del PDF final.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}.pdf"

    # Crea documento PDF base.
    pdf = SimpleDocTemplate(str(destino), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)

    # Inicializa lista de elementos del PDF.
    elementos = []

    # Construye modelo semántico único para render documental.
    modelo = construir_modelo_semantico_informe(resultado)

    # Obtiene estilos de referencia.
    estilos = getSampleStyleSheet()

    # Crea estilo de portada.
    estilo_portada = ParagraphStyle(name="Portada", parent=estilos["Title"], alignment=1, textColor=colors.HexColor("#1F4E78"))

    # Inserta portada.
    elementos.append(Spacer(1, 80))

    # Inserta título de portada.
    elementos.append(Paragraph("INFORME DE AUDITORÍA SEO", estilo_portada))

    # Extrae metadatos semánticos para portada consistente.
    meta_modelo = modelo.get("meta", {})

    # Inserta datos generales de portada.
    elementos.append(Spacer(1, 16))
    datos_meta_portada = [
        ["Cliente", str(meta_modelo.get("cliente", resultado.cliente))],
        ["Gestor", str(meta_modelo.get("gestor", resultado.gestor))],
        ["Fecha de ejecución", str(meta_modelo.get("fecha_ejecucion", resultado.fecha_ejecucion))],
        ["Periodo desde", str(meta_modelo.get("periodo_desde") or "No disponible")],
        ["Periodo hasta", str(meta_modelo.get("periodo_hasta") or "No disponible")],
        ["Sitemap", str(meta_modelo.get("sitemap", resultado.sitemap))],
    ]
    tabla_meta_portada = Table(
        [
            [
                Paragraph(f"<b>{sanear_texto_para_pdf(etiqueta)}</b>", estilos["Normal"]),
                Paragraph(sanear_texto_para_pdf(valor), estilos["Normal"]),
            ]
            for etiqueta, valor in datos_meta_portada
        ],
        colWidths=[130, float(A4[0] - PDF_HORIZONTAL_MARGIN_POINTS - 130)],
    )
    tabla_meta_portada.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9D9D9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    elementos.append(tabla_meta_portada)
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(f"Periodo analizado: {sanear_texto_para_pdf(str(meta_modelo.get('periodo_texto', 'No disponible')))}", estilos["Normal"]))
    elementos.append(PageBreak())

    # Recorre secciones semánticas excluyendo portada ya renderizada.
    for seccion in modelo["secciones"]:
        if seccion["tipo_bloque"] == "portada":
            continue
        if seccion["titulo"] == "Anexo técnico":
            elementos.append(PageBreak())
        elementos.append(Paragraph(f"<b>{sanear_texto_para_pdf(str(seccion['titulo']))}</b>", estilos["Heading2"]))
        for tabla in seccion.get("tablas", []):
            elementos.extend(_renderizar_tabla_pdf(tabla, estilos))
        for tarjeta in seccion.get("tarjetas", []):
            elementos.append(Paragraph(sanear_texto_para_pdf(str(tarjeta.get("titulo", ""))), estilos["Normal"]))
            elementos.append(Paragraph(sanear_texto_para_pdf(str(tarjeta.get("subtitulo", ""))), estilos["Normal"]))
            for lista in tarjeta.get("listas", []):
                elementos.append(Paragraph(sanear_texto_para_pdf(str(lista.get("titulo", ""))), estilos["Normal"]))
                for item in lista.get("items", []):
                    elementos.append(Paragraph(sanear_texto_para_pdf(f"- {item}"), estilos["Normal"]))
        for linea in seccion.get("parrafos", []):
            elementos.append(Paragraph(sanear_texto_para_pdf(str(linea)), estilos["Normal"]))
        for nota in seccion.get("notas", []):
            elementos.append(Paragraph(sanear_texto_para_pdf(str(nota)), estilos["Normal"]))
        elementos.append(Spacer(1, 6))

    # Construye PDF final.
    pdf.build(elementos)

    # Devuelve ruta del PDF generado.
    return destino


# Exporta el informe IA en Markdown para edición y revisión humana.
def exportar_markdown_ia(resultado: ResultadoAuditoria, path_salida: Path) -> Path | None:
    """Genera un archivo Markdown con el resumen IA y metadatos."""

    # Sale sin crear archivo cuando no existe resumen IA.
    if not resultado.resumen_ia:
        # Retorna None cuando no hay contenido IA.
        return None

    # Garantiza la existencia de la carpeta de salida.
    asegurar_directorio(path_salida)

    # Define la ruta del archivo Markdown.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}_ia.md"

    # Resuelve periodo efectivo con helper común para evitar duplicación.
    periodo_desde, periodo_hasta = _resolver_periodo_analizado(resultado)

    # Construye etiqueta unificada de periodo.
    periodo_texto = f"{periodo_desde} - {periodo_hasta}" if periodo_desde and periodo_hasta else "No disponible"

    # Construye encabezado contextual de informe auxiliar IA.
    contenido = (
        "# Informe SEO IA (auxiliar interno)\n\n"
        f"- Cliente: {resultado.cliente}\n"
        f"- Gestor: {resultado.gestor}\n"
        f"- Fecha de ejecución: {resultado.fecha_ejecucion}\n"
        f"- Periodo analizado: {periodo_texto}\n"
        f"- Sitemap: {resultado.sitemap}\n\n"
        f"{resultado.resumen_ia}\n"
    )

    # Escribe el contenido del informe IA en UTF-8.
    destino.write_text(contenido, encoding="utf-8")

    # Devuelve la ruta del archivo generado.
    return destino


# Exporta un informe HTML reutilizable para visualización web.
def exportar_html(resultado: ResultadoAuditoria, path_salida: Path) -> Path:
    """Genera una versión HTML limpia con KPIs e incidencias principales."""

    # Garantiza que la carpeta de salida exista.
    asegurar_directorio(path_salida)

    # Define la ruta del archivo HTML.
    destino = path_salida / f"{construir_prefijo_archivo(resultado)}.html"

    # Calcula métricas globales para cabecera.
    metricas = calcular_metricas(resultado)

    # Construye modelo semántico único para render web.
    modelo = construir_modelo_semantico_informe(resultado)

    # Extrae metadatos editoriales para cabecera premium.
    meta_modelo = modelo.get("meta", {})

    # Obtiene filas de incidencias para el detalle técnico final.
    filas_ordenadas = sorted(construir_filas(resultado), key=lambda fila: (_peso_severidad(str(fila.get("severidad", "informativa"))), str(fila.get("url", ""))))

    # Inicializa render de secciones semánticas.
    secciones_html: list[str] = []

    # Recorre secciones y las convierte a HTML uniforme.
    for seccion in modelo["secciones"]:
        if seccion["tipo_bloque"] == "portada":
            continue
        bloques_html: list[str] = []
        for tabla in seccion.get("tablas", []):
            columnas = "".join(f"<th>{escape(sanitizar_texto_final_exportable(str(columna), formato='html'))}</th>" for columna in tabla.get("columnas", []))
            filas = "".join(f"<tr>{''.join(f'<td>{escape(sanitizar_texto_final_exportable(str(valor), formato='html'))}</td>' for valor in fila)}</tr>" for fila in tabla.get("filas", []))
            colspan = max(1, len(tabla.get("columnas", [])))
            bloques_html.append(
                f"<h3>{escape(sanitizar_texto_final_exportable(str(tabla.get('titulo', 'Tabla')), formato='html'))}</h3>"
                f"<table><thead><tr>{columnas}</tr></thead>"
                f"<tbody>{filas or f'<tr><td colspan=\"{colspan}\">No disponible</td></tr>'}</tbody></table>"
            )
        for tarjeta in seccion.get("tarjetas", []):
            listas = "".join(
                f"<p><b>{escape(sanitizar_texto_final_exportable(str(lista.get('titulo', '')), formato='html'))}</b></p><ul>{''.join(f'<li>{escape(sanitizar_texto_final_exportable(str(item), formato='html'))}</li>' for item in lista.get('items', []))}</ul>"
                for lista in tarjeta.get("listas", [])
            )
            bloques_html.append(f"<div class='tarjeta'><h3>{escape(sanitizar_texto_final_exportable(str(tarjeta.get('titulo', '')), formato='html'))}</h3><p>{escape(sanitizar_texto_final_exportable(str(tarjeta.get('subtitulo', '')), formato='html'))}</p>{listas}</div>")
        if seccion.get("parrafos"):
            bloques_html.append("<ul>" + "".join(f"<li>{escape(sanitizar_texto_final_exportable(str(linea), formato='html'))}</li>" for linea in seccion.get("parrafos", [])) + "</ul>")
        if seccion.get("notas"):
            bloques_html.append("".join(f"<p>{escape(sanitizar_texto_final_exportable(str(nota), formato='html'))}</p>" for nota in seccion.get("notas", [])))
        secciones_html.append(f"<div class='bloque'><h2>{escape(sanitizar_texto_final_exportable(str(seccion['titulo']), formato='html'))}</h2>{''.join(bloques_html)}</div>")

    # Construye contenido HTML básico y portable.
    contenido = f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>Informe SEO - {sanitizar_texto_final_exportable(resultado.cliente, formato='html')}</title>
  <style>
    body {{ font-family: Inter, Segoe UI, Arial, sans-serif; margin: 0; color: #1f2937; background: #f4f7fb; }}
    .contenedor {{ max-width: 1280px; margin: 0 auto; padding: 24px; }}
    .cabecera {{ background: linear-gradient(135deg, #1F4E78 0%, #0f2f4f 100%); color: #fff; border-radius: 14px; padding: 24px; box-shadow: 0 8px 20px rgba(15, 47, 79, 0.2); }}
    .cabecera h1 {{ margin: 0 0 10px 0; color: #fff; }}
    .metadatos {{ display: grid; grid-template-columns: repeat(2, minmax(220px, 1fr)); gap: 8px 16px; font-size: 14px; margin-top: 8px; }}
    .meta-ejecutiva {{ margin-top: 12px; background: rgba(255, 255, 255, 0.12); border: 1px solid rgba(255,255,255,0.22); border-radius: 12px; padding: 12px; }}
    .periodo-destacado {{ margin-top: 10px; display: inline-block; background: #FBBF24; color: #1F2937; padding: 6px 12px; border-radius: 999px; font-weight: 700; }}
    h2 {{ color: #1F4E78; margin-top: 24px; }}
    .kpis {{ display: grid; grid-template-columns: repeat(5, minmax(180px, 1fr)); gap: 12px; }}
    .kpi {{ background: #ffffff; padding: 12px; border-radius: 10px; border: 1px solid #dbe5f0; box-shadow: 0 3px 10px rgba(17, 24, 39, 0.06); }}
    .bloque {{ margin-top: 18px; padding: 16px; border: 1px solid #dce6f2; border-radius: 12px; background: #fff; box-shadow: 0 2px 8px rgba(17, 24, 39, 0.04); }}
    .tarjetas {{ display: grid; grid-template-columns: repeat(2, minmax(260px, 1fr)); gap: 12px; }}
    .tarjeta {{ background: #f8fbff; border: 1px solid #d7e4f4; border-radius: 10px; padding: 12px; }}
    .tarjeta h3 {{ margin: 0 0 8px 0; color: #1F4E78; font-size: 14px; }}
    .tarjeta ul {{ margin: 4px 0 8px 18px; padding: 0; }}
    table {{ border-collapse: collapse; width: 100%; margin-top: 12px; table-layout: fixed; }}
    th, td {{ border: 1px solid #d1d5db; padding: 8px; font-size: 12px; text-align: left; vertical-align: top; word-break: break-word; }}
    th {{ background: #1F4E78; color: white; }}
  </style>
</head>
<body>
  <div class="contenedor">
  <section class="cabecera">
    <h1>Informe de Auditoría SEO</h1>
    <div class="meta-ejecutiva">
      <div class="metadatos">
        <div><b>Cliente:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("cliente", resultado.cliente)), formato="html"))}</div>
        <div><b>Gestor:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("gestor", resultado.gestor)), formato="html"))}</div>
        <div><b>Fecha de ejecución:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("fecha_ejecucion", resultado.fecha_ejecucion)), formato="html"))}</div>
        <div><b>Periodo desde:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("periodo_desde") or "No disponible"), formato="html"))}</div>
        <div><b>Periodo hasta:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("periodo_hasta") or "No disponible"), formato="html"))}</div>
        <div><b>Sitemap:</b> {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("sitemap", resultado.sitemap)), formato="html"))}</div>
      </div>
      <div class="periodo-destacado">Periodo analizado: {escape(sanitizar_texto_final_exportable(str(meta_modelo.get("periodo_texto", "No disponible")), formato="html"))}</div>
    </div>
  </section>
  <h2>KPIs ejecutivos</h2>
  <div class="kpis">
    <div class="kpi"><b>Total URLs</b><br>{metricas["total_urls"]}</div>
    <div class="kpi"><b>Total incidencias</b><br>{metricas["total_incidencias"]}</div>
    <div class="kpi"><b>Incidencias agrupadas</b><br>{metricas.get("total_incidencias_agrupadas", 0)}</div>
    <div class="kpi"><b>URLs sanas</b><br>{metricas["urls_sanas"]}</div>
    <div class="kpi"><b>Score SEO</b><br>{metricas["score"]}</div>
  </div>
  {''.join(secciones_html)}
  <h2>Incidencias técnicas (detalle)</h2>
  <table>
    <thead><tr><th>URL</th><th>Severidad</th><th>Área</th><th>Problema</th><th>Recomendación</th></tr></thead>
    <tbody>
      {''.join(
        f"<tr style='background:{_color_pastel_severidad(str(fila.get('severidad','informativa')))}'>"
        f"<td>{escape(sanitizar_texto_final_exportable(str(fila.get('url','')), formato="html"))}</td>"
        f"<td>{escape(sanitizar_texto_final_exportable(str(fila.get('severidad','')), formato="html"))}</td>"
        f"<td>{escape(sanitizar_texto_final_exportable(str(fila.get('area','')), formato="html"))}</td>"
        f"<td>{escape(sanitizar_texto_final_exportable(str(fila.get('problema','')), formato="html"))}</td>"
        f"<td>{escape(sanitizar_texto_final_exportable(str(fila.get('recomendacion','')), formato="html"))}</td>"
        f"</tr>"
        for fila in filas_ordenadas[:120]
      )}
    </tbody>
  </table>
  </div>
</body>
</html>"""

    # Escribe contenido HTML en UTF-8.
    destino.write_text(contenido, encoding="utf-8")

    # Devuelve la ruta del HTML generado.
    return destino
