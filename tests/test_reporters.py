# Importa Path para construir rutas temporales de prueba.
from pathlib import Path

# Importa modelos del dominio para fabricar una auditoría mínima de prueba.
from seo_auditor.models import DatosAnalytics, DatosSearchConsole, DecisionIndexacion, HallazgoSeo, MetricaAnalyticsPagina, MetricaGscPagina, OportunidadRendimiento, ResultadoAuditoria, ResultadoRendimiento, ResultadoUrl
import seo_auditor.reporters as reporters_mod

# Importa funciones de reporters bajo prueba.
from seo_auditor.reporters import (
    _calcular_col_widths_pdf,
    _renderizar_bloque_dashboard,
    _renderizar_tabla_pdf,
    _renderizar_tabla_word,
    _resolver_subtablas_pdf,
    _construir_bloques_narrativos,
    _construir_quick_wins,
    PDF_HORIZONTAL_MARGIN_POINTS,
    calcular_score_prioridad_pagina,
    calcular_metricas,
    construir_modelo_semantico_informe,
    construir_cruces_gsc_analytics,
    construir_filas_contenido_consolidado,
    construir_jerarquia_visible,
    construir_secciones_desde_ia,
    exportar_excel,
    exportar_html,
    exportar_pdf,
    exportar_word,
    reemplazar_emojis_problematicos,
    sanitizar_texto_final_exportable,
    sanear_texto_para_pdf,
)

# Importa lector de libros Excel para validar KPIs.
from openpyxl import load_workbook
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table
import re


# Verifica que el saneamiento escape etiquetas potencialmente problemáticas.
def test_sanear_texto_para_pdf_escapa_marcado_html() -> None:
    """Comprueba que el texto con etiquetas HTML se escape para reportlab."""

    # Define un texto con etiquetas que pueden romper el parser de reportlab.
    texto = "Etiqueta <link rel='canonical' href='https://ejemplo.com'> y <b>negrita</b>."

    # Ejecuta el saneamiento antes de generar el PDF.
    saneado = sanear_texto_para_pdf(texto)

    # Verifica que la etiqueta se haya escapado correctamente.
    assert "&lt;link rel='canonical' href='https://ejemplo.com'&gt;" in saneado

    # Verifica que también se haya escapado el cierre de negrita.
    assert "&lt;b&gt;negrita&lt;/b&gt;" in saneado


# Verifica que jerarquía oculte Analytics aunque GSC esté activo.
def test_construir_jerarquia_visible_oculta_analytics_si_esta_inactivo() -> None:
    """Comprueba que la sección Analytics no aparezca cuando GA4 no está activo."""

    # Construye auditoría con GSC activo y Analytics inactivo.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
        search_console=DatosSearchConsole(activo=True),
        analytics=DatosAnalytics(activo=False),
    )

    # Calcula jerarquía visible.
    jerarquia = construir_jerarquia_visible(auditoria)

    # Verifica que la sección de Analytics no esté visible.
    assert "Comportamiento y conversión" not in jerarquia


# Verifica cruce GSC+Analytics normalizando URL completa vs pagePath.
def test_construir_cruces_gsc_analytics_normaliza_url_vs_path() -> None:
    """Comprueba que el cruce detecte la misma página con formatos distintos."""

    # Construye auditoría con ambas fuentes activas.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
        search_console=DatosSearchConsole(
            activo=True,
            paginas=[
                MetricaGscPagina(
                    url="https://ejemplo.com/producto/",
                    clicks=10,
                    impresiones=800,
                    ctr=0.0125,
                    posicion_media=8.1,
                )
            ],
        ),
        analytics=DatosAnalytics(
            activo=True,
            paginas=[
                MetricaAnalyticsPagina(
                    url="/producto",
                    sesiones=120,
                    usuarios=95,
                    rebote=0.72,
                    duracion_media=38,
                    conversiones=0,
                    calidad_trafico="baja",
                )
            ],
        ),
    )

    # Ejecuta cruce entre fuentes por URL.
    cruces = construir_cruces_gsc_analytics(auditoria)

    # Verifica que se haya cruzado la fila correctamente.
    assert len(cruces) == 1
    assert cruces[0]["url"] == "https://ejemplo.com/producto/"
    assert cruces[0]["sesiones"] == 120.0


# Verifica que resumen ejecutivo incluya explícitamente el periodo analizado.
def test_construir_bloques_narrativos_incluye_periodo_analizado() -> None:
    """Comprueba inserción de línea de periodo en el resumen ejecutivo."""

    # Construye auditoría mínima con periodo global informado.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
        periodo_date_from="2026-02-25",
        periodo_date_to="2026-03-24",
    )

    # Construye bloques narrativos finales.
    bloques = _construir_bloques_narrativos(auditoria)

    # Verifica línea explícita solicitada en resumen ejecutivo.
    assert "Periodo analizado: 2026-02-25 - 2026-03-24" in bloques["Resumen ejecutivo"]


# Verifica que la transformación de markdown IA cree secciones limpias.
def test_construir_secciones_desde_ia_limpia_markdown() -> None:
    """Comprueba conversión de markdown crudo a estructura de secciones."""

    # Define texto IA con markdown mixto.
    texto = "## Resumen ejecutivo\n**Texto clave**\n- Punto uno\n- Punto dos\n---\n### Roadmap\n1. Acción A"

    # Convierte texto en estructura intermedia.
    secciones = construir_secciones_desde_ia(texto)

    # Verifica que existan secciones con contenido.
    assert len(secciones) >= 2

    # Verifica que no persista markdown crudo en items.
    assert "**" not in " ".join(str(item) for seccion in secciones for item in seccion["items"])


# Verifica que el score SEO no sea excesivamente punitivo por defecto.
def test_calcular_metricas_score_ponderado() -> None:
    """Comprueba que la fórmula de score devuelva un valor razonable."""

    # Crea un hallazgo de severidad alta para la muestra.
    hallazgo = HallazgoSeo(
        tipo="indexación",
        severidad="alta",
        descripcion="La URL redirecciona.",
        recomendacion="Actualizar URL final.",
        area="Arquitectura",
        impacto="Alto",
        esfuerzo="Bajo",
        prioridad="P1",
    )

    # Construye un resultado URL con una incidencia.
    resultado_url = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=301,
        redirecciona=True,
        url_final="https://ejemplo.com/b",
        title="Página",
        h1="H1",
        meta_description="Meta",
        canonical="https://ejemplo.com/b",
        noindex=False,
        hallazgos=[hallazgo],
    )

    # Construye una auditoría mínima de ejemplo.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[resultado_url],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-19",
        gestor="Juan Antonio Sánchez Plaza",
    )

    # Calcula métricas agregadas.
    metricas = calcular_metricas(auditoria)

    # Verifica que el score esté en el rango operativo esperado.
    assert 5.0 <= float(metricas["score"]) <= 100.0

    # Verifica que exista el desglose de score por bloques.
    assert "score_bloques" in metricas

    # Verifica que exista total de incidencias agrupadas.
    assert "total_incidencias_agrupadas" in metricas


# Verifica que se rellenen secciones obligatorias cuando no haya IA.
def test_construir_bloques_narrativos_generar_fallback_completo() -> None:
    """Comprueba que todas las secciones narrativas tengan contenido de fallback."""

    # Crea un hallazgo representativo para alimentar secciones.
    hallazgo = HallazgoSeo(
        tipo="contenido",
        severidad="alta",
        descripcion="Falta title.",
        recomendacion="Definir title único.",
        area="Contenido",
        impacto="Alto",
        esfuerzo="Bajo",
        prioridad="P1",
    )

    # Construye un resultado URL mínimo.
    resultado_url = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="",
        h1="H1",
        meta_description="Meta",
        canonical="https://ejemplo.com/a",
        noindex=False,
        hallazgos=[hallazgo],
    )

    # Construye una auditoría sin resumen IA para forzar fallback.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[resultado_url],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
    )

    # Genera bloques narrativos por fallback.
    bloques = _construir_bloques_narrativos(auditoria)

    # Verifica que cada sección visible obligatoria tenga al menos una línea.
    assert all(len(bloques[seccion]) > 0 for seccion in construir_jerarquia_visible(auditoria))

    # Verifica que roadmap incluya fase de medio plazo.
    assert any("60" in linea or "medio plazo" in linea.lower() for linea in bloques["Roadmap"])

    # Verifica que exista bloque dedicado a gestión de indexación.
    assert len(bloques["Gestión de indexación"]) > 0


# Verifica que quick wins elimine duplicados y filas incompletas.
def test_construir_quick_wins_deduplica_y_filtra() -> None:
    """Comprueba que la lista de quick wins sea consistente y sin repeticiones torpes."""

    # Define filas con duplicado y fila incompleta.
    filas = [
        {"url": "https://ejemplo.com/a", "problema": "Falta title", "recomendacion": "Añadir title", "impacto": "Alto", "esfuerzo": "Bajo"},
        {"url": "https://ejemplo.com/a", "problema": "Falta title", "recomendacion": "Añadir title", "impacto": "Alto", "esfuerzo": "Bajo"},
        {"url": "", "problema": "Sin meta", "recomendacion": "Añadir meta", "impacto": "Alto", "esfuerzo": "Bajo"},
    ]

    # Construye quick wins deduplicados.
    quick_wins = _construir_quick_wins(filas, limite=10)

    # Verifica que solo quede una URL agrupada válida.
    assert len(quick_wins) == 1

    # Verifica que se agrupen problemas en lista.
    assert isinstance(quick_wins[0]["problemas"], list)

    # Verifica que se agrupen recomendaciones en lista.
    assert isinstance(quick_wins[0]["recomendaciones"], list)


# Verifica sustitución de emojis por etiquetas seguras para documentos.
def test_reemplazar_emojis_problematicos_sustituye_glifos() -> None:
    """Comprueba que los emojis se sustituyan por etiquetas textuales legibles."""

    # Define línea con emojis de estado frecuentes.
    texto = "✅ OK de indexación ⚠️ revisar canonical ❌ fallo crítico"

    # Ejecuta sustitución de compatibilidad documental.
    salida = reemplazar_emojis_problematicos(texto)

    # Verifica etiquetas de reemplazo esperadas.
    assert "Validado" in salida and "Revisión recomendada" in salida and "Incidencia detectada" in salida


# Verifica eliminación de placeholders editoriales sin resolver.
def test_sanitizar_texto_editorial_limpia_placeholders_mayusculas() -> None:
    """Comprueba que los placeholders con corchetes se conviertan a texto legible."""

    # Define texto con placeholder técnico sin resolver.
    texto = "Prioridad [ALTO_IMPACTO] y foco en [OBJETIVO]."

    # Ejecuta saneamiento editorial final.
    salida = sanear_texto_para_pdf(texto)

    # Verifica que no permanezcan corchetes ni tokens técnicos crudos.
    assert "[ALTO_IMPACTO]" not in salida and "[OBJETIVO]" not in salida


# Verifica política diferenciada de marcadores según formato de salida.
def test_reemplazar_emojis_problematicos_aplica_politica_por_formato() -> None:
    """Comprueba equivalencias editoriales distintas entre DOC y Excel."""

    # Define texto de entrada con iconografía común de estado.
    texto = "🔥 Resolver hoy ✅ Validado"

    # Ejecuta reemplazo para formato documental largo.
    salida_doc = reemplazar_emojis_problematicos(texto, formato="doc")

    # Ejecuta reemplazo para formato Excel corto.
    salida_excel = reemplazar_emojis_problematicos(texto, formato="excel")

    # Verifica equivalencias editoriales esperadas por formato.
    assert "Alta prioridad" in salida_doc and "Prioridad" in salida_excel and "OK" in salida_excel


# Verifica que la capa final bloquee placeholders residuales entre corchetes.
def test_sanitizar_texto_final_exportable_bloquea_placeholders_residuales() -> None:
    """Comprueba bloqueo final de tokens residuales tipo [TOKEN]."""

    # Define entrada con token técnico residual de plantilla.
    texto = "Plan [ALTA_PRIORIDAD] para [OBJETIVO_Q2]."

    # Aplica sanitización final para HTML.
    salida = sanitizar_texto_final_exportable(texto, formato="html")

    # Verifica ausencia de patrones técnicos en mayúsculas.
    assert re.search(r"\[[A-Z_]+\]", salida) is None


# Verifica que Word exportado no incluya tokens técnicos residuales.
def test_exportar_word_no_exporta_tokens_placeholder_residuales(tmp_path: Path) -> None:
    """Comprueba ausencia de patrones [A-Z_] en contenido DOCX final."""

    # Crea auditoría mínima con narrativa IA que incluye placeholders.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente [OBJETIVO]",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor [ALTA_PRIORIDAD]",
        resumen_ia="## Resumen ejecutivo\n- Activar [ALTA_PRIORIDAD] y [OBJETIVO_Q2].",
    )

    # Exporta documento Word a carpeta temporal.
    ruta_docx = exportar_word(auditoria, tmp_path)

    # Abre el DOCX generado para leer su contenido textual.
    contenido = "\n".join(parrafo.text for parrafo in Document(ruta_docx).paragraphs)

    # Verifica ausencia de placeholders técnicos en el documento final.
    assert re.search(r"\[[A-Z_]+\]", contenido) is None


# Verifica que Word pinte periodo y metadatos con fallback cuando no hay fechas.
def test_exportar_word_incluye_periodo_en_portada_con_fallback(tmp_path: Path) -> None:
    """Comprueba presencia de periodo y fallback de fechas en portada DOCX."""

    # Construye auditoría sin rango de fechas para forzar fallback.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente Word",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor Word",
    )

    # Exporta documento y recupera texto renderizado.
    ruta_docx = exportar_word(auditoria, tmp_path)
    texto = "\n".join(parrafo.text for parrafo in Document(ruta_docx).paragraphs)

    # Verifica bloque de periodo visible y fallback esperado.
    assert "Periodo analizado: No disponible" in texto
    assert "INFORME DE AUDITORÍA SEO" in texto


# Verifica que PDF renderice metadatos completos y periodo con fallback.
def test_exportar_pdf_incluye_bloque_meta_y_periodo_fallback(tmp_path: Path, monkeypatch) -> None:
    """Comprueba que la portada PDF incluya metadatos y periodo visible."""

    # Inicializa contenedor para capturar elementos enviados al builder.
    capturas: dict[str, object] = {}

    # Define plantilla fake para interceptar `build` sin parsear PDF binario.
    class PlantillaPdfFake:
        """Implementa interfaz mínima compatible con SimpleDocTemplate."""

        # Inicializa el destino en memoria.
        def __init__(self, *args, **kwargs) -> None:
            # Persiste argumentos para depuración.
            capturas["args"] = args
            capturas["kwargs"] = kwargs

        # Captura lista de elementos al construir el PDF.
        def build(self, elementos) -> None:
            # Guarda elementos generados por el exportador.
            capturas["elementos"] = elementos

    # Sustituye plantilla real por implementación fake.
    monkeypatch.setattr(reporters_mod, "SimpleDocTemplate", PlantillaPdfFake)

    # Construye auditoría sin fechas de periodo para forzar fallback.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente PDF",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor PDF",
    )

    # Ejecuta exportación PDF con la plantilla interceptada.
    exportar_pdf(auditoria, tmp_path)

    # Concatena textos planos de párrafos capturados.
    textos = " ".join(getattr(item, "text", "") for item in capturas.get("elementos", []))

    # Verifica presencia del periodo fallback en portada.
    assert "Periodo analizado: No disponible" in textos

    # Localiza tabla de metadatos añadida en portada.
    tabla_portada = next((item for item in capturas.get("elementos", []) if isinstance(item, Table)), None)

    # Valida existencia de tabla meta renderizada.
    assert tabla_portada is not None


# Verifica que la tabla Word aplique sanitización final en cabeceras y celdas.
def test_renderizar_tabla_word_aplica_sanitizacion_final_doc() -> None:
    """Comprueba que tablas DOCX no conserven placeholders ni emojis crudos."""

    # Crea documento Word en memoria para renderizar una tabla puntual.
    documento = Document()

    # Define tabla semántica con placeholders y emojis en cabecera y filas.
    tabla = {
        "columnas": ["Estado [ALTA_PRIORIDAD]", "Acción ✅"],
        "filas": [["Corregir [OBJETIVO_Q2]", "Aplicar 🔥 hoy"]],
    }

    # Renderiza tabla usando el helper de Word bajo prueba.
    _renderizar_tabla_word(documento, tabla)

    # Concatena texto de celdas para validaciones de contenido.
    contenido = " ".join(celda.text for fila in documento.tables[0].rows for celda in fila.cells)

    # Verifica ausencia de placeholders técnicos.
    assert re.search(r"\[[A-Z_]+\]", contenido) is None

    # Verifica equivalencias editoriales de emojis en modo doc.
    assert "Validado" in contenido and "Alta prioridad" in contenido


# Verifica consolidación por URL en hoja de contenido.
def test_construir_filas_contenido_consolidado_no_duplica_urls() -> None:
    """Comprueba que la vista de contenido tenga una fila única por URL."""

    # Crea hallazgo de muestra para forzar varias filas técnicas.
    hallazgo = HallazgoSeo(
        tipo="contenido",
        severidad="alta",
        descripcion="Falta title.",
        recomendacion="Definir title único.",
        area="Contenido",
        impacto="Alto",
        esfuerzo="Bajo",
        prioridad="P1",
    )

    # Construye URL con dos hallazgos para simular duplicado por incidencia.
    resultado_url = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="",
        h1="H1",
        meta_description="Meta",
        canonical="https://ejemplo.com/a",
        noindex=False,
        hallazgos=[hallazgo, hallazgo],
    )

    # Construye auditoría mínima de prueba.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[resultado_url],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor",
    )

    # Calcula filas de contenido consolidadas.
    filas_consolidadas = construir_filas_contenido_consolidado(auditoria)

    # Verifica que exista una sola fila por URL.
    assert len(filas_consolidadas) == 1
    assert filas_consolidadas[0]["url"] == "https://ejemplo.com/a"


# Verifica consolidación de conteos y señales en filas de contenido.
def test_construir_filas_contenido_consolidado_agrega_conteos_y_senales() -> None:
    """Comprueba que la consolidación preserve conteos agregados por URL única."""

    # Crea hallazgos de severidad y tipo mixtos para probar agregaciones.
    hallazgo_critico_contenido = HallazgoSeo(
        tipo="contenido",
        severidad="crítica",
        descripcion="Contenido muy pobre.",
        recomendacion="Ampliar cobertura temática.",
        area="Contenido",
        impacto="Alto",
        esfuerzo="Medio",
        prioridad="P1",
    )
    hallazgo_alto_tecnico = HallazgoSeo(
        tipo="indexación",
        severidad="alta",
        descripcion="Noindex no deseado.",
        recomendacion="Revisar directiva robots.",
        area="Indexación",
        impacto="Alto",
        esfuerzo="Bajo",
        prioridad="P1",
    )

    # Construye dos resultados con la misma URL para validar merge consolidado.
    resultado_url_a1 = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="",
        h1="H1",
        meta_description="",
        canonical="https://ejemplo.com/a",
        noindex=True,
        thin_content=True,
        h1_unico=False,
        estructura_headings_correcta=False,
        imagenes_sin_alt=2,
        hallazgos=[hallazgo_critico_contenido],
    )
    resultado_url_a2 = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="Título",
        h1="H1",
        meta_description="Meta",
        canonical="https://ejemplo.com/a",
        noindex=False,
        hallazgos=[hallazgo_alto_tecnico],
    )

    # Construye auditoría con duplicidad de URL en orígen.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=2,
        resultados=[resultado_url_a1, resultado_url_a2],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor",
    )

    # Calcula filas consolidadas y obtiene la URL objetivo.
    filas_consolidadas = construir_filas_contenido_consolidado(auditoria)
    fila_a = next(fila for fila in filas_consolidadas if fila["url"] == "https://ejemplo.com/a")

    # Verifica agregación de conteos y áreas afectadas.
    assert fila_a["incidencias_url"] == 2
    assert fila_a["incidencias_criticas_altas"] == 2
    assert fila_a["incidencias_contenido"] == 1
    assert fila_a["areas_con_incidencias"] == "Contenido, Indexación"
    assert "thin_content" in fila_a["senales_clave"]
    assert "title_vacio" in fila_a["senales_clave"]


# Verifica que la hoja Contenido del Excel mantenga URL única y conteos coherentes.
def test_exportar_excel_contenido_unico_y_conteos_consistentes(tmp_path: Path) -> None:
    """Comprueba unicidad de URL en Contenido y consistencia contra Errores."""

    # Crea dos hallazgos para una misma URL.
    hallazgo_a1 = HallazgoSeo(
        tipo="contenido",
        severidad="alta",
        descripcion="Falta title.",
        recomendacion="Definir title único.",
        area="Contenido",
        impacto="Alto",
        esfuerzo="Bajo",
        prioridad="P1",
    )
    hallazgo_a2 = HallazgoSeo(
        tipo="indexación",
        severidad="media",
        descripcion="Canonical no coherente.",
        recomendacion="Corregir canonical.",
        area="Indexación",
        impacto="Medio",
        esfuerzo="Bajo",
        prioridad="P2",
    )

    # Construye resultados de entrada con dos URLs y una duplicada.
    resultado_a = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="A",
        h1="A",
        meta_description="A",
        canonical="https://ejemplo.com/a",
        noindex=False,
        hallazgos=[hallazgo_a1, hallazgo_a2],
    )
    resultado_b = ResultadoUrl(
        url="https://ejemplo.com/b",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/b",
        title="B",
        h1="B",
        meta_description="B",
        canonical="https://ejemplo.com/b",
        noindex=False,
        hallazgos=[],
    )
    resultado_a_duplicado = ResultadoUrl(
        url="https://ejemplo.com/a",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/a",
        title="A",
        h1="A",
        meta_description="A",
        canonical="https://ejemplo.com/a",
        noindex=False,
        hallazgos=[],
    )

    # Construye auditoría y exporta Excel.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=3,
        resultados=[resultado_a, resultado_b, resultado_a_duplicado],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor",
    )
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro y resuelve índices de columnas por nombre.
    libro = load_workbook(ruta_excel)
    hoja_contenido = libro["Contenido"]
    cabeceras_contenido = [celda.value for celda in hoja_contenido[1]]
    indice_url_contenido = cabeceras_contenido.index("url")
    indice_incidencias_url = cabeceras_contenido.index("incidencias_url")

    # Lee URLs y conteos de la hoja de contenido.
    filas_contenido = [fila for fila in hoja_contenido.iter_rows(min_row=2, values_only=True) if fila[indice_url_contenido]]
    urls_contenido = [fila[indice_url_contenido] for fila in filas_contenido]
    conteos_contenido = {fila[indice_url_contenido]: int(fila[indice_incidencias_url] or 0) for fila in filas_contenido}

    # Verifica unicidad estricta por URL en la hoja consolidada.
    assert len(urls_contenido) == len(set(urls_contenido))
    assert conteos_contenido["https://ejemplo.com/a"] == 2
    assert conteos_contenido["https://ejemplo.com/b"] == 0

    # Obtiene conteos desde la hoja de errores para validar consistencia.
    hoja_errores = libro["Errores"]
    cabeceras_errores = [celda.value for celda in hoja_errores[1]]
    indice_url_errores = cabeceras_errores.index("url")
    indice_problema_errores = cabeceras_errores.index("problema")
    conteos_errores: dict[str, int] = {}
    for fila in hoja_errores.iter_rows(min_row=2, values_only=True):
        url = fila[indice_url_errores]
        if not url:
            continue
        problema = str(fila[indice_problema_errores] or "").strip()
        if problema:
            conteos_errores[url] = conteos_errores.get(url, 0) + 1
        else:
            conteos_errores.setdefault(url, 0)

    # Verifica coherencia entre conteos de Contenido y Errores.
    assert conteos_contenido == conteos_errores


# Verifica que la fusión de campos escalares sea determinista e independiente del orden.
def test_construir_filas_contenido_consolidado_fusiona_escalares_de_forma_determinista() -> None:
    """Comprueba que la consolidación use estrategia explícita para campos escalares."""

    # Crea hallazgo para conservar conteo de incidencias durante la fusión.
    hallazgo = HallazgoSeo(
        tipo="contenido",
        severidad="alta",
        descripcion="Falta profundidad.",
        recomendacion="Ampliar texto.",
        area="Contenido",
        impacto="Alto",
        esfuerzo="Medio",
        prioridad="P1",
    )

    # Construye variante incompleta de la misma URL.
    incompleto = ResultadoUrl(
        url="https://ejemplo.com/c",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/c",
        title="",
        h1="",
        meta_description="",
        canonical="https://ejemplo.com/c",
        noindex=False,
        palabras=120,
        calidad_contenido="media",
        thin_content=True,
        densidad_texto=0.21,
        ratio_texto_html=0.15,
        h1_unico=False,
        estructura_headings_correcta=False,
        imagenes_sin_alt=3,
        lazy_load_detectado=False,
        hallazgos=[hallazgo],
    )

    # Construye variante completa de la misma URL.
    completo = ResultadoUrl(
        url="https://ejemplo.com/c",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com/c",
        title="Guía completa de producto",
        h1="Guía de producto",
        meta_description="Descripción extensa para producto clave",
        canonical="https://ejemplo.com/c",
        noindex=True,
        palabras=520,
        calidad_contenido="alta",
        thin_content=False,
        densidad_texto=0.55,
        ratio_texto_html=0.34,
        h1_unico=True,
        estructura_headings_correcta=True,
        imagenes_sin_alt=1,
        lazy_load_detectado=True,
        hallazgos=[],
    )

    # Ejecuta consolidación con orden A.
    auditoria_a = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=2,
        resultados=[incompleto, completo],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-30",
        gestor="Gestor",
    )
    fila_a = construir_filas_contenido_consolidado(auditoria_a)[0]

    # Ejecuta consolidación con orden inverso para validar determinismo.
    auditoria_b = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=2,
        resultados=[completo, incompleto],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-30",
        gestor="Gestor",
    )
    fila_b = construir_filas_contenido_consolidado(auditoria_b)[0]

    # Verifica igualdad completa de filas sin depender del orden de entrada.
    assert fila_a == fila_b

    # Verifica reglas explícitas de fusión escalar.
    assert fila_a["title"] == "Guía completa de producto"
    assert fila_a["h1"] == "Guía de producto"
    assert fila_a["meta_description"] == "Descripción extensa para producto clave"
    assert fila_a["palabras"] == 520
    assert fila_a["calidad_contenido"] == "media"
    assert fila_a["imagenes_sin_alt"] == 3
    assert fila_a["thin_content"] == "Sí"
    assert fila_a["noindex"] == "Sí"
    assert fila_a["h1_unico"] == "No"
    assert fila_a["estructura_headings_correcta"] == "No"
    assert fila_a["lazy_load_detectado"] == "Sí"


# Verifica trazabilidad de score de páginas prioritarias por componentes.
def test_calcular_score_prioridad_pagina_devuelve_componentes() -> None:
    """Comprueba que el cálculo de prioridad exponga desglose explicable."""

    # Construye estructuras mínimas de prueba para señales.
    class Obj:
        pass

    # Crea bloque mínimo GSC con oportunidad de CTR y posición.
    gsc = Obj()
    gsc.impresiones = 900.0
    gsc.ctr = 0.01
    gsc.posicion_media = 7.5

    # Crea bloque mínimo GA con sesiones sin conversión.
    ga = Obj()
    ga.sesiones = 150.0
    ga.conversiones = 0.0
    ga.rebote = 0.7

    # Crea bloque técnico con varios hallazgos.
    tecnico = Obj()
    tecnico.hallazgos = [1, 2, 3]

    # Ejecuta cálculo explicable de prioridad.
    evaluacion = calcular_score_prioridad_pagina(gsc=gsc, ga=ga, tecnico=tecnico)

    # Verifica score positivo y motivos presentes.
    assert float(evaluacion["score_prioridad"]) > 0
    assert len(evaluacion["motivos"]) >= 2
    assert "oportunidad_ctr" in evaluacion["componentes"]


# Verifica que el modelo semántico contenga tablas clave para alineación.
def test_construir_modelo_semantico_informe_incluye_tablas_clave() -> None:
    """Comprueba presencia de secciones tabulares en modelo neutral."""

    # Construye auditoría mínima para modelo semántico.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
    )

    # Construye modelo semántico reusable.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Resuelve títulos de secciones del modelo.
    titulos = {seccion["titulo"] for seccion in modelo["secciones"]}

    # Verifica secciones estructurales mínimas.
    assert "KPIs principales" in titulos
    assert "Anexo técnico" in titulos


# Verifica estructura semántica mínima con bloques explícitos obligatorios.
def test_construir_modelo_semantico_informe_bloques_explicitos_minimos() -> None:
    """Comprueba que el modelo publique bloques nuevos y secciones esenciales."""

    # Construye auditoría mínima para validar contrato semántico.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
    )

    # Genera modelo semántico del informe.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Verifica que metadatos explícitos del nuevo contrato existan.
    assert "metadatos" in modelo
    assert "periodo_texto" in modelo["metadatos"]

    # Localiza secciones obligatorias del informe.
    seccion_resumen = next((seccion for seccion in modelo["secciones"] if seccion["titulo"] == "Resumen ejecutivo"), None)
    seccion_kpis = next((seccion for seccion in modelo["secciones"] if seccion["titulo"] == "KPIs principales"), None)

    # Comprueba que las secciones obligatorias existan.
    assert seccion_resumen is not None
    assert seccion_kpis is not None

    # Verifica bloques semánticos nuevos mínimos.
    assert isinstance(seccion_resumen.get("resumen_ejecutivo", []), list)
    assert len(seccion_kpis.get("kpi_cards", [])) >= 3
    assert len(seccion_kpis.get("tablas_detalle", [])) >= 1


# Verifica compatibilidad retroactiva de bloques legacy frente al nuevo contrato.
def test_construir_modelo_semantico_informe_mantiene_compatibilidad_legacy() -> None:
    """Comprueba que se mantengan parrafos, tablas y tarjetas en paralelo."""

    # Construye auditoría mínima para validar compatibilidad.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
    )

    # Genera modelo semántico bajo prueba.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Recupera sección de KPI para comparar claves nuevas y legacy.
    seccion_kpis = next((seccion for seccion in modelo["secciones"] if seccion["titulo"] == "KPIs principales"), {})

    # Verifica coexistencia de bloques legacy y nuevos.
    assert len(seccion_kpis.get("tablas", [])) >= 1
    assert len(seccion_kpis.get("tablas_detalle", [])) >= 1
    assert len(seccion_kpis.get("kpi_cards", [])) >= 1


# Verifica que resumen_ejecutivo solo se use en la sección de resumen.
def test_construir_modelo_semantico_informe_limita_resumen_ejecutivo_a_seccion_resumen() -> None:
    """Comprueba integridad semántica del bloque resumen_ejecutivo."""

    # Construye auditoría mínima para validar contrato semántico.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
    )

    # Genera modelo semántico bajo prueba.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Recorre secciones para validar asignación selectiva.
    for seccion in modelo["secciones"]:
        if seccion["titulo"] == "Resumen ejecutivo":
            # Verifica presencia de bloque resumen en sección específica.
            assert isinstance(seccion.get("resumen_ejecutivo", []), list)
            continue

        # Verifica que el resto no reciba resumen ejecutivo por defecto.
        assert seccion.get("resumen_ejecutivo", []) == []


# Verifica bloque meta completo para exportadores documentales.
def test_construir_modelo_semantico_informe_incluye_meta_completo() -> None:
    """Comprueba que el modelo semántico publique metadatos completos."""

    # Crea auditoría con periodo explícito.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente Demo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor Demo",
        periodo_date_from="2026-02-01",
        periodo_date_to="2026-02-29",
    )

    # Genera modelo semántico bajo prueba.
    modelo = construir_modelo_semantico_informe(auditoria)
    meta = modelo["meta"]

    # Valida claves obligatorias del contrato meta.
    assert meta["cliente"] == "Cliente Demo"
    assert meta["gestor"] == "Gestor Demo"
    assert meta["fecha_ejecucion"] == "2026-03-20"
    assert meta["periodo_desde"] == "2026-02-01"
    assert meta["periodo_hasta"] == "2026-02-29"
    assert meta["sitemap"] == "https://ejemplo.com/sitemap.xml"


# Verifica fallback de periodo cuando no hay fechas en fuentes ni periodo persistido.
def test_construir_modelo_semantico_informe_periodo_fallback_sin_fechas() -> None:
    """Comprueba que el periodo se marque como no disponible sin fechas origen."""

    # Construye auditoría mínima sin periodo explícito.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente Demo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor Demo",
    )

    # Genera modelo semántico y obtiene bloque meta.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Valida fallback de periodo textual y extremos vacíos.
    assert modelo["meta"]["periodo_desde"] == ""
    assert modelo["meta"]["periodo_hasta"] == ""
    assert modelo["meta"]["periodo_texto"] == "No disponible"


# Verifica que rendimiento conserve tabla detallada y oportunidades.
def test_modelo_semantico_rendimiento_incluye_metricas_y_oportunidades() -> None:
    """Comprueba que no se pierda el detalle de PageSpeed en la capa semántica."""

    # Construye oportunidad de rendimiento de ejemplo.
    oportunidad = OportunidadRendimiento(
        id_oportunidad="op_1",
        titulo="Optimizar imágenes",
        descripcion="Reducir peso de recursos",
        ahorro_estimado="300 ms",
        severidad="alta",
    )

    # Construye resultado de rendimiento con métricas y oportunidad.
    rendimiento = ResultadoRendimiento(
        url="https://ejemplo.com",
        estrategia="mobile",
        performance_score=55.0,
        accessibility_score=90.0,
        best_practices_score=92.0,
        seo_score=88.0,
        lcp="2.5 s",
        cls="0.11",
        inp="180 ms",
        fcp="1.4 s",
        tbt="120 ms",
        speed_index="2.2 s",
        campo_lcp=None,
        campo_cls=None,
        campo_inp=None,
        oportunidades=[oportunidad],
    )

    # Construye auditoría mínima con bloque de rendimiento activo.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        rendimiento=[rendimiento],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-26",
        gestor="Gestor",
    )

    # Genera modelo semántico del informe.
    modelo = construir_modelo_semantico_informe(auditoria)

    # Obtiene sección de rendimiento del modelo.
    seccion_rendimiento = next((seccion for seccion in modelo["secciones"] if seccion["titulo"] == "Rendimiento y experiencia de usuario"), None)

    # Verifica existencia de sección y tablas esperadas.
    assert seccion_rendimiento is not None
    assert any(tabla.get("titulo") == "Rendimiento por métrica" for tabla in seccion_rendimiento.get("tablas", []))
    assert any(tabla.get("titulo") == "Oportunidades PageSpeed priorizadas" for tabla in seccion_rendimiento.get("tablas", []))


# Verifica que la media de score de dashboard use ejecuciones únicas.
def test_exportar_excel_score_medio_desde_ejecuciones_unicas(tmp_path: Path) -> None:
    """Comprueba que KPI de score medio no se sesgue por número de oportunidades."""

    # Crea resultado URL mínimo para construir auditoría válida.
    resultado_url = ResultadoUrl(
        url="https://ejemplo.com",
        tipo="page",
        estado_http=200,
        redirecciona=False,
        url_final="https://ejemplo.com",
        title="Home",
        h1="Inicio",
        meta_description="Meta",
        canonical="https://ejemplo.com",
        noindex=False,
        hallazgos=[],
    )

    # Crea oportunidades duplicadas en un mismo análisis móvil.
    oportunidades = [
        OportunidadRendimiento(id_oportunidad="op1", titulo="Reducir JS", descripcion="...", ahorro_estimado="500 ms", severidad="alta"),
        OportunidadRendimiento(id_oportunidad="op2", titulo="Optimizar imágenes", descripcion="...", ahorro_estimado="300 ms", severidad="media"),
    ]

    # Construye resultados de rendimiento por estrategia.
    rendimiento = [
        ResultadoRendimiento(
            url="https://ejemplo.com",
            estrategia="mobile",
            performance_score=50.0,
            accessibility_score=90.0,
            best_practices_score=90.0,
            seo_score=90.0,
            lcp="2,5 s",
            cls="0,10",
            inp="200 ms",
            fcp="1,2 s",
            tbt="100 ms",
            speed_index="2,0 s",
            campo_lcp=None,
            campo_cls=None,
            campo_inp=None,
            oportunidades=oportunidades,
        ),
        ResultadoRendimiento(
            url="https://ejemplo.com",
            estrategia="mobile",
            performance_score=100.0,
            accessibility_score=95.0,
            best_practices_score=95.0,
            seo_score=95.0,
            lcp="1,0 s",
            cls="0,01",
            inp="100 ms",
            fcp="0,8 s",
            tbt="0 ms",
            speed_index="1,0 s",
            campo_lcp=None,
            campo_cls=None,
            campo_inp=None,
            oportunidades=[],
        ),
    ]

    # Construye auditoría con rendimiento para exportación.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[resultado_url],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
        rendimiento=rendimiento,
        gestion_indexacion=[
            DecisionIndexacion(
                url="https://ejemplo.com",
                clasificacion="INDEXABLE",
                motivo="Sin señales de riesgo detectadas",
                accion_recomendada="Mantener monitorización",
                prioridad="Baja",
            )
        ],
    )

    # Exporta el Excel de prueba en carpeta temporal.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Abre libro generado para validación de KPI.
    libro = load_workbook(ruta_excel)

    # Obtiene hoja KPIs para validar el bloque ejecutivo mínimo.
    hoja_kpis = libro["KPIs"]

    # Verifica que exista hoja específica de indexación.
    assert "Indexacion" in libro.sheetnames

    # Busca fila del KPI de score rendimiento para evitar acoplamiento por posición.
    fila_score_rend = next((fila for fila in range(5, 60) if hoja_kpis[f"A{fila}"].value == "Score rendimiento"), None)

    # Verifica que la fila del KPI exista en la hoja de KPIs.
    assert fila_score_rend is not None

    # Verifica score rendimiento basado en ejecuciones únicas: (50 + 100) / 2 = 75.
    assert hoja_kpis[f"B{fila_score_rend}"].value == 75.0


# Verifica que Excel muestre periodo en KPIs y Dashboard.
def test_exportar_excel_incluye_periodo_en_kpis_y_dashboard(tmp_path: Path) -> None:
    """Comprueba visibilidad del periodo en cabeceras ejecutivas de Excel."""

    # Construye auditoría con periodo explícito.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente Excel",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor Excel",
        periodo_date_from="2026-03-01",
        periodo_date_to="2026-03-28",
    )

    # Exporta archivo Excel y abre hojas objetivo.
    ruta_excel = exportar_excel(auditoria, tmp_path)
    libro = load_workbook(ruta_excel)
    hoja_kpis = libro["KPIs"]
    hoja_dashboard = libro["Dashboard"]

    # Valida presencia del periodo en KPI superior y subcabecera dashboard.
    assert "Periodo: 2026-03-01 - 2026-03-28" in str(hoja_kpis["A2"].value or "")
    assert "Periodo: 2026-03-01 - 2026-03-28" in str(hoja_dashboard["A2"].value or "")


# Verifica fallback de periodo en cabeceras de Excel sin fechas disponibles.
def test_exportar_excel_periodo_fallback_sin_fechas(tmp_path: Path) -> None:
    """Comprueba que Excel use 'No disponible' cuando no hay periodo."""

    # Construye auditoría sin fechas de periodo.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Cliente Excel",
        fecha_ejecucion="2026-03-29",
        gestor="Gestor Excel",
    )

    # Exporta y carga libro para validación.
    ruta_excel = exportar_excel(auditoria, tmp_path)
    libro = load_workbook(ruta_excel)

    # Verifica fallback uniforme en KPIs y Dashboard.
    assert "Periodo: No disponible" in str(libro["KPIs"]["A2"].value or "")
    assert "Periodo: No disponible" in str(libro["Dashboard"]["A2"].value or "")


# Verifica que el dashboard conserve los gráficos esperados.
def test_exportar_excel_dashboard_contiene_graficos(tmp_path: Path) -> None:
    """Comprueba que el Dashboard exportado tenga gráficos visibles y cargados."""

    # Construye auditoría mínima para exportar Excel.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[
            ResultadoUrl(
                url="https://ejemplo.com",
                tipo="page",
                estado_http=200,
                redirecciona=False,
                url_final="https://ejemplo.com",
                title="Home",
                h1="Inicio",
                meta_description="Meta",
                canonical="https://ejemplo.com",
                noindex=False,
                hallazgos=[],
            )
        ],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
    )

    # Exporta libro para inspeccionar gráficos.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro recién exportado.
    libro = load_workbook(ruta_excel)

    # Obtiene hoja dashboard para validar gráficos.
    dashboard = libro["Dashboard"]

    # Verifica que existan al menos tres gráficos.
    assert len(dashboard._charts) >= 3


# Verifica que Excel incluya hojas nuevas de Search Console.
def test_exportar_excel_incluye_hojas_gsc(tmp_path: Path) -> None:
    """Comprueba que la exportación Excel cree hojas GSC aunque no haya datos."""

    # Construye auditoría mínima para exportar Excel.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[
            ResultadoUrl(
                url="https://ejemplo.com",
                tipo="page",
                estado_http=200,
                redirecciona=False,
                url_final="https://ejemplo.com",
                title="Home",
                h1="Inicio",
                meta_description="Meta",
                canonical="https://ejemplo.com",
                noindex=False,
                hallazgos=[],
            )
        ],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-24",
        gestor="Gestor",
    )

    # Exporta libro para inspeccionar estructura.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro recién exportado.
    libro = load_workbook(ruta_excel)

    # Verifica que existan hojas nuevas de Search Console.
    assert libro.sheetnames[0] == "KPIs"
    assert libro.sheetnames[1] == "Dashboard"
    assert "Search_Console_Paginas" in libro.sheetnames
    assert "Search_Console_Queries" in libro.sheetnames
    assert "Oportunidades_GSC" in libro.sheetnames
    assert "Analytics" in libro.sheetnames


# Verifica que el dashboard tenga congelación de panel y bloques ejecutivos.
def test_exportar_excel_dashboard_mejorado_legible(tmp_path: Path) -> None:
    """Comprueba mejoras visuales base del dashboard y navegación congelada."""

    # Construye auditoría mínima para validar dashboard.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-25",
        gestor="Gestor",
    )

    # Exporta Excel para inspección de dashboard.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro exportado.
    libro = load_workbook(ruta_excel)

    # Obtiene hoja principal.
    dashboard = libro["Dashboard"]

    # Verifica congelación de panel para navegación de KPIs.
    assert dashboard.freeze_panes == "A7"

    # Verifica existencia de bloque principal de score global + bloques.
    assert dashboard["A4"].value == "Score global y score por bloques"


# Verifica que la primera pantalla priorice solo bloques ejecutivos clave.
def test_exportar_excel_dashboard_primera_pantalla_con_bloques_clave(tmp_path: Path) -> None:
    """Comprueba presencia de bloques ancla y orden de lectura en la rejilla A-F / G-L."""

    # Construye auditoría mínima para validar layout ejecutivo.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-30",
        gestor="Gestor",
    )

    # Exporta Excel y abre dashboard.
    ruta_excel = exportar_excel(auditoria, tmp_path)
    dashboard = load_workbook(ruta_excel)["Dashboard"]

    # Verifica bloques clave en celdas ancla de primera pantalla.
    assert dashboard["A4"].value == "Score global y score por bloques"
    assert dashboard["G4"].value == "Severidades"
    assert dashboard["A14"].value == "Oportunidades principales"
    assert dashboard["G14"].value == "Top páginas prioritarias"

    # Valida orden de lectura ejecutiva (fila superior antes que fila inferior).
    assert dashboard["A4"].row < dashboard["A14"].row
    assert dashboard["G4"].row < dashboard["G14"].row


# Verifica que el detalle secundario quede desplazado fuera de la primera pantalla.
def test_exportar_excel_dashboard_mueve_detalle_secundario_a_bloques_inferiores(tmp_path: Path) -> None:
    """Comprueba que los bloques secundarios arranquen por debajo de la zona ejecutiva inicial."""

    # Construye auditoría mínima para validar estructura.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-30",
        gestor="Gestor",
    )

    # Exporta Excel y abre dashboard.
    ruta_excel = exportar_excel(auditoria, tmp_path)
    dashboard = load_workbook(ruta_excel)["Dashboard"]

    # Valida que los bloques secundarios estén debajo de la primera pantalla.
    assert dashboard["A24"].value == "Visibilidad orgánica real"
    assert dashboard["G24"].value == "Gestión de indexación"
    assert dashboard["A31"].value == "Comportamiento Analytics"


# Verifica retorno coherente de bloque dashboard cuando no hay líneas.
def test_renderizar_bloque_dashboard_retorna_fila_cabecera_si_esta_vacio() -> None:
    """Asegura que la fila final reportada sea la cabecera cuando no hay contenido."""

    # Crea libro temporal para probar helper visual.
    libro = Workbook()

    # Obtiene hoja activa para renderizar el bloque.
    hoja = libro.active

    # Renderiza bloque sin líneas de detalle.
    fila_final = _renderizar_bloque_dashboard(hoja, "D9", "Bloque vacío", [], "1D4ED8")

    # Verifica que se devuelva la fila de cabecera pintada.
    assert fila_final == 9


# Verifica que la hoja de errores mantenga color por severidad.
def test_exportar_excel_aplica_color_por_severidad(tmp_path: Path) -> None:
    """Comprueba que la fila de error reciba un color suave según severidad."""

    # Crea hallazgo de severidad crítica para validar color.
    hallazgo = HallazgoSeo(
        tipo="indexación",
        severidad="crítica",
        descripcion="Error 5xx.",
        recomendacion="Corregir servidor.",
        area="Infraestructura",
        impacto="Muy alto",
        esfuerzo="Medio",
        prioridad="P1",
    )

    # Construye resultado URL con incidencia crítica.
    resultado_url = ResultadoUrl(
        url="https://ejemplo.com",
        tipo="page",
        estado_http=500,
        redirecciona=False,
        url_final="https://ejemplo.com",
        title="Home",
        h1="Inicio",
        meta_description="Meta",
        canonical="https://ejemplo.com",
        noindex=False,
        hallazgos=[hallazgo],
    )

    # Construye auditoría para exportación.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[resultado_url],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
    )

    # Exporta Excel con datos críticos.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro exportado.
    libro = load_workbook(ruta_excel)

    # Obtiene hoja de errores.
    errores = libro["Errores"]

    # Obtiene color de fondo de la primera fila de datos.
    color = errores["A2"].fill.fgColor.rgb or ""

    # Verifica que se haya aplicado un color distinto a blanco.
    assert "FADBD8" in color


# Verifica que la sección de rendimiento no use valores None como datos reales.
def test_construir_bloques_narrativos_rendimiento_fallido_muestra_mensaje_profesional() -> None:
    """Comprueba que, ante fallo de PageSpeed, se renderice mensaje profesional y no valores vacíos."""

    # Construye auditoría mínima con estado fallido de PageSpeed.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
        fuentes_fallidas=["pagespeed"],
        pagespeed_estado={"https://ejemplo.com/": {"mobile": "timeout"}},
    )

    # Genera bloques narrativos.
    bloques = _construir_bloques_narrativos(auditoria)

    # Consolida sección de rendimiento.
    texto = " ".join(bloques["Rendimiento y experiencia de usuario"])

    # Verifica mensaje profesional de indisponibilidad.
    assert "No se pudieron obtener métricas de PageSpeed" in texto


# Verifica que la tabla de rendimiento se cree con rango válido.
def test_exportar_excel_tabla_rendimiento_valida(tmp_path: Path) -> None:
    """Comprueba que la tabla de rendimiento exista y tenga rango consistente."""

    # Construye auditoría con una fila de rendimiento para crear tabla.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-20",
        gestor="Gestor",
        rendimiento=[
            ResultadoRendimiento(
                url="https://ejemplo.com",
                estrategia="mobile",
                performance_score=80.0,
                accessibility_score=90.0,
                best_practices_score=90.0,
                seo_score=95.0,
                lcp="2,0 s",
                cls="0,05",
                inp="200 ms",
                fcp="1,0 s",
                tbt="80 ms",
                speed_index="1,5 s",
                campo_lcp=None,
                campo_cls=None,
                campo_inp=None,
            )
        ],
    )

    # Exporta Excel para inspección de tabla.
    ruta_excel = exportar_excel(auditoria, tmp_path)

    # Carga libro generado.
    libro = load_workbook(ruta_excel)

    # Obtiene hoja de rendimiento.
    rendimiento = libro["Rendimiento"]

    # Verifica existencia de tabla y rango válido esperado.
    assert "TablaRendimiento" in rendimiento.tables
    assert rendimiento.tables["TablaRendimiento"].ref == "A1:T2"


# Verifica que Word exporte correctamente la sección de comportamiento y conversión.
def test_exportar_word_comportamiento_conversion_sin_error(tmp_path: Path) -> None:
    """Comprueba que el DOCX se genere sin errores con Analytics activo."""

    # Construye auditoría con Analytics para activar sección de comportamiento.
    auditoria = ResultadoAuditoria(
        sitemap="https://ejemplo.com/sitemap.xml",
        total_urls=1,
        resultados=[],
        cliente="Ejemplo",
        fecha_ejecucion="2026-03-26",
        gestor="Gestor",
        analytics=DatosAnalytics(
            activo=True,
            paginas=[
                MetricaAnalyticsPagina(
                    url="/",
                    sesiones=120,
                    usuarios=95,
                    rebote=0.61,
                    duracion_media=84,
                    conversiones=3,
                    calidad_trafico="media",
                )
            ],
        ),
    )

    # Exporta archivo Word en carpeta temporal.
    ruta_word = exportar_word(auditoria, tmp_path)

    # Verifica que el archivo DOCX exista en disco.
    assert ruta_word.exists()


# Verifica que el cálculo de anchos PDF respete siempre el ancho útil disponible.
def test_calcular_col_widths_pdf_respeta_ancho_util_a4() -> None:
    """Comprueba que la suma de anchos coincida con A4 menos márgenes."""

    # Define ancho útil estándar de exportación PDF.
    ancho_util = float(A4[0] - PDF_HORIZONTAL_MARGIN_POINTS)

    # Construye tabla con columna narrativa extensa.
    columnas = ["URL", "Score", "Motivos", "Recomendación"]
    filas = [["https://ejemplo.com/a", "88", "Motivo extenso " * 4, "Recomendación extensa " * 5]]

    # Calcula anchos para tabla de prueba.
    anchos = _calcular_col_widths_pdf(columnas, filas, ancho_util=ancho_util)

    # Verifica cardinalidad y ajuste exacto al ancho utilizable.
    assert len(anchos) == len(columnas)
    assert round(sum(anchos), 3) == round(ancho_util, 3)


# Verifica que tablas anchas se dividan por semántica en bloques verticales.
def test_resolver_subtablas_pdf_divide_tablas_anchas() -> None:
    """Comprueba que se reduzca el ancho horizontal extremo por subtablas."""

    # Define tabla semántica ancha de páginas prioritarias.
    tabla = {
        "titulo": "Páginas prioritarias",
        "columnas": ["URL", "Score", "Impresiones", "CTR", "Sesiones", "Conversiones", "Motivos"],
        "filas": [["https://ejemplo.com/a", "80", "1000", "0.02", "120", "1", "Detalle narrativo muy largo"]],
    }

    # Resuelve subtablas para render PDF.
    subtables = _resolver_subtablas_pdf(tabla)

    # Verifica partición en dos bloques con menos columnas por bloque.
    assert len(subtables) == 2
    assert all(len(subtabla["columnas"]) <= 4 for subtabla in subtables)


# Verifica que el render PDF use ancho explícito y preserve wrap en celdas narrativas.
def test_renderizar_tabla_pdf_configura_col_widths_en_tablas_criticas() -> None:
    """Comprueba configuración explícita de ancho en tablas críticas."""

    # Construye tabla de rendimiento ancha para forzar subtablas.
    tabla = {
        "titulo": "Rendimiento por métrica",
        "columnas": ["URL / Estrategia", "Métrica", "Valor", "Observación"],
        "filas": [["https://ejemplo.com [mobile]", "LCP", "2.1 s", "Observación extensa " * 10]],
    }

    # Renderiza elementos PDF de la tabla.
    elementos = _renderizar_tabla_pdf(tabla, getSampleStyleSheet())

    # Filtra tablas reales del bloque renderizado.
    tablas = [elemento for elemento in elementos if isinstance(elemento, Table)]

    # Verifica que existan subtablas y todas usen colWidths explícito.
    assert len(tablas) == 2
    assert all(tabla_render._colWidths for tabla_render in tablas)
    assert max(len(tabla_render._colWidths) for tabla_render in tablas) <= 3
